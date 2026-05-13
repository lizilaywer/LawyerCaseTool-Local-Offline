# -*- coding: utf-8 -*-
"""法律工具中心对话框。"""

from __future__ import annotations

from datetime import date, datetime
from pathlib import Path
from typing import Callable, Dict, List, Optional, Tuple

from PySide6.QtCore import QDate, QEvent, QObject, QStringListModel, Qt, QUrl, Signal, QThread, QMarginsF
from PySide6.QtGui import QDesktopServices, QTextDocument
from PySide6.QtWidgets import (
    QApplication,
    QCalendarWidget,
    QCheckBox,
    QComboBox,
    QCompleter,
    QDateEdit,
    QDialog,
    QDoubleSpinBox,
    QFileDialog,
    QFormLayout,
    QFrame,
    QGridLayout,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMessageBox,
    QPushButton,
    QScrollArea,
    QSpinBox,
    QTabWidget,
    QTextEdit,
    QTextBrowser,
    QToolButton,
    QTreeWidget,
    QTreeWidgetItem,
    QVBoxLayout,
    QWidget,
)

from src.core.case_manager import CaseManager
try:
    from src.core.court_sms_service import (
        CourtSmsCaseMatch,
        CourtSmsDocument,
        CourtSmsHearingNotice,
        CourtSmsParseResult,
        CourtSmsService,
    )
    _COURT_SMS_IMPORT_ERROR: Optional[ModuleNotFoundError] = None
except ModuleNotFoundError as exc:
    CourtSmsCaseMatch = CourtSmsDocument = CourtSmsHearingNotice = CourtSmsParseResult = object
    CourtSmsService = None  # type: ignore[assignment]
    _COURT_SMS_IMPORT_ERROR = exc
from src.core.screenshot_pdf_merger import ScreenshotPdfMerger
from src.gui.widgets.screenshot_image_list import ScreenshotImageList
from src.core.legal_toolkit import (
    COMMON_CIVIL_CAUSES,
    LEGAL_REFERENCE_LINKS,
    PROCEDURAL_LIMIT_RULES,
    calculate_bankruptcy_administrator_fee,
    calculate_date_offset,
    calculate_delay_performance_interest,
    calculate_divorce_litigation_fee,
    calculate_execution_fee,
    calculate_labor_compensation,
    calculate_lawyer_fee,
    calculate_liquidated_damages,
    calculate_occupation_fee,
    calculate_personality_litigation_fee,
    calculate_preservation_fee,
    calculate_procedural_deadline,
    calculate_property_litigation_fee,
    calculate_simple_interest,
    calculate_traffic_injury_compensation,
    calculate_work_injury_death,
    calculate_work_injury_disability,
    compensation_years_for_age,
    money,
)
from src.core.lpr_data import calculate_lpr_interest, calculate_interest_days, get_lpr_manager
from src.gui.styles import APP_COLORS as COLORS, CHECK_ICON_PATH
from src.gui.widgets.docx_compare_widget import DocxCompareWidget
from src.gui.widgets.docx_auto_format_widget import DocxAutoFormatWidget
from src.gui.window_metrics import APP_SURFACE_DEFAULT_SIZE, APP_SURFACE_MIN_SIZE
from src.gui.widgets.transparent_form_layout import TransparentFormLayout


class BlankZeroMoneySpinBox(QDoubleSpinBox):
    """金额输入框：默认 0 显示为空，便于直接输入。"""

    def textFromValue(self, value: float) -> str:
        if abs(value - self.minimum()) < 10 ** (-max(0, self.decimals())):
            return ""
        return super().textFromValue(value)

    def valueFromText(self, text: str) -> float:
        if not str(text or "").strip():
            return self.minimum()
        return super().valueFromText(text)


class _CourtSmsReadWorker(QThread):
    """在后台线程执行法院短信解析、读取和下载。"""

    finished = Signal(object, object, object, object)  # parsed, documents, download_dir, hearing_notices
    error = Signal(str)

    def __init__(self, service, sms_text: str, parsed_result=None):
        super().__init__()
        self.service = service
        self.sms_text = sms_text
        self.parsed_result = parsed_result

    def run(self):
        try:
            parsed = self.parsed_result or self.service.parse_sms(self.sms_text)
            documents = self.service.fetch_documents(parsed)
            download_dir = self.service.download_documents(documents, parsed)
            hearing_notices = self.service.extract_hearing_notices(documents, parsed)
            self.finished.emit(parsed, documents, download_dir, hearing_notices)
        except Exception as exc:
            self.error.emit(str(exc))


class _CourtSmsAnalyzeWorker(QThread):
    """在后台线程执行庭审通知解析。"""

    finished = Signal(object, object, object, object)
    error = Signal(str)

    def __init__(self, service, documents, parse_result):
        super().__init__()
        self.service = service
        self.documents = documents
        self.parse_result = parse_result

    def run(self):
        try:
            hearing_notices = self.service.extract_hearing_notices(self.documents, self.parse_result)
            self.finished.emit(self.parse_result, self.documents, None, hearing_notices)
        except Exception as exc:
            self.error.emit(str(exc))


class _CourtSmsFileUploadWorker(QThread):
    """在后台线程处理上传的 PDF/图片文件，提取文本并识别庭审信息。"""

    finished = Signal(object, object)  # documents, hearing_notices
    error = Signal(str)

    def __init__(self, service, file_paths: List[str]):
        super().__init__()
        self.service = service
        self.file_paths = file_paths

    def run(self):
        try:
            documents, hearing_notices = self.service.extract_text_and_hearing_notices_from_files(
                self.file_paths
            )
            self.finished.emit(documents, hearing_notices)
        except Exception as exc:
            self.error.emit(str(exc))


class ToolCenterDialog(QDialog):
    """统一风格的法律工具中心。"""

    navigate_to_case_requested = Signal(str)
    navigate_to_calendar_requested = Signal(str)

    def __init__(self, parent=None, initial_case_id: str = "", embed_mode: bool = False):
        super().__init__(parent)
        self._embed_mode = embed_mode
        if embed_mode:
            self.setWindowFlags(Qt.Widget)
        self._case_manager = CaseManager()
        self._court_sms_service = CourtSmsService() if CourtSmsService is not None else None
        self._court_sms_available = self._court_sms_service is not None
        self._initial_case_id = str(initial_case_id or "").strip()
        self._sms_parse_result: Optional[CourtSmsParseResult] = None
        self._sms_documents: List[CourtSmsDocument] = []
        self._sms_matches: List[CourtSmsCaseMatch] = []
        self._sms_hearing_notices: List[CourtSmsHearingNotice] = []
        self._sms_download_dir: Optional[Path] = None
        self._last_store_dir: Optional[Path] = None
        self._court_sms_case_options: List[Dict[str, str]] = []
        self._court_sms_selected_case_id: str = ""
        self._court_sms_case_completer_model = QStringListModel(self)
        self._last_added_court_sms_case_id: str = ""
        self._last_added_court_sms_deadline_id: str = ""
        self._last_added_court_sms_deadline_date: str = ""

        # 截图合并功能状态
        self._screenshot_merger = ScreenshotPdfMerger()
        self._screenshot_image_list: Optional[ScreenshotImageList] = None
        self._screenshot_count_label: Optional[QLabel] = None
        self._screenshot_per_page_combo: Optional[QComboBox] = None
        self._screenshot_order_combo: Optional[QComboBox] = None
        self._screenshot_orientation_combo: Optional[QComboBox] = None
        self._screenshot_margin_spin: Optional[QSpinBox] = None
        self._screenshot_gap_spin: Optional[QSpinBox] = None
        self._screenshot_label_position_combo: Optional[QComboBox] = None
        self._screenshot_label_mode_combo: Optional[QComboBox] = None
        self._screenshot_label_prefix_edit: Optional[QLineEdit] = None
        self._screenshot_label_size_combo: Optional[QComboBox] = None
        self._screenshot_status_label: Optional[QLabel] = None
        self._screenshot_generate_btn: Optional[QPushButton] = None
        self._screenshot_merge_worker: Optional[QThread] = None
        self._screenshot_save_to_source: Optional[QCheckBox] = None
        self._screenshot_source_folder: Optional[Path] = None

        # 法院短信后台线程引用（防止 GC 导致崩溃）
        self._court_sms_read_worker: Optional[QThread] = None
        self._court_sms_analyze_worker: Optional[QThread] = None
        self._court_sms_upload_worker: Optional[QThread] = None

        self._setup_ui()
        self._bind_defaults()
        self._refresh_cause_reference()

    def _setup_ui(self) -> None:
        c = COLORS
        if not self._embed_mode:
            self.setWindowTitle("工具中心")
            self.setMinimumSize(*APP_SURFACE_MIN_SIZE)
            self.resize(*APP_SURFACE_DEFAULT_SIZE)
        self.setStyleSheet(f"""
            QDialog {{
                background: {c['surface_1']};
            }}
            QTabWidget::pane {{
                border: 1px solid {c['border']};
                border-radius: 12px;
                background: {c['surface_0']};
                top: -1px;
            }}
            QTabBar::tab {{
                background: transparent;
                color: {c['text_secondary']};
                border: 1px solid transparent;
                padding: 8px 14px;
                margin-right: 4px;
                border-top-left-radius: 10px;
                border-top-right-radius: 10px;
                font-size: 12px;
                font-weight: 600;
            }}
            QTabBar::tab:selected {{
                background: {c['surface_0']};
                color: {c['accent']};
                border-color: {c['border']};
            }}
            QScrollArea {{
                border: none;
                background: transparent;
            }}
            QScrollBar:vertical {{
                background: transparent;
                width: 8px;
            }}
            QScrollBar::handle:vertical {{
                background: {c['surface_3']};
                border-radius: 4px;
                min-height: 30px;
            }}
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{
                height: 0;
            }}
            QCheckBox {{
                background: transparent;
            }}
        """)

        root = QVBoxLayout(self)
        root.setContentsMargins(16, 14, 16, 16)
        root.setSpacing(12)

        self._tabs = QTabWidget()
        root.addWidget(self._tabs, 1)

        if self._court_sms_available:
            self._tabs.addTab(self._build_court_sms_tab(), "法院短信")
        else:
            self._tabs.addTab(self._build_court_sms_unavailable_tab(), "法院短信")
        self._tabs.addTab(self._build_fee_tab(), "费用")
        self._tabs.addTab(self._build_money_tab(), "利息与金额")
        self._tabs.addTab(self._build_compensation_tab(), "赔偿")
        self._tabs.addTab(self._build_procedure_tab(), "程序")
        self._tabs.addTab(self._build_reference_tab(), "参考与导航")
        self._tabs.addTab(self._build_screenshot_merge_tab(), "截图合并")
        self._tabs.addTab(self._build_docx_compare_tab(), "文档对比")
        self._tabs.addTab(self._build_auto_format_tab(), "自动排版")

    def _build_court_sms_unavailable_tab(self) -> QWidget:
        def populate(layout: QVBoxLayout) -> None:
            card, card_layout = self._create_card(
                "法院短信模块暂不可用",
                "工具中心已经恢复可打开，但当前环境缺少法院短信模块的运行依赖，因此该页签先以提示页形式展示。",
            )

            missing_name = getattr(_COURT_SMS_IMPORT_ERROR, "name", "") if _COURT_SMS_IMPORT_ERROR else ""
            if not missing_name:
                missing_name = "requests"

            detail = QLabel(
                f"缺少依赖：{missing_name}\n"
                "请优先按项目依赖清单安装核心依赖（例如执行 `pip install -r requirements.txt`），"
                "安装完成后即可恢复法院短信读取、文书暂存和案件匹配功能。"
            )
            detail.setWordWrap(True)
            detail.setStyleSheet(
                f"""
                background: rgba(248, 250, 252, 0.92);
                color: {COLORS['text_secondary']};
                border: 1px solid {COLORS['border']};
                border-radius: 12px;
                padding: 12px 14px;
                font-size: 12px;
                line-height: 1.6;
            """
            )
            card_layout.addWidget(detail)

            note = self._make_micro_hint("其余费用、程序、赔偿、截图合并等工具不受影响，可以继续正常使用。")
            card_layout.addWidget(note)
            layout.addWidget(card)

        return self._wrap_scroll_tab(populate)

    def _wrap_scroll_tab(self, builder: Callable[[QVBoxLayout], None]) -> QWidget:
        tab = QWidget()
        outer = QVBoxLayout(tab)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.Shape.NoFrame)

        content = QWidget()
        layout = QVBoxLayout(content)
        layout.setContentsMargins(0, 14, 14, 14)
        layout.setSpacing(12)
        builder(layout)
        layout.addStretch()

        scroll.setWidget(content)
        outer.addWidget(scroll)
        return tab

    def _create_card(self, title: str, subtitle: str = "") -> Tuple[QFrame, QVBoxLayout]:
        card = QFrame()
        card.setObjectName("toolCenterCard")
        card.setAutoFillBackground(True)
        card.setStyleSheet(f"""
            QFrame#toolCenterCard {{
                background-color: {COLORS['surface_0']};
                border: 1px solid rgba(226, 232, 240, 0.92);
                border-radius: 16px;
            }}
            QFrame#toolCenterCard QLabel {{
                background-color: {COLORS['surface_0']} !important;
            }}
            QFrame#toolCenterCard QCheckBox {{
                background-color: {COLORS['surface_0']} !important;
            }}
        """)
        layout = QVBoxLayout(card)
        layout.setContentsMargins(16, 14, 16, 14)
        layout.setSpacing(10)

        title_label = QLabel(title)
        title_label.setStyleSheet(f"""
            background: transparent;
            color: {COLORS['text_primary']};
            font-size: 14px;
            font-weight: 700;
        """)
        layout.addWidget(title_label)

        if subtitle:
            subtitle_label = QLabel(subtitle)
            subtitle_label.setWordWrap(True)
            subtitle_label.setStyleSheet(f"""
                background: transparent;
                color: {COLORS['text_muted']};
                font-size: 11px;
                line-height: 1.5;
            """)
            layout.addWidget(subtitle_label)

        return card, layout

    def _create_form(self) -> TransparentFormLayout:
        form = TransparentFormLayout()
        form.setLabelAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignTop)
        form.setFormAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignTop)
        form.setContentsMargins(0, 0, 0, 0)
        form.setHorizontalSpacing(12)
        form.setVerticalSpacing(8)
        return form

    def _make_button(self, text: str, accent: bool = False) -> QPushButton:
        btn = QPushButton(text)
        btn.setFixedHeight(30)
        btn.setFlat(True)
        if accent:
            btn.setStyleSheet(f"""
                QPushButton {{
                    background: {COLORS['accent_subtle']};
                    color: {COLORS['accent']};
                    border: 1px solid {COLORS['accent_light']};
                    border-radius: 10px;
                    padding: 0 14px;
                    font-size: 11px;
                    font-weight: 700;
                }}
                QPushButton:hover {{
                    background: rgba(219, 234, 254, 0.92);
                }}
            """)
        else:
            btn.setStyleSheet(f"""
                QPushButton {{
                    background: {COLORS['surface_1']};
                    color: {COLORS['text_secondary']};
                    border: 1px solid {COLORS['border']};
                    border-radius: 10px;
                    padding: 0 14px;
                    font-size: 11px;
                    font-weight: 700;
                }}
                QPushButton:hover {{
                    background: {COLORS['surface_2']};
                    color: {COLORS['text_primary']};
                }}
            """)
        return btn

    def _make_result_browser(self, min_height: int = 92, max_height: Optional[int] = None) -> QTextBrowser:
        browser = QTextBrowser()
        browser.setOpenExternalLinks(True)
        browser.setMinimumHeight(min_height)
        if max_height is not None:
            browser.setMaximumHeight(max_height)
        browser.setStyleSheet(f"""
            QTextBrowser {{
                background: rgba(248, 250, 252, 0.95);
                color: {COLORS['text_secondary']};
                border: 1px solid {COLORS['border']};
                border-radius: 12px;
                padding: 8px 10px;
                font-size: 12px;
                line-height: 1.6;
            }}
        """)
        return browser

    def _make_money_spin(
        self,
        maximum: float = 999999999.99,
        decimals: int = 2,
        blank_zero: bool = False,
    ) -> QDoubleSpinBox:
        widget = BlankZeroMoneySpinBox() if blank_zero else QDoubleSpinBox()
        widget.setRange(0, maximum)
        widget.setDecimals(decimals)
        widget.setSingleStep(100)
        widget.setButtonSymbols(QDoubleSpinBox.ButtonSymbols.NoButtons)
        widget.setMinimumHeight(34)
        widget.setGroupSeparatorShown(True)
        if blank_zero:
            widget.lineEdit().setPlaceholderText("请输入金额")
        return widget

    def _make_fee_result_browser(self) -> QTextBrowser:
        """费用页紧凑结果框。"""
        return self._make_result_browser(min_height=58, max_height=76)

    def _make_rate_spin(self, maximum: float = 1000.0, decimals: int = 4) -> QDoubleSpinBox:
        widget = self._make_money_spin(maximum=maximum, decimals=decimals)
        widget.setSingleStep(0.1)
        return widget

    def _make_int_spin(self, maximum: int = 99999) -> QSpinBox:
        widget = QSpinBox()
        widget.setRange(0, maximum)
        widget.setButtonSymbols(QSpinBox.ButtonSymbols.NoButtons)
        widget.setMinimumHeight(34)
        return widget

    def _make_date_edit(self) -> QDateEdit:
        widget = QDateEdit()
        widget.setCalendarPopup(True)
        widget.setDate(QDate.currentDate())
        widget.setMinimumHeight(34)
        # 独立创建日历控件并设置浅色样式，避免 QDateEdit 本体被样式表干扰
        calendar = QCalendarWidget()
        calendar.setStyleSheet("""
            QCalendarWidget { background-color: #ffffff; }
            QCalendarWidget QTableView {
                background-color: #ffffff; color: #1f2937;
                selection-background-color: #dbeafe; selection-color: #2563eb;
                border: none; outline: none;
            }
            QCalendarWidget QTableView::item { padding: 4px; border-radius: 4px; }
            QCalendarWidget QTableView::item:selected {
                background-color: #dbeafe; color: #2563eb; font-weight: 700;
            }
            QCalendarWidget QTableView::item:hover { background-color: #eff6ff; }
            QCalendarWidget QHeaderView::section {
                background-color: #f9fafb; color: #6b7280; border: none;
                padding: 6px 4px; font-size: 12px; font-weight: 600;
            }
            QCalendarWidget QWidget#qt_calendar_navigationbar {
                background-color: #f9fafb; border-bottom: 1px solid #e5e7eb; padding: 4px;
            }
            QCalendarWidget QSpinBox, QCalendarWidget QComboBox {
                background-color: #ffffff; color: #1f2937;
                border: 1px solid #d1d5db; border-radius: 4px;
                padding: 2px 6px; font-size: 13px; font-weight: 600;
            }
            QCalendarWidget QToolButton {
                background-color: transparent; color: #374151;
                border: none; border-radius: 4px; padding: 2px;
                font-size: 13px; font-weight: 700;
            }
            QCalendarWidget QToolButton:hover { background-color: #e5e7eb; }
        """)
        widget.setCalendarWidget(calendar)
        return widget

    def _make_line_edit(self, placeholder: str = "") -> QLineEdit:
        widget = QLineEdit()
        widget.setPlaceholderText(placeholder)
        widget.setMinimumHeight(34)
        return widget

    def _make_micro_hint(self, text: str = "") -> QLabel:
        label = QLabel(text)
        label.setWordWrap(True)
        label.setStyleSheet(f"""
            background: transparent;
            color: {COLORS['text_muted']};
            font-size: 11px;
            padding: 0 2px;
        """)
        return label

    def _make_multiline_edit(self, placeholder: str = "", min_height: int = 118) -> QTextEdit:
        widget = QTextEdit()
        widget.setPlaceholderText(placeholder)
        widget.setMinimumHeight(min_height)
        widget.setStyleSheet(f"""
            QTextEdit {{
                background: {COLORS['surface_1']};
                color: {COLORS['text_primary']};
                border: 1px solid {COLORS['border']};
                border-radius: 12px;
                padding: 9px 10px;
                font-size: 12px;
                line-height: 1.6;
            }}
            QTextEdit:focus {{
                border-color: {COLORS['accent_light']};
                background: {COLORS['surface_0']};
            }}
        """)
        return widget

    def _make_tree_widget(self, headers: Optional[List[str]] = None, min_height: int = 260) -> QTreeWidget:
        tree = QTreeWidget()
        tree.setRootIsDecorated(False)
        tree.setAlternatingRowColors(False)
        tree.setUniformRowHeights(True)
        tree.setMinimumHeight(min_height)
        resolved_headers = headers or ["文书名称", "法院", "时间", "状态"]
        tree.setColumnCount(len(resolved_headers))
        tree.setHeaderLabels(resolved_headers)
        tree.setStyleSheet(f"""
            QTreeWidget {{
                background: rgba(248, 250, 252, 0.92);
                color: {COLORS['text_secondary']};
                border: 1px solid {COLORS['border']};
                border-radius: 12px;
                padding: 4px;
                font-size: 12px;
            }}
            QTreeWidget::item {{
                height: 28px;
            }}
            QHeaderView::section {{
                background: transparent;
                color: {COLORS['text_muted']};
                padding: 6px 8px;
                border: none;
                border-bottom: 1px solid {COLORS['border']};
                font-size: 11px;
                font-weight: 700;
            }}
        """)
        return tree

    def _format_file_size(self, size_bytes: int) -> str:
        if size_bytes <= 0:
            return "已读取"
        units = ["B", "KB", "MB", "GB"]
        size = float(size_bytes)
        unit = units[0]
        for unit in units:
            if size < 1024 or unit == units[-1]:
                break
            size /= 1024
        if unit == "B":
            return f"{int(size)} {unit}"
        return f"{size:.1f} {unit}"

    def _format_money_html(self, title: str, amount) -> str:
        return (
            f"<div style='font-size:13px; color:{COLORS['text_primary']}; font-weight:700;'>{title}</div>"
            f"<div style='margin-top:6px; font-size:20px; color:{COLORS['accent']}; font-weight:800;'>¥ {amount}</div>"
        )

    def _build_fee_tab(self) -> QWidget:
        return self._wrap_scroll_tab(self._populate_fee_tab)

    def _build_court_sms_tab(self) -> QWidget:
        return self._wrap_scroll_tab(self._populate_court_sms_tab)

    def _populate_court_sms_tab(self, layout: QVBoxLayout) -> None:
        body = QGridLayout()
        body.setContentsMargins(0, 0, 0, 0)
        body.setHorizontalSpacing(12)
        body.setVerticalSpacing(18)
        body.addWidget(self._build_court_sms_input_card(), 0, 0)
        body.addWidget(self._build_court_sms_match_card(), 0, 1)
        body.addWidget(self._build_court_sms_documents_card(), 1, 0)
        body.addWidget(self._build_court_sms_hearing_card(), 1, 1)
        body.setColumnStretch(0, 10)
        body.setColumnStretch(1, 9)
        body.setRowStretch(0, 0)
        body.setRowStretch(1, 1)
        layout.addLayout(body, 1)

    def _populate_fee_tab(self, layout: QVBoxLayout) -> None:
        layout.addWidget(self._build_litigation_fee_card())
        layout.addWidget(self._build_execution_preservation_card())
        layout.addWidget(self._build_bankruptcy_fee_card())

    def _build_money_tab(self) -> QWidget:
        return self._wrap_scroll_tab(self._populate_money_tab)

    def _populate_money_tab(self, layout: QVBoxLayout) -> None:
        layout.addWidget(self._build_interest_card())
        layout.addWidget(self._build_liquidated_card())
        layout.addWidget(self._build_occupation_card())
        layout.addWidget(self._build_delay_interest_card())
        layout.addWidget(self._build_lawyer_fee_card())
        layout.addWidget(self._build_labor_compensation_card())

    def _build_compensation_tab(self) -> QWidget:
        return self._wrap_scroll_tab(self._populate_compensation_tab)

    def _populate_compensation_tab(self, layout: QVBoxLayout) -> None:
        layout.addWidget(self._build_work_injury_card())
        layout.addWidget(self._build_traffic_injury_card())

    def _build_procedure_tab(self) -> QWidget:
        return self._wrap_scroll_tab(self._populate_procedure_tab)

    def _populate_procedure_tab(self, layout: QVBoxLayout) -> None:
        layout.addWidget(self._build_procedural_limit_card())
        layout.addWidget(self._build_date_offset_card())

    def _build_reference_tab(self) -> QWidget:
        return self._wrap_scroll_tab(self._populate_reference_tab)

    def _populate_reference_tab(self, layout: QVBoxLayout) -> None:
        layout.addWidget(self._build_cause_reference_card())
        layout.addWidget(self._build_reference_links_card())

    def _build_court_sms_input_card(self) -> QWidget:
        card, layout = self._create_card(
            "法院短信读取",
            "粘贴短信后自动解析案号、法院、收件人和链接参数，并读取送达文书。"
            "也可直接拖拽或上传传票 PDF/图片，自动识别开庭信息。",
        )
        layout.setContentsMargins(14, 12, 14, 12)
        layout.setSpacing(8)

        self._court_sms_input = self._make_multiline_edit(
            "粘贴法院短信到此处，或直接拖拽传票 PDF/图片文件到此框内…",
            min_height=108,
        )
        self._court_sms_input.setMaximumHeight(126)
        self._court_sms_input.setAcceptDrops(True)
        self._court_sms_input.installEventFilter(self)
        # QTextEdit 的 viewport 才是实际接收拖放事件的控件
        self._court_sms_input.viewport().installEventFilter(self)
        layout.addWidget(self._court_sms_input)

        action_row = QHBoxLayout()
        action_row.setSpacing(8)
        paste_btn = self._make_button("粘贴短信")
        paste_btn.clicked.connect(self._paste_court_sms_from_clipboard)
        action_row.addWidget(paste_btn)

        parse_btn = self._make_button("智能解析并读取", accent=True)
        parse_btn.clicked.connect(self._parse_court_sms_only)
        action_row.addWidget(parse_btn)
        self._btn_parse_court_sms = parse_btn

        clear_btn = self._make_button("清空")
        clear_btn.clicked.connect(self._clear_court_sms_state)
        action_row.addWidget(clear_btn)
        self._btn_clear_court_sms = clear_btn

        upload_btn = self._make_button("上传文件/图片")
        upload_btn.clicked.connect(self._on_court_sms_upload_files)
        action_row.addWidget(upload_btn)
        self._btn_upload_court_sms = upload_btn

        action_row.addStretch()
        layout.addLayout(action_row)

        self._court_sms_summary = self._make_result_browser(112)
        layout.addWidget(self._court_sms_summary)
        return card

    def _build_court_sms_documents_card(self) -> QWidget:
        card, layout = self._create_card(
            "文书清单",
            "双击文书可直接打开；未暂存时会打开法院原始链接。",
        )
        layout.setContentsMargins(14, 12, 14, 12)
        layout.setSpacing(8)

        self._court_sms_doc_tree = self._make_tree_widget(min_height=230)
        self._court_sms_doc_tree.itemDoubleClicked.connect(self._open_court_sms_document)
        layout.addWidget(self._court_sms_doc_tree)

        self._court_sms_doc_hint = QLabel("尚未开始读取法院文书。点击“智能解析并读取”后会自动暂存文书。")
        self._court_sms_doc_hint.setStyleSheet(
            f"background: transparent; color: {COLORS['text_muted']}; font-size: 11px;"
        )
        layout.addWidget(self._court_sms_doc_hint)
        return card

    def _build_court_sms_match_card(self) -> QWidget:
        card, layout = self._create_card(
            "案件匹配与确认存放",
            "按案号、法院、当事人给出建议案件；确认后才会实际落盘。",
        )
        layout.setContentsMargins(14, 12, 14, 12)
        layout.setSpacing(8)

        form = self._create_form()
        self._court_sms_case_combo = QComboBox()
        self._court_sms_case_combo.setEditable(True)
        self._court_sms_case_combo.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self._court_sms_case_combo.setMaxVisibleItems(12)
        self._court_sms_case_combo.setMinimumHeight(36)
        self._court_sms_case_combo.currentIndexChanged.connect(self._refresh_court_sms_match_preview)
        self._court_sms_case_combo.activated.connect(self._select_court_sms_case_from_combo)
        combo_line_edit = self._court_sms_case_combo.lineEdit()
        if combo_line_edit:
            combo_line_edit.setPlaceholderText("建议案件，可直接输入搜索或下拉选择")
            combo_line_edit.installEventFilter(self)
            combo_line_edit.textEdited.connect(self._handle_court_sms_case_search_input)
            combo_line_edit.returnPressed.connect(self._commit_court_sms_case_search_text)

        popup_style = f"""
            QAbstractItemView {{
                background: {COLORS['surface_0']};
                color: {COLORS['text_primary']};
                border: 1px solid {COLORS['border']};
                border-radius: 12px;
                outline: none;
                padding: 3px;
                font-size: 11px;
                selection-background-color: {COLORS['accent_subtle']};
                selection-color: {COLORS['accent']};
            }}
            QAbstractItemView::item {{
                min-height: 26px;
                padding: 0 6px;
                margin: 1px 3px;
                border-radius: 7px;
            }}
            QAbstractItemView::item:hover {{
                background: rgba(241, 245, 249, 0.92);
            }}
        """
        self._court_sms_case_combo.view().setStyleSheet(popup_style)
        if hasattr(self._court_sms_case_combo.view(), "setSpacing"):
            self._court_sms_case_combo.view().setSpacing(0)

        combo_completer = QCompleter(self._court_sms_case_completer_model, self)
        combo_completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        combo_completer.setFilterMode(Qt.MatchFlag.MatchContains)
        combo_completer.setCompletionMode(QCompleter.CompletionMode.PopupCompletion)
        combo_completer.activated[str].connect(self._select_court_sms_case_from_text)
        self._court_sms_case_combo.setCompleter(combo_completer)
        if combo_completer.popup():
            combo_completer.popup().setStyleSheet(popup_style)
            if hasattr(combo_completer.popup(), "setSpacing"):
                combo_completer.popup().setSpacing(0)
        self._court_sms_target_folder = self._make_line_edit("例如：法院送达文书/（2026）皖1722民初273号_石台县人民法院")
        self._court_sms_custom_save = QCheckBox("自定义保存到任意文件夹")
        self._court_sms_custom_save.toggled.connect(self._toggle_court_sms_custom_save_mode)
        self._court_sms_custom_dir = self._make_line_edit("可选择任意文件夹作为保存目标")
        self._court_sms_custom_dir.textChanged.connect(self._refresh_court_sms_match_preview)
        self._court_sms_custom_dir_btn = self._make_button("选择文件夹")
        self._court_sms_custom_dir_btn.clicked.connect(self._choose_court_sms_custom_directory)
        custom_dir_row = QHBoxLayout()
        custom_dir_row.setSpacing(8)
        custom_dir_row.addWidget(self._court_sms_custom_dir, 1)
        custom_dir_row.addWidget(self._court_sms_custom_dir_btn)

        form.addRow("建议案件", self._court_sms_case_combo)
        self._court_sms_case_hint = self._make_micro_hint("点击“智能解析并读取”后，系统会在这里给出建议案件。")
        form.addRow("", self._court_sms_case_hint)
        form.addRow("保存目录", self._court_sms_target_folder)
        form.addRow("", self._court_sms_custom_save)
        form.addRow("自定义目录", custom_dir_row)
        layout.addLayout(form)

        self._court_sms_match_browser = self._make_result_browser(86, max_height=112)
        layout.addWidget(self._court_sms_match_browser)

        action_row = QHBoxLayout()
        action_row.setSpacing(8)
        open_temp_btn = self._make_button("打开暂存目录")
        open_temp_btn.clicked.connect(self._open_court_sms_staging_folder)
        action_row.addWidget(open_temp_btn)

        save_btn = self._make_button("确认存放到案件", accent=True)
        save_btn.clicked.connect(self._confirm_store_court_documents)
        action_row.addWidget(save_btn)

        open_saved_btn = self._make_button("打开已存入文件夹")
        open_saved_btn.clicked.connect(self._open_court_sms_saved_folder)
        action_row.addWidget(open_saved_btn)
        action_row.addStretch()
        layout.addLayout(action_row)

        self._court_sms_store_result = self._make_result_browser(64, max_height=86)
        layout.addWidget(self._court_sms_store_result)
        return card

    def _build_court_sms_hearing_card(self) -> QWidget:
        card, layout = self._create_card(
            "庭审文书识别",
            "自动识别传票、出庭通知书或开庭通知书，可加入期限提醒。",
        )
        layout.setContentsMargins(14, 12, 14, 12)
        layout.setSpacing(8)

        self._court_sms_hearing_tree = self._make_tree_widget(
            ["文书", "类型", "被传唤人", "开庭时间", "地点", "状态"],
            min_height=92,
        )
        self._court_sms_hearing_tree.setMaximumHeight(118)
        self._court_sms_hearing_tree.itemSelectionChanged.connect(self._refresh_court_sms_hearing_preview)
        layout.addWidget(self._court_sms_hearing_tree)

        self._court_sms_hearing_hint = self._make_micro_hint("点击“智能解析并读取”后，这里会自动识别庭审类文书。")
        layout.addWidget(self._court_sms_hearing_hint)

        self._court_sms_hearing_browser = self._make_result_browser(122, max_height=150)
        layout.addWidget(self._court_sms_hearing_browser)

        # 可编辑字段表单（双击编辑）
        self._hearing_edit_frame = QFrame()
        self._hearing_edit_frame.hide()
        edit_layout = QVBoxLayout(self._hearing_edit_frame)
        edit_layout.setContentsMargins(0, 6, 0, 0)
        edit_layout.setSpacing(4)

        edit_hint = QLabel('双击字段可编辑，修改后点击"加入期限"生效')
        edit_hint.setStyleSheet(f"color: {COLORS['text_muted']}; font-size: 10px; background: transparent; border: none;")
        edit_layout.addWidget(edit_hint)

        self._hearing_edit_fields = {}
        _EDITABLE_FIELDS = [
            ("case_number", "案号"),
            ("summoned_person", "被传唤人"),
            ("hearing_date", "开庭日期"),
            ("hearing_time", "开庭时间"),
            ("hearing_place", "开庭地点"),
            ("cause", "案由"),
            ("court_name", "法院"),
        ]
        for field_key, field_label in _EDITABLE_FIELDS:
            row = QHBoxLayout()
            row.setSpacing(8)
            label = QLabel(f"{field_label}：")
            label.setFixedWidth(70)
            label.setStyleSheet(f"color: {COLORS['text_secondary']}; font-size: 12px; background: transparent; border: none;")
            row.addWidget(label)

            value_label = QLabel("")
            value_label.setStyleSheet(f"color: {COLORS['text_primary']}; font-size: 12px; font-weight: 500; background: transparent; border: none; padding: 2px 4px; border-radius: 4px;")
            value_label.setCursor(Qt.CursorShape.IBeamCursor)
            value_label.setToolTip(f"双击编辑{field_label}")
            value_label.setWordWrap(True)
            row.addWidget(value_label, 1)

            value_edit = QLineEdit()
            value_edit.hide()
            value_edit.setStyleSheet(f"""
                QLineEdit {{
                    background: {COLORS['surface_0']};
                    color: {COLORS['text_primary']};
                    border: 1px solid {COLORS['accent']};
                    border-radius: 6px;
                    padding: 2px 6px;
                    font-size: 12px;
                }}
            """)
            row.addWidget(value_edit, 1)

            # 双击切换编辑
            value_label.mouseDoubleClickEvent = lambda ev, lbl=value_label, edt=value_edit: self._toggle_hearing_field_edit(lbl, edt)
            value_edit.editingFinished.connect(lambda lbl=value_label, edt=value_edit: self._finish_hearing_field_edit(lbl, edt))

            edit_layout.addLayout(row)
            self._hearing_edit_fields[field_key] = (value_label, value_edit)

        layout.addWidget(self._hearing_edit_frame)

        action_row = QHBoxLayout()
        action_row.setSpacing(6)
        self._btn_refresh_court_sms_hearing = self._make_button("重新识别")
        self._btn_refresh_court_sms_hearing.clicked.connect(self._analyze_court_sms_hearing_notices)
        action_row.addWidget(self._btn_refresh_court_sms_hearing)

        self._btn_add_court_sms_hearing_deadline = self._make_button("加入期限", accent=True)
        self._btn_add_court_sms_hearing_deadline.clicked.connect(self._add_selected_court_sms_hearing_deadline)
        action_row.addWidget(self._btn_add_court_sms_hearing_deadline)

        self._btn_unlink_court_sms_case = self._make_button("不关联案件")
        self._btn_unlink_court_sms_case.setToolTip("不关联到具体案件，仅将文书保存到自定义文件夹")
        self._btn_unlink_court_sms_case.clicked.connect(self._on_unlink_court_sms_case)
        action_row.addWidget(self._btn_unlink_court_sms_case)

        self._btn_open_court_sms_case = self._make_button("回到案件")
        self._btn_open_court_sms_case.clicked.connect(self._open_court_sms_case_from_hearing)
        self._btn_open_court_sms_case.setEnabled(False)
        action_row.addWidget(self._btn_open_court_sms_case)

        self._btn_open_court_sms_calendar = self._make_button("查看全部期限")
        self._btn_open_court_sms_calendar.clicked.connect(self._open_court_sms_calendar_from_hearing)
        self._btn_open_court_sms_calendar.setEnabled(False)
        action_row.addWidget(self._btn_open_court_sms_calendar)
        action_row.addStretch()
        layout.addLayout(action_row)

        self._court_sms_hearing_result_hint = self._make_micro_hint("识别结果会跟随当前所选案件动态判断是否可直接加入期限。")
        layout.addWidget(self._court_sms_hearing_result_hint)
        return card

    def eventFilter(self, watched, event):  # type: ignore[override]
        # 案件搜索 combo 点击
        combo = getattr(self, "_court_sms_case_combo", None)
        line_edit = combo.lineEdit() if combo else None
        if (
            line_edit
            and watched is line_edit
            and event.type() == QEvent.Type.MouseButtonPress
            and not self._court_sms_custom_save.isChecked()
        ):
            self._begin_court_sms_case_search()

        # 法院短信输入框拖拽文件（QTextEdit 本体 + viewport 都要拦截）
        sms_input = getattr(self, "_court_sms_input", None)
        sms_viewport = sms_input.viewport() if sms_input else None
        if sms_input and (watched is sms_input or watched is sms_viewport):
            if event.type() == QEvent.Type.DragEnter:
                mime = event.mimeData()
                if mime.hasUrls():
                    urls = mime.urls()
                    if any(u.toLocalFile().lower().endswith(
                        (".pdf", ".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif")
                    ) for u in urls):
                        event.acceptProposedAction()
                        return True
            elif event.type() == QEvent.Type.DragMove:
                event.acceptProposedAction()
                return True
            elif event.type() == QEvent.Type.Drop:
                mime = event.mimeData()
                paths = [
                    u.toLocalFile() for u in mime.urls()
                    if u.toLocalFile().lower().endswith(
                        (".pdf", ".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif")
                    )
                ]
                if paths:
                    self._process_uploaded_court_sms_files(paths)
                    return True

        return super().eventFilter(watched, event)

    def _build_litigation_fee_card(self) -> QWidget:
        card, layout = self._create_card(
            "诉讼费用",
            "优先覆盖财产案件、离婚案件、人格权纠纷等常用口径。涉及法定区间的事项，支持手动输入法院实际确定的基础收费。",
        )
        form = self._create_form()

        self._litigation_type = QComboBox()
        self._litigation_type.addItem("财产案件", "property")
        self._litigation_type.addItem("离婚案件", "divorce")
        self._litigation_type.addItem("人格权纠纷", "personality")
        self._litigation_type.addItem("劳动争议", "labor")
        self._litigation_type.addItem("行政案件", "administrative")
        self._litigation_type.addItem("知识产权（无争议金额）", "ip_no_amount")
        self._litigation_type.addItem("管辖权异议", "jurisdiction")
        self._litigation_type.currentIndexChanged.connect(self._update_litigation_mode_hint)

        self._litigation_amount = self._make_money_spin(blank_zero=True)
        self._litigation_base = self._make_money_spin(maximum=10000, decimals=2)
        self._litigation_base.setValue(50)

        form.addRow("案件类型", self._litigation_type)
        form.addRow("金额 / 财产价值", self._litigation_amount)
        form.addRow("基础收费", self._litigation_base)
        layout.addLayout(form)

        self._litigation_hint = QLabel("")
        self._litigation_hint.setWordWrap(True)
        self._litigation_hint.setStyleSheet(f"background:transparent;color:{COLORS['text_muted']};font-size:11px;")
        layout.addWidget(self._litigation_hint)

        action_row = QHBoxLayout()
        action_row.addStretch()
        btn = self._make_button("计算诉讼费用", accent=True)
        btn.clicked.connect(self._calculate_litigation_fee)
        action_row.addWidget(btn)
        layout.addLayout(action_row)

        self._litigation_result = self._make_fee_result_browser()
        layout.addWidget(self._litigation_result)
        return card

    def _build_execution_preservation_card(self) -> QWidget:
        card, layout = self._create_card(
            "执行费 / 保全费",
            "按照《诉讼费用交纳办法》的常用申请费口径。无金额模式可填入法院实际确定的固定收费。",
        )
        form = self._create_form()

        self._apply_fee_type = QComboBox()
        self._apply_fee_type.addItem("申请执行费", "execution")
        self._apply_fee_type.addItem("财产保全费", "preservation")

        self._apply_amount = self._make_money_spin(blank_zero=True)
        self._apply_no_amount = QCheckBox("无金额 / 固定收费模式")
        self._apply_fixed = self._make_money_spin(maximum=5000, decimals=2)
        self._apply_fixed.setValue(50)

        form.addRow("费用类型", self._apply_fee_type)
        form.addRow("标的金额", self._apply_amount)
        form.addRow("", self._apply_no_amount)
        form.addRow("固定收费", self._apply_fixed)
        layout.addLayout(form)

        btn = self._make_button("计算申请费", accent=True)
        btn.clicked.connect(self._calculate_apply_fee)
        layout.addWidget(btn, 0, Qt.AlignmentFlag.AlignRight)

        self._apply_result = self._make_fee_result_browser()
        layout.addWidget(self._apply_result)
        return card

    def _build_bankruptcy_fee_card(self) -> QWidget:
        card, layout = self._create_card(
            "破产管理人报酬",
            "按《企业破产案件管理人报酬规定》的分段上限口径计算参考值。实际报酬由法院结合工作量、案件复杂程度等调整。",
        )
        form = self._create_form()
        self._bankruptcy_assets = self._make_money_spin(blank_zero=True)
        self._bankruptcy_factor = self._make_rate_spin(maximum=3, decimals=2)
        self._bankruptcy_factor.setValue(1.0)
        form.addRow("可分配 / 清偿财产价值", self._bankruptcy_assets)
        form.addRow("调整系数", self._bankruptcy_factor)
        layout.addLayout(form)

        btn = self._make_button("计算报酬参考", accent=True)
        btn.clicked.connect(self._calculate_bankruptcy_fee)
        layout.addWidget(btn, 0, Qt.AlignmentFlag.AlignRight)

        self._bankruptcy_result = self._make_fee_result_browser()
        layout.addWidget(self._bankruptcy_result)
        return card

    def _build_interest_card(self) -> QWidget:
        card, layout = self._create_card(
            "利息计算",
            "支持 LPR 分段计息与固定利率计息，自动匹配全国银行间同业拆借中心公布的贷款市场报价利率。",
        )
        # 增加卡片底部边距，避免圆角区域截断按钮下沿边框
        layout.setContentsMargins(16, 14, 16, 22)
        # 左右分栏：左侧输入，右侧结果
        content_split = QHBoxLayout()
        content_split.setSpacing(16)

        # ─── 左侧：输入区域 ───
        left_panel = QWidget()
        left_panel.setStyleSheet("background: transparent;")
        left_layout = QVBoxLayout(left_panel)
        left_layout.setSpacing(6)
        left_layout.setContentsMargins(0, 0, 0, 0)

        form = self._create_form()
        self._interest_principal = self._make_money_spin(blank_zero=True)
        self._interest_principal.lineEdit().setPlaceholderText("请输入金额")
        form.addRow("计算基数", self._interest_principal)

        self._interest_rate_type_combo = QComboBox()
        self._interest_rate_type_combo.addItem("全国银行间同业拆借中心公布的贷款市场报价利率(LPR)", "lpr")
        form.addRow("利率类型", self._interest_rate_type_combo)

        # 日期行
        date_widget = QWidget()
        date_grid = QGridLayout(date_widget)
        date_grid.setHorizontalSpacing(6)
        date_grid.setVerticalSpacing(4)
        date_grid.setContentsMargins(0, 0, 0, 0)
        self._interest_start_date = self._make_date_edit()
        self._interest_end_date = self._make_date_edit()
        date_grid.addWidget(QLabel("起始日期"), 0, 0)
        date_grid.addWidget(self._interest_start_date, 0, 1)
        date_grid.addWidget(QLabel("截止日期"), 1, 0)
        date_grid.addWidget(self._interest_end_date, 1, 1)
        form.addRow("", date_widget)

        # 起止日期选项（2×2）
        opt_widget = QWidget()
        opt_grid = QGridLayout(opt_widget)
        opt_grid.setHorizontalSpacing(8)
        opt_grid.setVerticalSpacing(4)
        opt_grid.setContentsMargins(0, 0, 0, 0)
        self._interest_date_opt_both = QCheckBox("起止日期均计算在内")
        self._interest_date_opt_start_only = QCheckBox("起始日期计算在内，截止日期不计算在内")
        self._interest_date_opt_end_only = QCheckBox("起始日期不计算在内，截止日期计算在内")
        self._interest_date_opt_neither = QCheckBox("起止日期均不计算在内")
        self._interest_date_opt_both.setChecked(True)
        opt_grid.addWidget(self._interest_date_opt_both, 0, 0)
        opt_grid.addWidget(self._interest_date_opt_start_only, 0, 1)
        opt_grid.addWidget(self._interest_date_opt_neither, 1, 0)
        opt_grid.addWidget(self._interest_date_opt_end_only, 1, 1)
        form.addRow("起止日期选项", opt_widget)

        # 一年为
        basis_widget = QWidget()
        basis_row = QHBoxLayout(basis_widget)
        basis_row.setSpacing(6)
        basis_row.setContentsMargins(0, 0, 0, 0)
        self._interest_year_basis_360 = QCheckBox("360天")
        self._interest_year_basis_365 = QCheckBox("365天")
        self._interest_year_basis_365.setChecked(True)
        basis_row.addWidget(self._interest_year_basis_360)
        basis_row.addWidget(self._interest_year_basis_365)
        basis_row.addStretch()
        form.addRow("一年为", basis_widget)

        left_layout.addLayout(form)

        # LPR 选项
        lpr_group = QWidget()
        lpr_layout = QVBoxLayout(lpr_group)
        lpr_layout.setSpacing(4)
        lpr_layout.setContentsMargins(0, 0, 0, 0)

        lpr_label = QLabel("LPR选项")
        lpr_label.setStyleSheet(f"color: {COLORS['text_primary']}; font-weight: 600;")
        lpr_layout.addWidget(lpr_label)

        # LPR计算方式
        mode_row = QHBoxLayout()
        mode_row.setSpacing(6)
        self._interest_lpr_mode_seg = QCheckBox("分段LPR")
        self._interest_lpr_mode_fixed = QCheckBox("指定LPR")
        self._interest_lpr_mode_seg.setChecked(True)
        mode_row.addWidget(QLabel("计算方式:"))
        mode_row.addWidget(self._interest_lpr_mode_seg)
        mode_row.addWidget(self._interest_lpr_mode_fixed)
        mode_row.addStretch()
        lpr_layout.addLayout(mode_row)

        # LPR档次 + 指定利率
        term_row = QHBoxLayout()
        term_row.setSpacing(6)
        self._interest_lpr_term_1y = QCheckBox("1年期LPR")
        self._interest_lpr_term_5y = QCheckBox("5年期以上LPR")
        self._interest_lpr_term_1y.setChecked(True)
        term_row.addWidget(QLabel("LPR档次:"))
        term_row.addWidget(self._interest_lpr_term_1y)
        term_row.addWidget(self._interest_lpr_term_5y)
        term_row.addSpacing(12)
        self._interest_lpr_fixed_rate = self._make_rate_spin(maximum=100, decimals=4)
        self._interest_lpr_fixed_rate.setVisible(False)
        term_row.addWidget(QLabel("指定利率(%):"))
        term_row.addWidget(self._interest_lpr_fixed_rate)
        term_row.addStretch()
        lpr_layout.addLayout(term_row)

        # 调整方式
        adjust_row = QHBoxLayout()
        adjust_row.setSpacing(6)
        self._interest_lpr_adjust_combo = QComboBox()
        self._interest_lpr_adjust_combo.addItem("倍数", "multiple")
        self._interest_lpr_adjust_combo.addItem("加计", "add")
        self._interest_lpr_adjust_combo.addItem("浮动", "float")
        self._interest_lpr_adjust_value = self._make_rate_spin(maximum=1000, decimals=4)
        adjust_row.addWidget(QLabel("调整:"))
        adjust_row.addWidget(self._interest_lpr_adjust_combo)
        adjust_row.addWidget(self._interest_lpr_adjust_value)
        adjust_row.addStretch()
        lpr_layout.addLayout(adjust_row)

        left_layout.addWidget(lpr_group)

        left_layout.addWidget(self._make_micro_hint("例如：加计50%，相当于上浮50%，相当于1.5倍"))
        left_layout.addStretch()

        # 按钮行
        btn_row = QHBoxLayout()
        btn_row.addStretch()
        calc_btn = QToolButton()
        calc_btn.setText("计算利息")
        calc_btn.setFixedHeight(28)
        calc_btn.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextOnly)
        calc_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        calc_btn.setAttribute(Qt.WidgetAttribute.WA_LayoutUsesWidgetRect)
        calc_btn.setStyleSheet(f"""
            QToolButton {{
                background: {COLORS['accent_subtle']};
                color: {COLORS['accent']};
                border: 1px solid {COLORS['accent_light']};
                border-radius: 8px;
                padding: 0 14px;
                font-size: 11px;
                font-weight: 700;
            }}
            QToolButton:hover {{
                background: rgba(219, 234, 254, 0.92);
            }}
        """)
        calc_btn.clicked.connect(self._calculate_interest)
        btn_row.addWidget(calc_btn)
        left_layout.addLayout(btn_row)
        left_layout.addSpacing(3)

        content_split.addWidget(left_panel, 1)

        # ─── 右侧：结果区域 ───
        right_panel = QWidget()
        right_panel.setStyleSheet("background: transparent;")
        right_layout = QVBoxLayout(right_panel)
        right_layout.setSpacing(8)
        right_layout.setContentsMargins(0, 0, 0, 0)

        self._interest_result = self._make_result_browser()
        self._interest_result.setMinimumWidth(320)
        right_layout.addWidget(self._interest_result, 1)

        # 导出按钮
        export_row = QHBoxLayout()
        export_row.addStretch()
        self._interest_export_btn = QToolButton()
        self._interest_export_btn.setText("导出结果")
        self._interest_export_btn.setFixedHeight(28)
        self._interest_export_btn.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextOnly)
        self._interest_export_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self._interest_export_btn.setAttribute(Qt.WidgetAttribute.WA_LayoutUsesWidgetRect)
        self._interest_export_btn.setStyleSheet(f"""
            QToolButton {{
                background: {COLORS['surface_1']};
                color: {COLORS['text_secondary']};
                border: 1px solid {COLORS['border']};
                border-radius: 8px;
                padding: 0 14px;
                font-size: 11px;
                font-weight: 700;
            }}
            QToolButton:hover {{
                background: {COLORS['surface_2']};
                color: {COLORS['text_primary']};
            }}
        """)
        self._interest_export_btn.setEnabled(False)
        self._interest_export_btn.clicked.connect(self._export_interest_result)
        export_row.addWidget(self._interest_export_btn)
        right_layout.addLayout(export_row)
        right_layout.addSpacing(3)

        content_split.addWidget(right_panel, 1)
        layout.addLayout(content_split)

        # 互斥处理
        def _on_date_opt_changed(changed_cb):
            for cb in (self._interest_date_opt_both, self._interest_date_opt_start_only,
                       self._interest_date_opt_end_only, self._interest_date_opt_neither):
                if cb != changed_cb:
                    cb.blockSignals(True)
                    cb.setChecked(False)
                    cb.blockSignals(False)
            if not changed_cb.isChecked():
                changed_cb.blockSignals(True)
                changed_cb.setChecked(True)
                changed_cb.blockSignals(False)

        self._interest_date_opt_both.stateChanged.connect(lambda: _on_date_opt_changed(self._interest_date_opt_both))
        self._interest_date_opt_start_only.stateChanged.connect(lambda: _on_date_opt_changed(self._interest_date_opt_start_only))
        self._interest_date_opt_end_only.stateChanged.connect(lambda: _on_date_opt_changed(self._interest_date_opt_end_only))
        self._interest_date_opt_neither.stateChanged.connect(lambda: _on_date_opt_changed(self._interest_date_opt_neither))

        def _on_basis_changed(changed_cb):
            other = self._interest_year_basis_365 if changed_cb == self._interest_year_basis_360 else self._interest_year_basis_360
            other.blockSignals(True)
            other.setChecked(False)
            other.blockSignals(False)
            if not changed_cb.isChecked():
                changed_cb.blockSignals(True)
                changed_cb.setChecked(True)
                changed_cb.blockSignals(False)

        self._interest_year_basis_360.stateChanged.connect(lambda: _on_basis_changed(self._interest_year_basis_360))
        self._interest_year_basis_365.stateChanged.connect(lambda: _on_basis_changed(self._interest_year_basis_365))

        def _on_lpr_mode_changed(changed_cb):
            other = self._interest_lpr_mode_fixed if changed_cb == self._interest_lpr_mode_seg else self._interest_lpr_mode_seg
            other.blockSignals(True)
            other.setChecked(False)
            other.blockSignals(False)
            if not changed_cb.isChecked():
                changed_cb.blockSignals(True)
                changed_cb.setChecked(True)
                changed_cb.blockSignals(False)
            self._interest_lpr_fixed_rate.setVisible(self._interest_lpr_mode_fixed.isChecked())

        self._interest_lpr_mode_seg.stateChanged.connect(lambda: _on_lpr_mode_changed(self._interest_lpr_mode_seg))
        self._interest_lpr_mode_fixed.stateChanged.connect(lambda: _on_lpr_mode_changed(self._interest_lpr_mode_fixed))

        def _on_lpr_term_changed(changed_cb):
            other = self._interest_lpr_term_5y if changed_cb == self._interest_lpr_term_1y else self._interest_lpr_term_1y
            other.blockSignals(True)
            other.setChecked(False)
            other.blockSignals(False)
            if not changed_cb.isChecked():
                changed_cb.blockSignals(True)
                changed_cb.setChecked(True)
                changed_cb.blockSignals(False)

        self._interest_lpr_term_1y.stateChanged.connect(lambda: _on_lpr_term_changed(self._interest_lpr_term_1y))
        self._interest_lpr_term_5y.stateChanged.connect(lambda: _on_lpr_term_changed(self._interest_lpr_term_5y))

        self._last_interest_result: Optional[Dict[str, Any]] = None
        return card

    def _build_liquidated_card(self) -> QWidget:
        card, layout = self._create_card(
            "违约金计算",
            "支持固定金额 + 按日比例或年化比例的混合计算，便于合同违约责任的快速测算。",
        )
        form = self._create_form()
        self._liquidated_base = self._make_money_spin()
        self._liquidated_daily_rate = self._make_rate_spin(maximum=100, decimals=4)
        self._liquidated_annual_rate = self._make_rate_spin(maximum=500, decimals=4)
        self._liquidated_days = self._make_int_spin(36500)
        self._liquidated_fixed = self._make_money_spin()
        form.addRow("计费基数", self._liquidated_base)
        form.addRow("按日比例（%）", self._liquidated_daily_rate)
        form.addRow("年化比例（%）", self._liquidated_annual_rate)
        form.addRow("期间天数", self._liquidated_days)
        form.addRow("固定违约金", self._liquidated_fixed)
        layout.addLayout(form)

        btn = self._make_button("计算违约金", accent=True)
        btn.clicked.connect(self._calculate_liquidated)
        layout.addWidget(btn, 0, Qt.AlignmentFlag.AlignRight)

        self._liquidated_result = self._make_result_browser()
        layout.addWidget(self._liquidated_result)
        return card

    def _build_occupation_card(self) -> QWidget:
        card, layout = self._create_card(
            "占有使用费",
            "适合房屋、场地、设备等占有使用费的简算。可按月、按日及固定金额合并计算。",
        )
        form = self._create_form()
        self._occupation_monthly = self._make_money_spin()
        self._occupation_months = self._make_money_spin(maximum=1200, decimals=2)
        self._occupation_daily = self._make_money_spin()
        self._occupation_days = self._make_int_spin(36500)
        self._occupation_fixed = self._make_money_spin()
        form.addRow("月度标准", self._occupation_monthly)
        form.addRow("月数", self._occupation_months)
        form.addRow("日标准", self._occupation_daily)
        form.addRow("天数", self._occupation_days)
        form.addRow("其他固定金额", self._occupation_fixed)
        layout.addLayout(form)

        btn = self._make_button("计算占有使用费", accent=True)
        btn.clicked.connect(self._calculate_occupation)
        layout.addWidget(btn, 0, Qt.AlignmentFlag.AlignRight)

        self._occupation_result = self._make_result_browser()
        layout.addWidget(self._occupation_result)
        return card

    def _build_delay_interest_card(self) -> QWidget:
        card, layout = self._create_card(
            "迟延履行利息",
            "按照一般债务利息 + 加倍部分债务利息拆分展示。加倍部分按日万分之一点七五的司法解释口径测算。",
        )
        form = self._create_form()
        self._delay_principal = self._make_money_spin()
        self._delay_days = self._make_int_spin(36500)
        self._delay_normal_rate = self._make_rate_spin(maximum=500, decimals=4)
        self._delay_day_basis = QComboBox()
        self._delay_day_basis.addItem("按 365 天", 365)
        self._delay_day_basis.addItem("按 360 天", 360)
        form.addRow("未履行金额", self._delay_principal)
        form.addRow("迟延天数", self._delay_days)
        form.addRow("一般债务年利率（%）", self._delay_normal_rate)
        form.addRow("一般利息折算口径", self._delay_day_basis)
        layout.addLayout(form)

        btn = self._make_button("计算迟延履行利息", accent=True)
        btn.clicked.connect(self._calculate_delay_interest)
        layout.addWidget(btn, 0, Qt.AlignmentFlag.AlignRight)

        self._delay_result = self._make_result_browser(108)
        layout.addWidget(self._delay_result)
        return card

    def _build_lawyer_fee_card(self) -> QWidget:
        card, layout = self._create_card(
            "律师费参考",
            "律师费并无法定全国统一收费标准，本工具按固定收费、计时收费、按标的比例三种常见约定模式测算，仅作约定阶段参考。",
        )
        form = self._create_form()
        self._lawyer_mode = QComboBox()
        self._lawyer_mode.addItem("固定收费", "fixed")
        self._lawyer_mode.addItem("计时收费", "hourly")
        self._lawyer_mode.addItem("按标的比例", "proportional")
        self._lawyer_mode.currentIndexChanged.connect(self._update_lawyer_mode_hint)
        self._lawyer_fixed = self._make_money_spin()
        self._lawyer_hour_rate = self._make_money_spin()
        self._lawyer_hours = self._make_money_spin(maximum=10000, decimals=2)
        self._lawyer_claim_amount = self._make_money_spin()
        self._lawyer_rate = self._make_rate_spin(maximum=100, decimals=4)
        form.addRow("收费模式", self._lawyer_mode)
        form.addRow("固定收费", self._lawyer_fixed)
        form.addRow("小时费率", self._lawyer_hour_rate)
        form.addRow("工时", self._lawyer_hours)
        form.addRow("标的金额", self._lawyer_claim_amount)
        form.addRow("比例（%）", self._lawyer_rate)
        layout.addLayout(form)

        self._lawyer_hint = QLabel("")
        self._lawyer_hint.setWordWrap(True)
        self._lawyer_hint.setStyleSheet(f"background:transparent;color:{COLORS['text_muted']};font-size:11px;")
        layout.addWidget(self._lawyer_hint)

        btn = self._make_button("计算律师费参考", accent=True)
        btn.clicked.connect(self._calculate_lawyer_fee)
        layout.addWidget(btn, 0, Qt.AlignmentFlag.AlignRight)

        self._lawyer_result = self._make_result_browser()
        layout.addWidget(self._lawyer_result)
        return card

    def _build_labor_compensation_card(self) -> QWidget:
        card, layout = self._create_card(
            "经济补偿金 / 赔偿金",
            "依据《劳动合同法》常用口径：经济补偿按工作年限计算；违法解除赔偿金通常按经济补偿的二倍计算。高工资情形支持 3 倍社平工资和 12 年封顶规则。",
        )
        form = self._create_form()
        self._labor_monthly_wage = self._make_money_spin()
        self._labor_local_avg = self._make_money_spin()
        self._labor_years = self._make_int_spin(80)
        self._labor_extra_months = self._make_int_spin(11)
        form.addRow("劳动者月工资", self._labor_monthly_wage)
        form.addRow("当地上年度职工月平均工资", self._labor_local_avg)
        form.addRow("完整工作年数", self._labor_years)
        form.addRow("不足一年另计月数", self._labor_extra_months)
        layout.addLayout(form)

        btn = self._make_button("计算补偿金 / 赔偿金", accent=True)
        btn.clicked.connect(self._calculate_labor_compensation)
        layout.addWidget(btn, 0, Qt.AlignmentFlag.AlignRight)

        self._labor_result = self._make_result_browser(110)
        layout.addWidget(self._labor_result)
        return card

    def _build_work_injury_card(self) -> QWidget:
        card, layout = self._create_card(
            "工伤赔偿",
            "一次性伤残补助金按《工伤保险条例》法定倍数计算；一次性医疗补助金、就业补助金等地方差异部分由你输入当地标准。工亡部分同时支持丧葬补助金、一次性工亡补助金和供养亲属抚恤金月额测算。",
        )

        grid = QGridLayout()
        grid.setHorizontalSpacing(18)
        grid.setVerticalSpacing(12)

        disability_card, disability_layout = self._create_card("伤残待遇", "")
        disability_form = self._create_form()
        self._injury_wage = self._make_money_spin()
        self._injury_level = self._make_int_spin(10)
        self._injury_level.setRange(1, 10)
        self._injury_local_medical = self._make_money_spin()
        self._injury_local_employment = self._make_money_spin()
        disability_form.addRow("本人月工资", self._injury_wage)
        disability_form.addRow("伤残等级（1-10）", self._injury_level)
        disability_form.addRow("当地一次性医疗补助金", self._injury_local_medical)
        disability_form.addRow("当地一次性就业补助金", self._injury_local_employment)
        disability_layout.addLayout(disability_form)
        disability_btn = self._make_button("计算伤残待遇", accent=True)
        disability_btn.clicked.connect(self._calculate_work_injury_disability)
        disability_layout.addWidget(disability_btn, 0, Qt.AlignmentFlag.AlignRight)
        self._injury_disability_result = self._make_result_browser(120)
        disability_layout.addWidget(self._injury_disability_result)

        death_card, death_layout = self._create_card("工亡待遇", "")
        death_form = self._create_form()
        self._death_employee_wage = self._make_money_spin()
        self._death_local_avg = self._make_money_spin()
        self._death_national_income = self._make_money_spin()
        self._death_spouse_count = self._make_int_spin(5)
        self._death_spouse_count.setValue(1)
        self._death_other_count = self._make_int_spin(10)
        self._death_extra_supported = self._make_int_spin(10)
        death_form.addRow("职工本人工资", self._death_employee_wage)
        death_form.addRow("统筹地区上年度月平均工资", self._death_local_avg)
        death_form.addRow("上一年度全国城镇居民人均可支配收入", self._death_national_income)
        death_form.addRow("配偶人数", self._death_spouse_count)
        death_form.addRow("其他供养亲属人数", self._death_other_count)
        death_form.addRow("孤寡 / 孤儿等加发人数", self._death_extra_supported)
        death_layout.addLayout(death_form)
        death_btn = self._make_button("计算工亡待遇", accent=True)
        death_btn.clicked.connect(self._calculate_work_injury_death)
        death_layout.addWidget(death_btn, 0, Qt.AlignmentFlag.AlignRight)
        self._injury_death_result = self._make_result_browser(120)
        death_layout.addWidget(self._injury_death_result)

        grid.addWidget(disability_card, 0, 0)
        grid.addWidget(death_card, 0, 1)
        layout.addLayout(grid)
        return card

    def _build_traffic_injury_card(self) -> QWidget:
        card, layout = self._create_card(
            "交通事故人身损害",
            "按《人身损害赔偿司法解释》的常见项目拆分计算。各地年度标准差异较大的项目，如残疾赔偿金、死亡赔偿金基数，请按当地最新标准手动输入。",
        )
        form = self._create_form()
        self._traffic_medical = self._make_money_spin()
        self._traffic_rehab = self._make_money_spin()
        self._traffic_followup = self._make_money_spin()
        self._traffic_meal = self._make_money_spin()
        self._traffic_nutrition = self._make_money_spin()
        self._traffic_nursing_days = self._make_int_spin(36500)
        self._traffic_nursing_daily = self._make_money_spin()
        self._traffic_lost_days = self._make_int_spin(36500)
        self._traffic_lost_daily = self._make_money_spin()
        self._traffic_transport = self._make_money_spin()
        self._traffic_accommodation = self._make_money_spin()
        self._traffic_age = self._make_int_spin(120)
        self._traffic_disability_percent = self._make_rate_spin(maximum=100, decimals=2)
        self._traffic_disability_base = self._make_money_spin()
        self._traffic_death_base = self._make_money_spin()
        self._traffic_assistive = self._make_money_spin()
        self._traffic_mental = self._make_money_spin()
        self._traffic_funeral = self._make_money_spin()
        self._traffic_dependents = self._make_money_spin()

        form.addRow("医疗费", self._traffic_medical)
        form.addRow("康复费", self._traffic_rehab)
        form.addRow("后续治疗费", self._traffic_followup)
        form.addRow("住院伙食补助费", self._traffic_meal)
        form.addRow("营养费", self._traffic_nutrition)
        form.addRow("护理天数", self._traffic_nursing_days)
        form.addRow("护理费日标准", self._traffic_nursing_daily)
        form.addRow("误工天数", self._traffic_lost_days)
        form.addRow("误工费日标准", self._traffic_lost_daily)
        form.addRow("交通费", self._traffic_transport)
        form.addRow("住宿费", self._traffic_accommodation)
        form.addRow("受害人年龄", self._traffic_age)
        form.addRow("伤残系数（%）", self._traffic_disability_percent)
        form.addRow("残疾赔偿金年标准", self._traffic_disability_base)
        form.addRow("死亡赔偿金年标准", self._traffic_death_base)
        form.addRow("辅助器具费", self._traffic_assistive)
        form.addRow("精神损害抚慰金", self._traffic_mental)
        form.addRow("丧葬费", self._traffic_funeral)
        form.addRow("被扶养人生活费", self._traffic_dependents)
        layout.addLayout(form)

        self._traffic_year_hint = QLabel("")
        self._traffic_year_hint.setWordWrap(True)
        self._traffic_year_hint.setStyleSheet(f"background:transparent;color:{COLORS['text_muted']};font-size:11px;")
        layout.addWidget(self._traffic_year_hint)

        btn = self._make_button("计算人身损害赔偿", accent=True)
        btn.clicked.connect(self._calculate_traffic_injury)
        layout.addWidget(btn, 0, Qt.AlignmentFlag.AlignRight)

        self._traffic_result = self._make_result_browser(130)
        layout.addWidget(self._traffic_result)
        return card

    def _build_procedural_limit_card(self) -> QWidget:
        card, layout = self._create_card(
            "审限计算",
            "按常见法定审限规则推算。遇公告、鉴定、中止等情形，实际审限可能依法顺延或扣除。",
        )
        form = self._create_form()
        self._limit_start = self._make_date_edit()
        self._limit_rule = QComboBox()
        for key, item in PROCEDURAL_LIMIT_RULES.items():
            self._limit_rule.addItem(str(item["label"]), key)
        self._limit_exclude_start = QCheckBox("次日起算")
        self._limit_exclude_start.setChecked(True)
        self._limit_move_weekend = QCheckBox("顺延至下一个工作日（仅排周末）")
        self._limit_move_weekend.setChecked(True)
        form.addRow("起算日期", self._limit_start)
        form.addRow("审限规则", self._limit_rule)
        form.addRow("", self._limit_exclude_start)
        form.addRow("", self._limit_move_weekend)
        layout.addLayout(form)

        btn = self._make_button("计算审限", accent=True)
        btn.clicked.connect(self._calculate_procedural_limit)
        layout.addWidget(btn, 0, Qt.AlignmentFlag.AlignRight)

        self._limit_result = self._make_result_browser()
        layout.addWidget(self._limit_result)
        return card

    def _build_date_offset_card(self) -> QWidget:
        card, layout = self._create_card(
            "日期 / 期日推算",
            "支持天、月、年混合推算。可选择次日起算，以及遇周末自动顺延。当前不自动识别法定节假日调休。",
        )
        form = self._create_form()
        self._date_offset_start = self._make_date_edit()
        self._date_offset_days = self._make_int_spin(36500)
        self._date_offset_months = self._make_int_spin(240)
        self._date_offset_years = self._make_int_spin(50)
        self._date_offset_exclude = QCheckBox("次日起算")
        self._date_offset_weekend = QCheckBox("顺延到下一个工作日（仅排周末）")
        self._date_offset_weekend.setChecked(True)
        form.addRow("起算日期", self._date_offset_start)
        form.addRow("增加天数", self._date_offset_days)
        form.addRow("增加月数", self._date_offset_months)
        form.addRow("增加年数", self._date_offset_years)
        form.addRow("", self._date_offset_exclude)
        form.addRow("", self._date_offset_weekend)
        layout.addLayout(form)

        btn = self._make_button("推算日期", accent=True)
        btn.clicked.connect(self._calculate_date_offset)
        layout.addWidget(btn, 0, Qt.AlignmentFlag.AlignRight)

        self._date_offset_result = self._make_result_browser()
        layout.addWidget(self._date_offset_result)
        return card

    def _build_cause_reference_card(self) -> QWidget:
        card, layout = self._create_card(
            "民事案由参考",
            "这里做的是高频案由速查与检索，便于起草、立案和内部分类；正式适用时建议结合最高法最新《民事案件案由规定》核对。",
        )
        top = QHBoxLayout()
        self._cause_search = self._make_line_edit("搜索常见案由，如：买卖、借款、交通事故、股东、劳动...")
        self._cause_search.textChanged.connect(self._refresh_cause_reference)
        top.addWidget(self._cause_search, 1)
        refresh_btn = self._make_button("刷新")
        refresh_btn.clicked.connect(self._refresh_cause_reference)
        top.addWidget(refresh_btn)
        layout.addLayout(top)

        self._cause_browser = self._make_result_browser(260)
        self._cause_browser.setOpenExternalLinks(True)
        layout.addWidget(self._cause_browser)
        return card

    def _build_reference_links_card(self) -> QWidget:
        card, layout = self._create_card(
            "法律导航链接",
            "以下链接均为常见法律检索、审判公开、执行公开和司法服务入口，可直接点击跳转。",
        )
        browser = self._make_result_browser(220)
        browser.setOpenExternalLinks(True)
        items = [
            f"<li><a href='{url}'>{title}</a></li>"
            for title, url in LEGAL_REFERENCE_LINKS
        ]
        browser.setHtml(
            "<h3 style='margin:0 0 8px 0;'>常用导航</h3>"
            "<ul style='margin-top:6px; line-height:1.8;'>"
            + "".join(items)
            + "</ul>"
            "<hr/>"
            "<p><b>说明</b>：诉讼费用、执行费、保全费、工伤、经济补偿、破产管理人报酬等计算，均应结合最新法律、司法解释、地方法院口径和本案具体情形复核。</p>"
        )
        layout.addWidget(browser)
        return card

    def _bind_defaults(self) -> None:
        self._update_litigation_mode_hint()
        self._update_lawyer_mode_hint()
        self._update_traffic_year_hint()
        self._traffic_age.valueChanged.connect(self._update_traffic_year_hint)
        self._clear_court_sms_state()

    def _update_litigation_mode_hint(self) -> None:
        mode = self._litigation_type.currentData()
        hints = {
            "property": "财产案件按标的额分段累进计算。",
            "divorce": "离婚案件基础收费法定区间 50-300 元；涉及财产分割超过 20 万元部分按 0.5% 计收。",
            "personality": "人格权纠纷基础收费法定区间 100-500 元；赔偿额超过 5 万元后再按比例加收。",
            "labor": "劳动争议案件每件通常交纳 10 元。",
            "administrative": "行政案件每件通常交纳 50 元。",
            "ip_no_amount": "知识产权无争议金额案件通常按 500-1000 元区间，由法院确定。",
            "jurisdiction": "管辖权异议案件通常按 50-100 元区间，由法院确定。",
        }
        defaults = {
            "property": 50,
            "divorce": 50,
            "personality": 100,
            "labor": 10,
            "administrative": 50,
            "ip_no_amount": 500,
            "jurisdiction": 50,
        }
        self._litigation_base.setValue(defaults.get(mode, 50))
        self._litigation_hint.setText(hints.get(mode, ""))

    def _update_lawyer_mode_hint(self) -> None:
        mode = self._lawyer_mode.currentData()
        text = {
            "fixed": "固定收费适合标准化业务或分阶段委托。",
            "hourly": "计时收费适合顾问、谈判、尽调等按工时计费的案件。",
            "proportional": "按标的比例常见于争议解决或风险代理参考，比例须结合当地监管和业务类型审慎适用。",
        }.get(mode, "")
        self._lawyer_hint.setText(text)

    def _update_traffic_year_hint(self) -> None:
        years = compensation_years_for_age(self._traffic_age.value())
        self._traffic_year_hint.setText(
            f"按当前年龄推算，残疾赔偿金 / 死亡赔偿金常见年限参考为 {years} 年；如当地司法口径不同，可自行按当地标准调整年标准数额。"
        )

    def _clear_court_sms_state(self) -> None:
        self._sms_parse_result = None
        self._sms_documents = []
        self._sms_matches = []
        self._sms_hearing_notices = []
        self._sms_download_dir = None
        self._last_added_court_sms_case_id = ""
        self._last_added_court_sms_deadline_id = ""
        self._last_added_court_sms_deadline_date = ""
        if hasattr(self, "_court_sms_input"):
            self._court_sms_input.clear()
        if hasattr(self, "_court_sms_summary"):
            self._court_sms_summary.setHtml(
                "<p><b>待处理</b>：请先粘贴法院短信。系统会优先识别案号、法院、收件人和链接参数。</p>"
            )
        if hasattr(self, "_court_sms_doc_tree"):
            self._court_sms_doc_tree.clear()
        if hasattr(self, "_court_sms_doc_hint"):
            self._court_sms_doc_hint.setText("尚未开始读取法院文书。点击“智能解析并读取”后会自动暂存文书。")
        if hasattr(self, "_court_sms_hearing_tree"):
            self._court_sms_hearing_tree.clear()
        if hasattr(self, "_court_sms_hearing_hint"):
            self._court_sms_hearing_hint.setText("点击“智能解析并读取”后，这里会自动识别庭审类文书。")
        if hasattr(self, "_court_sms_hearing_browser"):
            self._court_sms_hearing_browser.setHtml(
                "<p>尚未识别到庭审文书。若下载的 PDF 中包含传票或出庭通知书，这里会自动整理为可加入案件期限的提醒。</p>"
            )
        if hasattr(self, "_court_sms_hearing_result_hint"):
            self._court_sms_hearing_result_hint.setText("识别结果会跟随当前所选案件动态判断是否可直接加入期限。")
        if hasattr(self, "_court_sms_case_combo"):
            self._court_sms_case_combo.blockSignals(True)
            self._court_sms_case_combo.clear()
            self._court_sms_case_combo.addItem("请先读取短信文书", "")
            self._court_sms_case_combo.blockSignals(False)
            if self._court_sms_case_combo.lineEdit():
                self._court_sms_case_combo.lineEdit().clear()
        if hasattr(self, "_court_sms_case_hint"):
            self._court_sms_case_hint.setText("点击“智能解析并读取”后，系统会在这里给出建议案件。")
        self._court_sms_case_completer_model.setStringList([])
        self._court_sms_case_options = []
        self._court_sms_selected_case_id = ""
        if hasattr(self, "_court_sms_target_folder"):
            self._court_sms_target_folder.clear()
        if hasattr(self, "_court_sms_custom_save"):
            self._court_sms_custom_save.setChecked(False)
        if hasattr(self, "_court_sms_custom_dir"):
            self._court_sms_custom_dir.clear()
        if hasattr(self, "_court_sms_custom_dir_btn"):
            self._court_sms_custom_dir_btn.setEnabled(False)
        self._sync_court_sms_busy_state()
        self._refresh_court_sms_hearing_action_buttons(None)
        if hasattr(self, "_court_sms_match_browser"):
            self._court_sms_match_browser.setHtml(
                "<p>当前还没有自动匹配结果。点击“智能解析并读取”后，这里会显示建议案件和匹配依据。</p>"
            )
        if hasattr(self, "_court_sms_store_result"):
            self._court_sms_store_result.setHtml(
                "<p>点击“智能解析并读取”后，系统会自动读取并暂存文书，再由你确认存放到案件目录。</p>"
            )
        self._refresh_court_sms_hearing_preview()

    def _toggle_court_sms_custom_save_mode(self, enabled: bool) -> None:
        if hasattr(self, "_court_sms_case_combo"):
            self._court_sms_case_combo.setEnabled(not enabled)
        if hasattr(self, "_court_sms_custom_dir"):
            self._court_sms_custom_dir.setEnabled(enabled)
        if hasattr(self, "_court_sms_custom_dir_btn"):
            self._court_sms_custom_dir_btn.setEnabled(enabled)
        if hasattr(self, "_btn_unlink_court_sms_case"):
            self._btn_unlink_court_sms_case.setEnabled(not enabled)
        if not enabled and self._court_sms_case_options:
            self._populate_court_sms_case_combo(selected_case_id=self._court_sms_selected_case_id)
        self._update_court_sms_case_hint()
        self._refresh_court_sms_match_preview()

    def _on_unlink_court_sms_case(self) -> None:
        """点击'不关联案件'：勾选自定义保存，默认保存到桌面。"""
        if hasattr(self, "_court_sms_custom_save"):
            self._court_sms_custom_save.setChecked(True)
        desktop = str(Path.home() / "Desktop")
        if hasattr(self, "_court_sms_custom_dir"):
            self._court_sms_custom_dir.setText(desktop)

    def _toggle_hearing_field_edit(self, label: QLabel, line_edit: QLineEdit) -> None:
        """双击标签时切换为编辑模式。"""
        label.setVisible(False)
        line_edit.setText(label.text() if label.text() != "未识别" else "")
        line_edit.setVisible(True)
        line_edit.setFocus()
        line_edit.selectAll()

    def _finish_hearing_field_edit(self, label: QLabel, line_edit: QLineEdit) -> None:
        """编辑完成后切回显示模式。"""
        new_val = line_edit.text().strip()
        label.setText(new_val or "未识别")
        line_edit.setVisible(False)
        label.setVisible(True)

    def _sync_hearing_edit_fields_to_notice(self, notice: CourtSmsHearingNotice) -> None:
        """将可编辑字段中的当前值写回到 notice 对象。"""
        if not hasattr(self, "_hearing_edit_fields"):
            return
        _FIELD_ATTR = {
            "case_number": "case_number",
            "summoned_person": "summoned_person",
            "hearing_date": "hearing_date",
            "hearing_time": "hearing_time",
            "hearing_place": "hearing_place",
            "cause": "cause",
            "court_name": "court_name",
        }
        for key, attr in _FIELD_ATTR.items():
            lbl, edt = self._hearing_edit_fields[key]
            val = edt.text().strip() if edt.isVisible() else lbl.text().strip()
            if val == "未识别":
                val = ""
            setattr(notice, attr, val)

    def _choose_court_sms_custom_directory(self) -> None:
        selected = QFileDialog.getExistingDirectory(
            self,
            "选择自定义保存目录",
            self._court_sms_custom_dir.text().strip() or "",
        )
        if selected:
            self._court_sms_custom_dir.setText(selected)

    def _paste_court_sms_from_clipboard(self) -> None:
        clipboard = QApplication.clipboard()
        text = clipboard.text().strip()
        if not text:
            QMessageBox.information(self, "剪贴板为空", "当前剪贴板里没有可识别的短信内容。")
            return
        self._court_sms_input.setPlainText(text)

    # ------------------------------------------------------------------
    # 文件上传（PDF/图片）
    # ------------------------------------------------------------------

    def _on_court_sms_upload_files(self) -> None:
        """通过文件对话框选择 PDF/图片文件。"""
        if self._has_active_background_workers():
            QMessageBox.information(self, "请稍候", "当前仍有后台任务正在进行，请等待完成后再继续操作。")
            return
        paths, _ = QFileDialog.getOpenFileNames(
            self, "选择传票或法院文书文件", "",
            "法院文书 (*.pdf *.jpg *.jpeg *.png *.bmp *.tiff *.tif);;PDF 文件 (*.pdf);;图片文件 (*.jpg *.jpeg *.png *.bmp);;所有文件 (*)",
        )
        if paths:
            self._process_uploaded_court_sms_files(paths)

    def _process_uploaded_court_sms_files(self, file_paths: List[str]) -> None:
        """启动后台线程处理上传的文件。"""
        if self._has_active_background_workers():
            QMessageBox.information(self, "请稍候", "当前仍有后台任务正在进行，请等待完成后再继续操作。")
            return

        if not self._court_sms_available:
            QMessageBox.warning(self, "服务不可用", "法院短信服务依赖未就绪，无法解析文件。")
            return

        # 清空之前的 SMS 解析状态，但保留案件匹配信息
        self._sms_documents = []
        self._sms_hearing_notices = []
        self._sms_parse_result = None
        self._sms_download_dir = None
        self._court_sms_doc_tree.clear()
        self._court_sms_hearing_tree.clear()

        self._court_sms_summary.setHtml(
            "<p><b>状态</b>：正在识别上传的文件，请稍候…</p>"
        )
        self._court_sms_doc_hint.setText("正在识别上传的文件…")
        self._court_sms_hearing_hint.setText("正在识别庭审信息…")

        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
        self._sync_court_sms_busy_state()

        self._court_sms_upload_worker = _CourtSmsFileUploadWorker(
            self._court_sms_service, file_paths
        )
        self._court_sms_upload_worker.finished.connect(self._on_court_sms_file_upload_finished)
        self._court_sms_upload_worker.error.connect(self._on_court_sms_file_upload_error)
        self._court_sms_upload_worker.start()

    def _on_court_sms_file_upload_finished(
        self, documents: List, hearing_notices: List
    ) -> None:
        """文件上传解析完成回调。"""
        worker = self._court_sms_upload_worker
        self._court_sms_upload_worker = None
        QApplication.restoreOverrideCursor()

        self._sms_documents = documents
        self._sms_hearing_notices = hearing_notices

        # 更新文书列表
        self._populate_court_sms_documents()
        # 更新庭审识别
        self._populate_court_sms_hearing_notices()
        # 更新案件匹配
        self._refresh_case_matches()

        # 更新摘要
        file_count = len(documents)
        hearing_count = len(hearing_notices)
        summary_html = f"<p><b>上传文件</b>：{file_count} 份</p>"
        if hearing_count > 0:
            summary_html += f"<p><b>识别到庭审信息</b>：{hearing_count} 条</p>"
            for notice in hearing_notices:
                summary_html += (
                    f"<p>  - {notice.notice_type or '文书'}：{notice.case_number or '未知案号'} "
                    f"| 被传唤人：{notice.summoned_person or '未知'}"
                    f"| {notice.display_time} | {notice.hearing_place or '未知地点'}</p>"
                )
        else:
            summary_html += '<p><b>识别结果</b>：未发现传票或出庭通知。如文件包含开庭信息，可点击"重新识别"重试。</p>'
        self._court_sms_summary.setHtml(summary_html)

        self._court_sms_store_result.setHtml(
            '<p>文件已识别。请核对建议案件与保存目录后，点击"确认存放到案件"。</p>'
            if documents else
            "<p>未识别到有效文书文件。</p>"
        )

        self._sync_court_sms_busy_state()
        if worker is not None:
            worker.deleteLater()

    def _on_court_sms_file_upload_error(self, error_msg: str) -> None:
        """文件上传解析失败回调。"""
        worker = self._court_sms_upload_worker
        self._court_sms_upload_worker = None
        QApplication.restoreOverrideCursor()
        self._sync_court_sms_busy_state()
        if worker is not None:
            worker.deleteLater()
        self._court_sms_doc_hint.setText("文件识别失败。")
        self._court_sms_summary.setHtml(f"<p><b>识别失败</b>：{error_msg}</p>")
        QMessageBox.warning(self, "文件识别失败", error_msg)

    def _parse_court_sms_only(self) -> None:
        if self._has_active_background_workers():
            QMessageBox.information(self, "请稍候", "当前仍有后台任务正在进行，请等待完成后再继续操作。")
            return
        sms_text = self._court_sms_input.toPlainText().strip()

        # 检测输入是否为文件路径（拖拽文件场景）
        file_paths = self._extract_file_paths_from_text(sms_text)
        if file_paths:
            self._process_uploaded_court_sms_files(file_paths)
            return

        if not sms_text:
            QMessageBox.information(self, "请先输入内容", "请粘贴法院短信或拖拽传票 PDF/图片到输入框，再开始智能解析并读取。")
            return
        try:
            parsed = self._court_sms_service.parse_sms(sms_text)
        except Exception as exc:
            QMessageBox.warning(self, "短信解析失败", str(exc))
            return

        self._sms_parse_result = parsed
        self._sms_documents = []
        self._sms_download_dir = None
        self._sms_hearing_notices = []
        self._sms_matches = []
        self._court_sms_doc_tree.clear()
        self._court_sms_doc_hint.setText("短信已解析，正在自动读取并暂存法院文书，请稍候…")
        self._court_sms_target_folder.setText(self._court_sms_service.suggest_relative_folder(parsed))

        case_number = parsed.case_number or "未识别"
        recipient = parsed.recipient_name or "未识别"
        court_name = parsed.court_name or "未识别"
        link_html = (
            f"<a href='{parsed.link}'>{parsed.link}</a>"
            if parsed.link
            else "未识别"
        )
        self._court_sms_summary.setHtml(
            f"<p><b>法院</b>：{court_name}</p>"
            f"<p><b>收件人</b>：{recipient}</p>"
            f"<p><b>案号</b>：{case_number}</p>"
            f"<p><b>链接参数</b>：qdbh={parsed.qdbh} / sdbh={parsed.sdbh} / sdsin={parsed.sdsin}</p>"
            f"<p><b>短信链接</b>：{link_html}</p>"
            "<p><b>状态</b>：已解析成功，正在自动读取网页文书…</p>"
        )
        self._court_sms_store_result.setHtml(
            "<p>短信已解析，系统正在自动读取并暂存文书。读取完成后，这里会继续提示你确认存放到案件目录。</p>"
        )
        self._read_and_stage_court_documents(pre_parsed=parsed)
        self._populate_court_sms_hearing_notices()
        self._refresh_case_matches()

    @staticmethod
    def _extract_file_paths_from_text(text: str) -> List[str]:
        """从文本中提取存在的文件路径（拖拽文件场景）。

        支持格式：纯路径、file:// URL、富文本中的路径。
        """
        from urllib.parse import unquote, urlparse
        _FILE_EXTS = (".pdf", ".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif")
        paths = []
        for line in text.replace("\r\n", "\n").split("\n"):
            line = line.strip().strip('"').strip("'")
            if not line:
                continue
            # 处理 file:// URL
            if line.startswith("file://"):
                try:
                    parsed_url = urlparse(line)
                    line = unquote(parsed_url.path)
                except Exception:
                    pass
            p = Path(line)
            if p.exists() and p.is_file() and p.suffix.lower() in _FILE_EXTS:
                paths.append(str(p))
        return paths

    def _on_court_sms_read_finished(self, parsed, documents, download_dir, hearing_notices) -> None:
        """后台读取法院文书完成后的回调。"""
        worker = self._court_sms_read_worker
        self._court_sms_read_worker = None
        QApplication.restoreOverrideCursor()
        self._sms_parse_result = parsed
        self._sms_documents = documents
        self._sms_hearing_notices = hearing_notices
        self._sms_download_dir = download_dir
        self._court_sms_target_folder.setText(self._court_sms_service.suggest_relative_folder(parsed))

        self._populate_court_sms_documents()
        self._populate_court_sms_hearing_notices()
        self._refresh_case_matches()

        self._court_sms_summary.setHtml(
            f"<p><b>法院</b>：{parsed.court_name or '未识别'}</p>"
            f"<p><b>收件人</b>：{parsed.recipient_name or '未识别'}</p>"
            f"<p><b>案号</b>：{parsed.case_number or '未识别'}</p>"
            f"<p><b>文书数量</b>：{len(documents)} 份</p>"
            f"<p><b>暂存目录</b>：{download_dir}</p>"
        )
        self._court_sms_store_result.setHtml(
            "<p>文书已下载到本地暂存。请核对建议案件与保存目录后，再点击“确认存放到案件”。</p>"
        )
        self._sync_court_sms_busy_state()
        if worker is not None:
            worker.deleteLater()

    def _on_court_sms_read_error(self, error_msg: str) -> None:
        """后台读取法院文书失败后的回调。"""
        worker = self._court_sms_read_worker
        self._court_sms_read_worker = None
        QApplication.restoreOverrideCursor()
        self._sync_court_sms_busy_state()
        if worker is not None:
            worker.deleteLater()
        self._court_sms_doc_hint.setText("短信已解析，但自动读取文书失败。请检查短信链接后重新点击“智能解析并读取”。")
        self._court_sms_store_result.setHtml(
            "<p>自动读取暂存失败。你可以修改短信内容或稍后重新点击“智能解析并读取”重试。</p>"
        )
        QMessageBox.warning(self, "读取文书失败", error_msg)

    def _read_and_stage_court_documents(self, pre_parsed=None) -> None:
        if self._has_active_background_workers():
            QMessageBox.information(self, "正在处理中", "法院文书仍在后台读取或识别中，请稍候再试。")
            return
        sms_text = self._court_sms_input.toPlainText().strip()
        if not sms_text:
            QMessageBox.information(self, "请先粘贴短信", "先把法院短信完整粘贴到输入区，再开始智能解析并读取。")
            return

        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
        self._sync_court_sms_busy_state()
        if pre_parsed is not None:
            self._court_sms_doc_hint.setText("短信已解析，正在自动读取并暂存法院文书，请稍候…")
        else:
            self._court_sms_doc_hint.setText("正在读取并暂存法院文书，请稍候…")
        self._court_sms_read_worker = _CourtSmsReadWorker(self._court_sms_service, sms_text, pre_parsed)
        self._court_sms_read_worker.finished.connect(self._on_court_sms_read_finished)
        self._court_sms_read_worker.error.connect(self._on_court_sms_read_error)
        self._court_sms_read_worker.start()

    def _populate_court_sms_documents(self) -> None:
        self._court_sms_doc_tree.clear()
        for document in self._sms_documents:
            status = self._format_file_size(document.size_bytes) if document.local_path else "仅已读取链接"
            item = QTreeWidgetItem([
                document.name,
                document.court_name or "-",
                document.created_at[:19].replace("T", " ") if document.created_at else "-",
                status,
            ])
            item.setData(0, Qt.ItemDataRole.UserRole, document.local_path or document.download_url)
            item.setData(0, Qt.ItemDataRole.UserRole + 1, bool(document.local_path))
            self._court_sms_doc_tree.addTopLevelItem(item)

        count = len(self._sms_documents)
        if self._sms_download_dir:
            self._court_sms_doc_hint.setText(
                f"已暂存 {count} 份文书。双击可打开单个文件，或在右侧确认存放到案件目录。"
            )
        else:
            self._court_sms_doc_hint.setText(f"已读取 {count} 份文书。")

        for column in range(self._court_sms_doc_tree.columnCount()):
            self._court_sms_doc_tree.resizeColumnToContents(column)

    def _on_court_sms_analyze_finished(self, _parsed, _documents, _download_dir, hearing_notices) -> None:
        """后台识别庭审通知完成后的回调。"""
        worker = self._court_sms_analyze_worker
        self._court_sms_analyze_worker = None
        QApplication.restoreOverrideCursor()
        self._sms_hearing_notices = hearing_notices
        self._populate_court_sms_hearing_notices()
        self._sync_court_sms_busy_state()
        if worker is not None:
            worker.deleteLater()

    def _on_court_sms_analyze_error(self, error_msg: str) -> None:
        """后台识别庭审通知失败后的回调。"""
        worker = self._court_sms_analyze_worker
        self._court_sms_analyze_worker = None
        QApplication.restoreOverrideCursor()
        self._sync_court_sms_busy_state()
        if worker is not None:
            worker.deleteLater()
        QMessageBox.warning(self, "识别失败", error_msg)

    def _analyze_court_sms_hearing_notices(self) -> None:
        if self._has_active_background_workers():
            QMessageBox.information(self, "正在处理中", "后台任务尚未结束，请等待当前读取或识别完成。")
            return
        if not self._sms_documents:
            QMessageBox.information(self, "请先智能解析", "请先点击“智能解析并读取”，系统会自动读取并暂存法院文书后再识别庭审信息。")
            return
        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
        self._sync_court_sms_busy_state()
        self._court_sms_hearing_hint.setText("正在识别庭审类文书，请稍候…")
        self._court_sms_analyze_worker = _CourtSmsAnalyzeWorker(self._court_sms_service, self._sms_documents, self._sms_parse_result)
        self._court_sms_analyze_worker.finished.connect(self._on_court_sms_analyze_finished)
        self._court_sms_analyze_worker.error.connect(self._on_court_sms_analyze_error)
        self._court_sms_analyze_worker.start()

    def _has_running_worker(self, worker: Optional[QThread]) -> bool:
        return bool(worker is not None and worker.isRunning())

    def _has_active_background_workers(self) -> bool:
        return any(
            self._has_running_worker(worker)
            for worker in (
                self._court_sms_read_worker,
                self._court_sms_analyze_worker,
                self._court_sms_upload_worker,
                self._screenshot_merge_worker,
            )
        )

    def _sync_court_sms_busy_state(self) -> None:
        reading = self._has_running_worker(self._court_sms_read_worker)
        analyzing = self._has_running_worker(self._court_sms_analyze_worker)
        uploading = self._has_running_worker(self._court_sms_upload_worker)
        busy = reading or analyzing or uploading
        if hasattr(self, "_btn_parse_court_sms"):
            self._btn_parse_court_sms.setEnabled(not busy)
        if hasattr(self, "_btn_clear_court_sms"):
            self._btn_clear_court_sms.setEnabled(not busy)
        if hasattr(self, "_btn_upload_court_sms"):
            self._btn_upload_court_sms.setEnabled(not busy)
        if hasattr(self, "_btn_refresh_court_sms_hearing"):
            can_refresh = bool(self._sms_documents) and not busy
            self._btn_refresh_court_sms_hearing.setEnabled(can_refresh)

    def _populate_court_sms_hearing_notices(self) -> None:
        if not hasattr(self, "_court_sms_hearing_tree"):
            return

        self._court_sms_hearing_tree.clear()
        for index, notice in enumerate(self._sms_hearing_notices):
            item = QTreeWidgetItem([
                notice.document_name,
                notice.notice_type or "庭审文书",
                notice.summoned_person or "未识别",
                notice.display_time,
                notice.hearing_place or "未识别",
                self._court_sms_hearing_notice_status(notice)[0],
            ])
            item.setData(0, Qt.ItemDataRole.UserRole, index)
            self._court_sms_hearing_tree.addTopLevelItem(item)

        for column in range(self._court_sms_hearing_tree.columnCount()):
            self._court_sms_hearing_tree.resizeColumnToContents(column)

        if self._sms_hearing_notices:
            self._court_sms_hearing_tree.setCurrentItem(self._court_sms_hearing_tree.topLevelItem(0))
            self._court_sms_hearing_hint.setText(
                f"已识别到 {len(self._sms_hearing_notices)} 份庭审类文书。选择一项后可直接加入期限提醒。"
            )
        else:
            self._court_sms_hearing_hint.setText(
                "暂未识别到可加入期限的庭审文书。若 PDF 为扫描件或不含传票/出庭通知书文本，当前不会自动生成提醒。"
            )
        self._sync_court_sms_busy_state()
        self._refresh_court_sms_hearing_preview()

    def _refresh_case_matches(self) -> None:
        if not self._sms_parse_result and not self._sms_hearing_notices:
            # 没有 SMS 解析结果也没有庭审识别，尝试展示全部案件供手动选择
            self._populate_court_sms_case_combo_for_upload()
            return

        # 上传文件场景：从庭审通知构造合成 parse_result 进行匹配
        if not self._sms_parse_result and self._sms_hearing_notices:
            first = self._sms_hearing_notices[0]
            self._sms_parse_result = CourtSmsParseResult(
                raw_text="",
                link="",
                case_number=first.case_number or "",
                court_name=first.court_name or "",
            )

        cases = self._case_manager.get_all_cases()
        self._sms_matches = self._court_sms_service.match_cases(
            self._sms_parse_result,
            cases,
            preferred_case_id=self._initial_case_id,
        )

        selected_case_id = self._court_sms_selected_case_id or self._initial_case_id
        if not selected_case_id and self._sms_matches:
            selected_case_id = self._sms_matches[0].case_id
        self._court_sms_selected_case_id = selected_case_id or ""

        self._court_sms_case_options = []
        matched_ids = {match.case_id for match in self._sms_matches}
        for case in cases:
            case_id = str(case.get("id", "")).strip()
            if not case_id:
                continue
            display = f"{case.get('name', '')}  ·  案件项目"
            search_terms = [
                case.get("name", ""),
                str(case.get("path", "")).strip(),
                str(case.get("category", "")).strip(),
                " ".join(case.get("tags", []) or []),
            ]
            for field in case.get("info_fields", []) or []:
                search_terms.append(str(field.get("label", "")).strip())
                search_terms.append(str(field.get("value", "")).strip())
            for value in (case.get("variables", {}) or {}).values():
                search_terms.append(str(value).strip())

            self._court_sms_case_options.append({
                "case_id": case_id,
                "display": display,
                "search": " ".join(item for item in search_terms if item),
                "kind": "manual",
                "score": "0",
            })

        for match in reversed(self._sms_matches):
            case_record = self._case_manager.get_case(match.case_id)
            if not case_record:
                continue
            display = f"{match.case_name}  ·  匹配分 {match.score}"
            search_terms = [
                match.case_name,
                str(case_record.get("path", "")).strip(),
                " ".join(match.reasons),
            ]
            for field in case_record.get("info_fields", []) or []:
                search_terms.append(str(field.get("label", "")).strip())
                search_terms.append(str(field.get("value", "")).strip())
            for value in (case_record.get("variables", {}) or {}).values():
                search_terms.append(str(value).strip())

            self._court_sms_case_options = [
                option for option in self._court_sms_case_options if option["case_id"] != match.case_id
            ]
            self._court_sms_case_options.insert(0, {
                "case_id": match.case_id,
                "display": display,
                "search": " ".join(item for item in search_terms if item),
                "kind": "suggested",
                "score": str(match.score),
            })

        self._populate_court_sms_case_combo(selected_case_id=selected_case_id)
        self._refresh_court_sms_match_preview()

    def _populate_court_sms_case_combo_for_upload(self) -> None:
        """上传文件但没有匹配线索时，列出全部案件供手动选择。"""
        cases = self._case_manager.get_all_cases()
        self._court_sms_case_options = []
        for case in cases:
            case_id = str(case.get("id", "")).strip()
            if not case_id:
                continue
            self._court_sms_case_options.append({
                "case_id": case_id,
                "display": f"{case.get('name', '')}  ·  案件项目",
                "search": case.get("name", ""),
                "kind": "manual",
                "score": "0",
            })
        self._populate_court_sms_case_combo(selected_case_id=self._initial_case_id)
        self._court_sms_case_hint.setText(
            "上传文件未自动匹配到案件，请手动选择或搜索案件。"
        )

    def _populate_court_sms_case_combo(self, selected_case_id: str = "") -> None:
        combo = self._court_sms_case_combo
        combo.blockSignals(True)
        combo.clear()
        completion_items: List[str] = []
        if self._court_sms_case_options:
            for option in self._court_sms_case_options:
                combo.addItem(option["display"], option["case_id"])
                completion_items.append(option["display"])
        else:
            combo.addItem("未找到匹配案件项目", "")
            completion_items.append("未找到匹配案件项目")

        index = combo.findData(selected_case_id)
        if index < 0 and combo.count() > 0:
            index = 0
        combo.setCurrentIndex(index)
        combo.blockSignals(False)
        self._court_sms_case_completer_model.setStringList(completion_items)
        self._update_court_sms_case_hint()

    def _begin_court_sms_case_search(self) -> None:
        combo = self._court_sms_case_combo
        if combo.currentIndex() < 0:
            return
        line_edit = combo.lineEdit()
        combo.blockSignals(True)
        combo.setCurrentIndex(-1)
        combo.blockSignals(False)
        if line_edit:
            line_edit.clear()
            line_edit.setFocus()
        self._refresh_court_sms_match_preview()
        self._update_court_sms_case_hint()

    def _handle_court_sms_case_search_input(self, text: str) -> None:
        if self._court_sms_custom_save.isChecked():
            return
        normalized = text.strip()
        combo = self._court_sms_case_combo
        line_edit = combo.lineEdit()
        if normalized:
            combo.blockSignals(True)
            combo.setCurrentIndex(-1)
            combo.blockSignals(False)
            if line_edit and line_edit.text() != text:
                line_edit.setText(text)
                line_edit.setCursorPosition(len(text))
            self._refresh_court_sms_match_preview()

        completer = combo.completer()
        if not completer:
            return
        if normalized:
            filtered_options = [
                option["display"]
                for option in self._court_sms_case_options
                if normalized.lower() in option["display"].lower()
                or normalized.lower() in option["search"].lower()
            ]
            self._court_sms_case_completer_model.setStringList(filtered_options)
            completer.setCompletionPrefix(normalized)
            completer.complete()
        else:
            self._court_sms_case_completer_model.setStringList(
                [option["display"] for option in self._court_sms_case_options]
            )
        self._update_court_sms_case_hint()

    def _commit_court_sms_case_search_text(self) -> None:
        if self._court_sms_custom_save.isChecked():
            return
        line_edit = self._court_sms_case_combo.lineEdit()
        text = line_edit.text().strip() if line_edit else ""
        self._select_court_sms_case_from_text(text)

    def _select_court_sms_case_from_combo(self) -> None:
        case_id = str(self._court_sms_case_combo.currentData() or "").strip()
        self._court_sms_selected_case_id = case_id
        self._populate_court_sms_case_combo(selected_case_id=case_id)
        self._refresh_court_sms_match_preview()
        self._update_court_sms_case_hint()

    def _select_court_sms_case_from_text(self, text: str) -> None:
        display_text = str(text or "").strip()
        if not display_text:
            return
        index = self._court_sms_case_combo.findText(display_text, Qt.MatchFlag.MatchExactly)
        if index < 0:
            index = next(
                (
                    option_index
                    for option_index, option in enumerate(self._court_sms_case_options)
                    if option["display"] == display_text
                    or display_text.lower() in option["display"].lower()
                    or display_text.lower() in option["search"].lower()
                ),
                -1,
            )
        if index < 0:
            return

        self._court_sms_case_combo.setCurrentIndex(index)
        case_id = str(self._court_sms_case_combo.itemData(index) or "").strip()
        self._court_sms_selected_case_id = case_id
        self._refresh_court_sms_match_preview()
        self._update_court_sms_case_hint()

    def _refresh_court_sms_match_preview(self) -> None:
        if hasattr(self, "_court_sms_custom_save") and self._court_sms_custom_save.isChecked():
            custom_dir = self._court_sms_custom_dir.text().strip()
            if not custom_dir:
                self._court_sms_match_browser.setHtml(
                    "<p>当前已切换到 <b>自定义保存</b> 模式。请选择任意文件夹作为保存目标，确认后文书会复制到该目录。</p>"
                )
            else:
                relative_folder = self._court_sms_target_folder.text().strip()
                final_path = str(Path(custom_dir) / Path(relative_folder)) if relative_folder else custom_dir
                self._court_sms_match_browser.setHtml(
                    f"<p><b>保存方式</b>：自定义文件夹</p>"
                    f"<p><b>目标目录</b>：{custom_dir}</p>"
                    f"<p><b>最终路径</b>：{final_path}</p>"
                    "<p>说明：此模式下不会依赖案件项目，仅按你指定的目录保存。</p>"
                )
            self._refresh_court_sms_hearing_preview()
            return

        case_id = self._get_active_court_sms_case_id()
        if not case_id:
            search_text = self._get_court_sms_case_search_text()
            if search_text:
                self._court_sms_match_browser.setHtml(
                    "<p>正在搜索案件项目。请从下方候选结果中选择一个案件，再确认存放。</p>"
                )
            else:
                self._court_sms_match_browser.setHtml(
                    "<p>尚未选择案件。若自动匹配没有命中，可在下拉框中手动选择目标案件。</p>"
                )
            self._refresh_court_sms_hearing_preview()
            return

        selected_match = next((item for item in self._sms_matches if item.case_id == case_id), None)
        case_record = self._case_manager.get_case(case_id)
        if not case_record:
            self._court_sms_match_browser.setHtml("<p>当前案件已不存在，请重新选择。</p>")
            return

        reasons = (
            "".join(f"<li>{reason}</li>" for reason in selected_match.reasons)
            if selected_match
            else "<li>手动选择案件</li>"
        )
        self._court_sms_match_browser.setHtml(
            f"<p><b>案件名称</b>：{case_record.get('name', '')}</p>"
            f"<p><b>案件路径</b>：{case_record.get('path', '') or '未关联目录'}</p>"
            f"<p><b>期限状态</b>：{case_record.get('folder_status', '') or '未知'}</p>"
            f"<p><b>匹配依据</b>：</p><ul style='margin-top:4px;'>{reasons}</ul>"
        )
        self._refresh_court_sms_hearing_preview()

    def _open_court_sms_document(self, item: QTreeWidgetItem, _column: int) -> None:
        target = str(item.data(0, Qt.ItemDataRole.UserRole) or "").strip()
        if not target:
            return
        if bool(item.data(0, Qt.ItemDataRole.UserRole + 1)):
            QDesktopServices.openUrl(QUrl.fromLocalFile(target))
        else:
            QDesktopServices.openUrl(QUrl(target))

    def _open_court_sms_staging_folder(self) -> None:
        if not self._sms_download_dir or not self._sms_download_dir.exists():
            QMessageBox.information(self, "尚无暂存目录", "请先点击“智能解析并读取”，系统会自动读取并暂存法院文书。")
            return
        QDesktopServices.openUrl(QUrl.fromLocalFile(str(self._sms_download_dir)))

    def _open_court_sms_saved_folder(self) -> None:
        if not self._last_store_dir or not self._last_store_dir.exists():
            QMessageBox.information(self, "尚无已存入目录", "请先确认存放法院文书。")
            return
        QDesktopServices.openUrl(QUrl.fromLocalFile(str(self._last_store_dir)))

    def _confirm_store_court_documents(self) -> None:
        if not self._sms_documents:
            QMessageBox.information(self, "请先智能解析", "当前还没有已暂存的法院文书。请先点击“智能解析并读取”。")
            return

        relative_folder = self._court_sms_target_folder.text().strip()
        use_custom_dir = self._court_sms_custom_save.isChecked() if hasattr(self, "_court_sms_custom_save") else False

        if use_custom_dir:
            custom_dir = self._court_sms_custom_dir.text().strip()
            if not custom_dir:
                QMessageBox.information(self, "请选择目录", "请先选择要保存到的任意文件夹。")
                return
            final_target = str(Path(custom_dir) / Path(relative_folder)) if relative_folder else custom_dir
            confirm_text = (
                f"即将把 {len(self._sms_documents)} 份文书复制到自定义目录：\n{final_target}\n\n是否继续？"
            )
            target_label = "自定义目录"
            target_name = custom_dir
        else:
            case_id = self._get_active_court_sms_case_id()
            if not case_id:
                QMessageBox.information(self, "请选择案件", "请先选择要存放到哪个案件项目。")
                return

            case_record = self._case_manager.get_case(case_id)
            if not case_record:
                QMessageBox.warning(self, "案件不存在", "当前选择的案件已不存在，请重新选择。")
                return

            if not relative_folder:
                QMessageBox.information(self, "请输入目录", "请输入案件内的目标保存目录。")
                return

            confirm_text = (
                f"即将把 {len(self._sms_documents)} 份文书复制到：\n{case_record.get('name', '')}\n\n"
                f"相对目录：{relative_folder}\n\n是否继续？"
            )
            target_label = "案件"
            target_name = case_record.get("name", "")

        confirm = QMessageBox.question(self, "确认存放法院文书", confirm_text)
        if confirm != QMessageBox.StandardButton.Yes:
            return

        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
        try:
            if use_custom_dir:
                saved_paths = self._court_sms_service.save_documents_to_directory(
                    self._court_sms_custom_dir.text().strip(),
                    self._sms_documents,
                    relative_folder,
                )
            else:
                saved_paths = self._court_sms_service.save_documents_to_case(case_record, self._sms_documents, relative_folder)
        except Exception as exc:
            QMessageBox.warning(self, "保存失败", str(exc))
            return
        finally:
            QApplication.restoreOverrideCursor()

        if saved_paths:
            self._last_store_dir = Path(saved_paths[0]).parent
        self._court_sms_store_result.setHtml(
            f"<p><b>已成功存放</b>：{len(saved_paths)} 份文书</p>"
            f"<p><b>{target_label}</b>：{target_name}</p>"
            f"<p><b>目录</b>：{self._last_store_dir if self._last_store_dir else '-'}</p>"
        )
        QMessageBox.information(
            self,
            "存放完成",
            f"已将 {len(saved_paths)} 份法院文书复制到{target_label}。",
        )

    def _get_selected_court_sms_hearing_notice(self) -> Optional[CourtSmsHearingNotice]:
        if not hasattr(self, "_court_sms_hearing_tree"):
            return None
        item = self._court_sms_hearing_tree.currentItem()
        if not item:
            return self._sms_hearing_notices[0] if self._sms_hearing_notices else None
        index = item.data(0, Qt.ItemDataRole.UserRole)
        if isinstance(index, int) and 0 <= index < len(self._sms_hearing_notices):
            return self._sms_hearing_notices[index]
        return None

    def _find_existing_court_sms_deadline(self, case_id: str, notice: CourtSmsHearingNotice) -> Optional[Dict[str, Any]]:
        case = self._case_manager.get_case(case_id)
        if not case:
            return None

        for deadline in case.get("deadlines", []) or []:
            if notice.document_name and str(deadline.get("source_document_name", "")).strip() == notice.document_name:
                return deadline
            if notice.document_path and str(deadline.get("source_document_path", "")).strip() == notice.document_path:
                return deadline
            if (
                str(deadline.get("date", "")).strip() == notice.hearing_date
                and str(deadline.get("time", "")).strip() == (notice.hearing_time or "09:00")
                and str(deadline.get("source_case_number", "")).strip() == notice.case_number
            ):
                return deadline
        return None

    def _find_existing_global_court_sms_deadline(self, notice: CourtSmsHearingNotice) -> Optional[Dict[str, Any]]:
        """检查全局期限中是否已存在相同的庭审提醒。"""
        for deadline in self._case_manager.get_global_deadlines():
            if notice.document_name and str(deadline.get("source_document_name", "")).strip() == notice.document_name:
                return deadline
            if notice.document_path and str(deadline.get("source_document_path", "")).strip() == notice.document_path:
                return deadline
            if (
                str(deadline.get("date", "")).strip() == notice.hearing_date
                and str(deadline.get("time", "")).strip() == (notice.hearing_time or "09:00")
                and str(deadline.get("source_case_number", "")).strip() == notice.case_number
            ):
                return deadline
        return None

    def _court_sms_hearing_notice_status(self, notice: CourtSmsHearingNotice) -> Tuple[str, str]:
        active_case_id = self._get_active_court_sms_case_id()
        if active_case_id:
            existing = self._find_existing_court_sms_deadline(active_case_id, notice)
            if existing:
                notice.added_case_id = active_case_id
                notice.added_deadline_id = str(existing.get("id", "")).strip()
                return "已在期限中", COLORS["success"]
            if notice.added_case_id == active_case_id and notice.added_deadline_id:
                return "已加入", COLORS["success"]
            if notice.added_case_id and notice.added_case_id != active_case_id:
                return "已加入其他案件", COLORS["warning"]
            return "可加入", COLORS["accent"]

        # 未关联案件：检查全局期限
        if notice.added_deadline_id and not notice.added_case_id:
            return "已加入", COLORS["success"]
        existing_global = self._find_existing_global_court_sms_deadline(notice)
        if existing_global:
            notice.added_deadline_id = str(existing_global.get("id", "")).strip()
            return "已在期限中", COLORS["success"]
        return "可加入", COLORS["accent"]

    def _refresh_court_sms_hearing_preview(self) -> None:
        if not hasattr(self, "_court_sms_hearing_browser"):
            return

        notice = self._get_selected_court_sms_hearing_notice()
        if not notice:
            self._court_sms_hearing_browser.setHtml(
                "<p>尚未识别到庭审文书。点击“智能解析并读取”后，若文书中包含传票或出庭通知书，这里会展示可加入期限的开庭提醒。</p>"
            )
            if hasattr(self, "_hearing_edit_frame"):
                self._hearing_edit_frame.hide()
            if hasattr(self, "_btn_add_court_sms_hearing_deadline"):
                self._btn_add_court_sms_hearing_deadline.setEnabled(False)
            self._refresh_court_sms_hearing_action_buttons(None)
            return

        status_text, status_color = self._court_sms_hearing_notice_status(notice)
        detail_lines = [
            f"<p><b>文书</b>：{notice.document_name}</p>",
            f"<p><b>类型</b>：{notice.notice_type or '庭审文书'}</p>",
            f"<p><b>当前状态</b>：<span style='color:{status_color}; font-weight:700;'>{status_text}</span></p>",
        ]
        if notice.signer:
            detail_lines.append(f"<p><b>签发人</b>：{notice.signer}</p>")
        if notice.contact_person or notice.contact_phone:
            contact_text = notice.contact_person or "未识别"
            if notice.contact_phone:
                contact_text = f"{contact_text} / {notice.contact_phone}" if notice.contact_person else notice.contact_phone
            detail_lines.append(f"<p><b>联系人</b>：{contact_text}</p>")
        self._court_sms_hearing_browser.setHtml("".join(detail_lines))

        # 填充可编辑字段
        if hasattr(self, "_hearing_edit_frame"):
            field_values = {
                "case_number": notice.case_number or "",
                "summoned_person": notice.summoned_person or "",
                "hearing_date": notice.hearing_date or "",
                "hearing_time": notice.hearing_time or "",
                "hearing_place": notice.hearing_place or "",
                "cause": notice.cause or "",
                "court_name": notice.court_name or "",
            }
            for key, (lbl, edt) in self._hearing_edit_fields.items():
                val = field_values.get(key, "")
                lbl.setText(val or "未识别")
                lbl.setVisible(True)
                edt.setText(val)
                edt.setVisible(False)
            self._hearing_edit_frame.show()

        can_add = (
            status_text not in {"已在期限中", "已加入"}
            and not self._has_active_background_workers()
        )
        if hasattr(self, "_btn_add_court_sms_hearing_deadline"):
            self._btn_add_court_sms_hearing_deadline.setEnabled(can_add)
        self._refresh_court_sms_hearing_action_buttons(notice)
        if hasattr(self, "_court_sms_hearing_result_hint"):
            if can_add:
                self._court_sms_hearing_result_hint.setText(
                    "核对无误后，可把当前识别结果直接加入期限提醒。"
                )
            else:
                self._court_sms_hearing_result_hint.setText(
                    f"当前状态：{status_text}。如需加入期限，请先确认目标案件并避免重复添加。"
                )

        if hasattr(self, "_court_sms_hearing_tree"):
            for row in range(self._court_sms_hearing_tree.topLevelItemCount()):
                item = self._court_sms_hearing_tree.topLevelItem(row)
                index = item.data(0, Qt.ItemDataRole.UserRole)
                if not isinstance(index, int) or not (0 <= index < len(self._sms_hearing_notices)):
                    continue
                item_notice = self._sms_hearing_notices[index]
                item.setText(4, self._court_sms_hearing_notice_status(item_notice)[0])

    def _resolve_court_sms_hearing_navigation_target(
        self,
        notice: Optional[CourtSmsHearingNotice] = None,
    ) -> Tuple[str, str, str]:
        selected_notice = notice or self._get_selected_court_sms_hearing_notice()
        if selected_notice:
            case_id = str(selected_notice.added_case_id or "").strip()
            deadline_id = str(selected_notice.added_deadline_id or "").strip()
            deadline_date = str(selected_notice.hearing_date or "").strip()
            active_case_id = self._get_active_court_sms_case_id()
            if not case_id and active_case_id:
                existing = self._find_existing_court_sms_deadline(active_case_id, selected_notice)
                if existing:
                    case_id = active_case_id
                    deadline_id = str(existing.get("id", "")).strip()
                    deadline_date = str(existing.get("date", "") or deadline_date).strip()
                    selected_notice.added_case_id = case_id
                    selected_notice.added_deadline_id = deadline_id
            if case_id or deadline_date:
                return case_id, deadline_id, deadline_date

        return (
            self._last_added_court_sms_case_id,
            self._last_added_court_sms_deadline_id,
            self._last_added_court_sms_deadline_date,
        )

    def _refresh_court_sms_hearing_action_buttons(
        self,
        notice: Optional[CourtSmsHearingNotice] = None,
    ) -> None:
        target_case_id, _deadline_id, target_date = self._resolve_court_sms_hearing_navigation_target(notice)
        if hasattr(self, "_btn_open_court_sms_case"):
            self._btn_open_court_sms_case.setEnabled(bool(target_case_id))
        if hasattr(self, "_btn_open_court_sms_calendar"):
            self._btn_open_court_sms_calendar.setEnabled(True)

    def _add_selected_court_sms_hearing_deadline(self) -> None:
        notice = self._get_selected_court_sms_hearing_notice()
        if not notice:
            QMessageBox.information(self, "暂无可加入事项", "当前没有识别到可加入期限的庭审文书。")
            return

        # 将用户编辑过的字段值写回 notice
        self._sync_hearing_edit_fields_to_notice(notice)

        case_id = self._get_active_court_sms_case_id()
        if case_id:
            # 加入关联案件的期限
            case_record = self._case_manager.get_case(case_id)
            if not case_record:
                QMessageBox.warning(self, "案件不存在", "当前选择的案件已不存在，请重新选择。")
                return

            existing = self._find_existing_court_sms_deadline(case_id, notice)
            if existing:
                notice.added_case_id = case_id
                notice.added_deadline_id = str(existing.get("id", "")).strip()
                self._last_added_court_sms_case_id = case_id
                self._last_added_court_sms_deadline_id = str(existing.get("id", "")).strip()
                self._last_added_court_sms_deadline_date = str(existing.get("date", "") or notice.hearing_date or "").strip()
                self._refresh_court_sms_hearing_preview()
                QMessageBox.information(self, "已存在", "该庭审提醒已在当前案件的期限中，无需重复添加。")
                return

            deadline_data = self._court_sms_service.build_deadline_from_hearing_notice(notice)
            deadline_id = self._case_manager.add_deadline(case_id, deadline_data)
            if not deadline_id:
                QMessageBox.warning(self, "添加失败", "写入案件期限提醒时出现问题。")
                return

            notice.added_case_id = case_id
            notice.added_deadline_id = deadline_id
            self._last_added_court_sms_case_id = case_id
            self._last_added_court_sms_deadline_id = deadline_id
            self._last_added_court_sms_deadline_date = str(deadline_data.get("date", "") or notice.hearing_date or "").strip()
            self._refresh_court_sms_hearing_preview()
            self._notify_case_deadline_updated(case_id)
            QMessageBox.information(
                self,
                "已加入案件期限",
                f"已将“{deadline_data.get('title', '开庭安排')}”加入到案件“{case_record.get('name', '')}”的期限提醒。",
            )
            return

        # 未关联案件：加入全局期限
        existing_global = self._find_existing_global_court_sms_deadline(notice)
        if existing_global:
            notice.added_deadline_id = str(existing_global.get("id", "")).strip()
            self._last_added_court_sms_deadline_id = notice.added_deadline_id
            self._last_added_court_sms_deadline_date = str(existing_global.get("date", "") or notice.hearing_date or "").strip()
            self._refresh_court_sms_hearing_preview()
            QMessageBox.information(self, "已存在", "该庭审提醒已在全局期限中，无需重复添加。")
            return

        deadline_data = self._court_sms_service.build_deadline_from_hearing_notice(notice)
        deadline_id = self._case_manager.add_global_deadline(deadline_data)
        if not deadline_id:
            QMessageBox.warning(self, "添加失败", "写入全局期限提醒时出现问题。")
            return

        notice.added_deadline_id = deadline_id
        self._last_added_court_sms_deadline_id = deadline_id
        self._last_added_court_sms_deadline_date = str(deadline_data.get("date", "") or notice.hearing_date or "").strip()
        self._refresh_court_sms_hearing_preview()
        QMessageBox.information(
            self,
            "已加入期限",
            f"已将“{deadline_data.get('title', '开庭安排')}”加入到全局期限提醒（未关联案件）。",
        )

    def _open_court_sms_case_from_hearing(self) -> None:
        case_id, _deadline_id, _deadline_date = self._resolve_court_sms_hearing_navigation_target()
        if not case_id:
            QMessageBox.information(self, "暂无案件可跳转", "请先把当前庭审提醒加入案件期限，或选择已匹配的案件。")
            return
        self.navigate_to_case_requested.emit(case_id)
        self.accept()

    def _open_court_sms_calendar_from_hearing(self) -> None:
        _case_id, _deadline_id, deadline_date = self._resolve_court_sms_hearing_navigation_target()
        self.navigate_to_calendar_requested.emit(deadline_date)
        self.accept()

    def _notify_case_deadline_updated(self, case_id: str) -> None:
        parent = self.parent()
        visited = 0
        while parent is not None and visited < 8:
            if (
                hasattr(parent, "_case_id")
                and str(getattr(parent, "_case_id", "")).strip() == case_id
                and hasattr(parent, "_refresh_case_from_store")
            ):
                try:
                    parent._refresh_case_from_store()
                except Exception:
                    pass
                return
            parent = parent.parent() if hasattr(parent, "parent") else None
            visited += 1

    def _get_court_sms_case_search_text(self) -> str:
        line_edit = self._court_sms_case_combo.lineEdit()
        if not line_edit:
            return ""
        text = line_edit.text().strip()
        current_text = self._court_sms_case_combo.currentText().strip()
        current_case_id = str(self._court_sms_case_combo.currentData() or "").strip()
        if current_case_id and text == current_text:
            return ""
        return text

    def _get_active_court_sms_case_id(self) -> str:
        if getattr(self, "_court_sms_custom_save", None) and self._court_sms_custom_save.isChecked():
            return ""
        current_case_id = str(self._court_sms_case_combo.currentData() or "").strip()
        if current_case_id:
            return current_case_id
        if not self._get_court_sms_case_search_text():
            return self._court_sms_selected_case_id
        return ""

    def _update_court_sms_case_hint(self) -> None:
        if not hasattr(self, "_court_sms_case_hint"):
            return
        if getattr(self, "_court_sms_custom_save", None) and self._court_sms_custom_save.isChecked():
            self._court_sms_case_hint.setText("当前已切换为自定义保存，文书不会关联到案件项目。")
            return
        if not self._court_sms_case_options:
            self._court_sms_case_hint.setText("点击“智能解析并读取”后，系统会在这里给出建议案件。")
            return

        search_text = self._get_court_sms_case_search_text()
        total_count = len(self._court_sms_case_options)
        if search_text:
            matched_count = sum(
                1
                for option in self._court_sms_case_options
                if search_text.lower() in option["display"].lower()
                or search_text.lower() in option["search"].lower()
            )
            self._court_sms_case_hint.setText(
                f"搜索“{search_text}”：匹配到 {matched_count} 个案件项目。可回车确认，或直接点候选项。"
            )
            return

        active_case_id = self._get_active_court_sms_case_id()
        matched_case = next((item for item in self._sms_matches if item.case_id == active_case_id), None)
        if matched_case:
            self._court_sms_case_hint.setText(
                f"已带出最匹配案件，匹配分 {matched_case.score}。点输入框可改搜，点右侧箭头可查看全部 {total_count} 个案件。"
            )
            return
        if active_case_id:
            self._court_sms_case_hint.setText(
                f"当前为手动选择案件。点输入框可重新搜索，点右侧箭头可查看全部 {total_count} 个案件。"
            )
            return
        self._court_sms_case_hint.setText(
            f"点输入框直接搜索案件，点右侧箭头查看全部 {total_count} 个案件。"
        )

    def _refresh_cause_reference(self) -> None:
        keyword = self._cause_search.text().strip().lower() if hasattr(self, "_cause_search") else ""
        sections: List[str] = []
        for section, causes in COMMON_CIVIL_CAUSES.items():
            filtered = [item for item in causes if not keyword or keyword in item.lower() or keyword in section.lower()]
            if keyword and keyword not in section.lower() and not filtered:
                continue
            body = "".join(f"<li>{item}</li>" for item in filtered)
            sections.append(
                f"<h4 style='margin:10px 0 4px 0;'>{section}</h4><ul style='margin:0 0 8px 16px; line-height:1.7;'>{body}</ul>"
            )
        if not sections:
            sections.append("<p>未找到匹配项，可换关键词或直接查看最高法最新案由规定原文。</p>")
        sections.append("<p><a href='https://www.court.gov.cn/'>到最高人民法院官网检索最新《民事案件案由规定》及修改决定</a></p>")
        self._cause_browser.setHtml("".join(sections))

    def _calculate_litigation_fee(self) -> None:
        mode = self._litigation_type.currentData()
        amount = self._litigation_amount.value()
        base = self._litigation_base.value()
        if mode == "property":
            fee = calculate_property_litigation_fee(amount)
        elif mode == "divorce":
            fee = calculate_divorce_litigation_fee(base, amount)
        elif mode == "personality":
            fee = calculate_personality_litigation_fee(base, amount)
        elif mode == "labor":
            fee = money("10")
        elif mode == "administrative":
            fee = money("50")
        else:
            fee = money(base)
        self._litigation_result.setHtml(
            self._format_money_html("诉讼费用参考", fee)
            + "<p>提示：存在法定区间的案件类型，应以受理法院最终核定金额为准。</p>"
        )

    def _calculate_apply_fee(self) -> None:
        fee_type = self._apply_fee_type.currentData()
        if fee_type == "execution":
            fee = calculate_execution_fee(
                self._apply_amount.value(),
                no_amount_mode=self._apply_no_amount.isChecked(),
                base_fee=self._apply_fixed.value(),
            )
        else:
            fee = calculate_preservation_fee(
                self._apply_amount.value(),
                no_amount_mode=self._apply_no_amount.isChecked(),
                base_fee=self._apply_fixed.value(),
            )
        title = "申请执行费参考" if fee_type == "execution" else "财产保全费参考"
        self._apply_result.setHtml(self._format_money_html(title, fee))

    def _calculate_bankruptcy_fee(self) -> None:
        fee = calculate_bankruptcy_administrator_fee(
            self._bankruptcy_assets.value(),
            adjustment_factor=self._bankruptcy_factor.value(),
        )
        self._bankruptcy_result.setHtml(
            self._format_money_html("破产管理人报酬参考", fee)
            + "<p>提示：本结果按分段上限口径计算，再乘调整系数；实际报酬由人民法院裁定。</p>"
        )

    def _calculate_interest(self) -> None:
        principal = self._interest_principal.value()
        start = self._interest_start_date.date().toString("yyyy-MM-dd")
        end = self._interest_end_date.date().toString("yyyy-MM-dd")

        include_start = self._interest_date_opt_both.isChecked() or self._interest_date_opt_start_only.isChecked()
        include_end = self._interest_date_opt_both.isChecked() or self._interest_date_opt_end_only.isChecked()

        day_basis = 360 if self._interest_year_basis_360.isChecked() else 365

        term = "1y" if self._interest_lpr_term_1y.isChecked() else "5y"
        lpr_mode = "segment" if self._interest_lpr_mode_seg.isChecked() else "fixed"
        fixed_rate = self._interest_lpr_fixed_rate.value()
        adjust_mode = str(self._interest_lpr_adjust_combo.currentData() or "multiple")
        adjust_value = self._interest_lpr_adjust_value.value()

        try:
            result = calculate_lpr_interest(
                principal=principal,
                start_date=start,
                end_date=end,
                term=term,
                day_basis=day_basis,
                include_start=include_start,
                include_end=include_end,
                lpr_mode=lpr_mode,
                fixed_rate=fixed_rate,
                adjust_mode=adjust_mode,
                adjust_value=adjust_value,
            )
        except Exception as e:
            self._interest_result.setHtml(f"<p style='color:red;'>计算出错：{e}</p>")
            self._interest_export_btn.setEnabled(False)
            self._last_interest_result = None
            return

        self._last_interest_result = result
        self._last_interest_result["principal"] = principal

        total = result["total_interest"]
        days = result["total_days"]
        segments = result.get("segments", [])

        html = f"""
        <h3 style="margin:0 0 10px 0;color:{COLORS['text_primary']};">利息计算结果</h3>
        <table style="width:100%;border-collapse:collapse;font-size:13px;margin-bottom:10px;">
          <tr style="background:{COLORS['surface_1']};">
            <th style="padding:6px 8px;text-align:left;border-bottom:1px solid {COLORS['border']};">项目</th>
            <th style="padding:6px 8px;text-align:right;border-bottom:1px solid {COLORS['border']};">数值</th>
          </tr>
          <tr>
            <td style="padding:6px 8px;border-bottom:1px solid {COLORS['border']};">计算基数</td>
            <td style="padding:6px 8px;text-align:right;border-bottom:1px solid {COLORS['border']};">¥ {principal:,.2f}</td>
          </tr>
          <tr style="background:#f9fafb;">
            <td style="padding:6px 8px;border-bottom:1px solid {COLORS['border']};"><b>总利息</b></td>
            <td style="padding:6px 8px;text-align:right;border-bottom:1px solid {COLORS['border']};color:{COLORS['accent']};"><b>¥ {total:,.2f}</b></td>
          </tr>
          <tr>
            <td style="padding:6px 8px;border-bottom:1px solid {COLORS['border']};">计息天数</td>
            <td style="padding:6px 8px;text-align:right;border-bottom:1px solid {COLORS['border']};">{days} 天</td>
          </tr>
        </table>
        """

        if segments:
            html += f'<h4 style="margin:10px 0 6px 0;color:{COLORS["text_primary"]};">分段明细</h4>'
            html += '<table style="width:100%;border-collapse:collapse;font-size:12px;">'
            html += f'<tr style="background:{COLORS["surface_1"]};"><th style="padding:6px 8px;text-align:left;border-bottom:1px solid {COLORS["border"]};">起止日期</th><th style="padding:6px 8px;text-align:right;border-bottom:1px solid {COLORS["border"]};">天数</th><th style="padding:6px 8px;text-align:right;border-bottom:1px solid {COLORS["border"]};">利率</th><th style="padding:6px 8px;text-align:right;border-bottom:1px solid {COLORS["border"]};">利息</th></tr>'
            for i, seg in enumerate(segments):
                bg = "background:#f9fafb;" if i % 2 == 1 else ""
                html += f'<tr style="{bg}"><td style="padding:6px 8px;border-bottom:1px solid {COLORS["border"]};">{seg["start"]} ~ {seg["end"]}</td><td style="padding:6px 8px;text-align:right;border-bottom:1px solid {COLORS["border"]};">{seg["days"]}</td><td style="padding:6px 8px;text-align:right;border-bottom:1px solid {COLORS["border"]};">{seg["rate"]:.4f}%</td><td style="padding:6px 8px;text-align:right;border-bottom:1px solid {COLORS["border"]};">¥ {seg["interest"]:,.2f}</td></tr>'
            html += '</table>'

        self._interest_result.setHtml(html)
        self._interest_export_btn.setEnabled(True)

    def _export_interest_result(self) -> None:
        if not self._last_interest_result:
            return
        path, _ = QFileDialog.getSaveFileName(
            self, "导出利息计算结果", "利息计算结果.html",
            "HTML 文件 (*.html);;Word 文档 (*.docx);;PDF 文档 (*.pdf)"
        )
        if not path:
            return
        if path.lower().endswith(".docx"):
            try:
                from docx import Document
                from docx.shared import Pt, RGBColor, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                doc = Document()

                # 添加标题
                title = doc.add_paragraph()
                run = title.add_run("利息计算结果")
                run.bold = True
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                principal = self._last_interest_result.get("principal", 0)
                total = self._last_interest_result["total_interest"]
                days = self._last_interest_result["total_days"]
                segments = self._last_interest_result.get("segments", [])

                # 汇总表
                summary = doc.add_table(rows=1, cols=2)
                summary.style = "Table Grid"
                hdr = summary.rows[0].cells
                hdr[0].text = "项目"
                hdr[1].text = "数值"
                for cell in hdr:
                    for paragraph in cell.paragraphs:
                        for r in paragraph.runs:
                            r.bold = True

                for label, value in [
                    ("计算基数", f"¥ {principal:,.2f}"),
                    ("总利息", f"¥ {total:,.2f}"),
                    ("计息天数", f"{days} 天"),
                ]:
                    row = summary.add_row().cells
                    row[0].text = label
                    row[1].text = value
                    if label == "总利息":
                        for paragraph in row[1].paragraphs:
                            for r in paragraph.runs:
                                r.bold = True

                if segments:
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    p.add_run("分段明细").bold = True

                    table = doc.add_table(rows=1, cols=4)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = "起止日期"
                    hdr_cells[1].text = "天数"
                    hdr_cells[2].text = "利率"
                    hdr_cells[3].text = "利息"
                    for cell in hdr_cells:
                        for paragraph in cell.paragraphs:
                            for r in paragraph.runs:
                                r.bold = True
                    for seg in segments:
                        row_cells = table.add_row().cells
                        row_cells[0].text = f"{seg['start']} ~ {seg['end']}"
                        row_cells[1].text = str(seg["days"])
                        row_cells[2].text = f"{seg['rate']:.4f}%"
                        row_cells[3].text = f"¥ {seg['interest']:,.2f}"

                doc.save(path)
                QMessageBox.information(self, "导出成功", f"已保存到：{path}")
            except Exception as e:
                QMessageBox.warning(self, "导出失败", f"Word 导出出错：{e}")
        elif path.lower().endswith(".pdf"):
            try:
                from PySide6.QtPrintSupport import QPrinter
                from PySide6.QtGui import QTextDocument

                printer = QPrinter(QPrinter.PrinterMode.HighResolution)
                printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
                printer.setOutputFileName(path)
                printer.setPageMargins(QMarginsF(15, 15, 15, 15), QPrinter.Unit.Millimeter)

                doc = QTextDocument()
                doc.setHtml(self._interest_result.toHtml())
                doc.print(printer)
                QMessageBox.information(self, "导出成功", f"已保存到：{path}")
            except Exception as e:
                QMessageBox.warning(self, "导出失败", f"PDF 导出出错：{e}")
        else:
            try:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(self._interest_result.toHtml())
                QMessageBox.information(self, "导出成功", f"已保存到：{path}")
            except Exception as e:
                QMessageBox.warning(self, "导出失败", f"HTML 导出出错：{e}")

    def _calculate_liquidated(self) -> None:
        result = calculate_liquidated_damages(
            self._liquidated_base.value(),
            self._liquidated_days.value(),
            daily_rate_pct=self._liquidated_daily_rate.value(),
            annual_rate_pct=self._liquidated_annual_rate.value(),
            fixed_amount=self._liquidated_fixed.value(),
        )
        self._liquidated_result.setHtml(self._format_money_html("违约金结果", result))

    def _calculate_occupation(self) -> None:
        result = calculate_occupation_fee(
            self._occupation_monthly.value(),
            self._occupation_months.value(),
            daily_fee=self._occupation_daily.value(),
            days=self._occupation_days.value(),
            fixed_amount=self._occupation_fixed.value(),
        )
        self._occupation_result.setHtml(self._format_money_html("占有使用费结果", result))

    def _calculate_delay_interest(self) -> None:
        result = calculate_delay_performance_interest(
            self._delay_principal.value(),
            self._delay_days.value(),
            normal_annual_rate_pct=self._delay_normal_rate.value(),
            day_basis=int(self._delay_day_basis.currentData()),
        )
        self._delay_result.setHtml(
            self._format_money_html("迟延履行利息合计", result["total"])
            + f"<p>一般债务利息：¥ {result['general_interest']}</p>"
            + f"<p>加倍部分债务利息：¥ {result['doubled_interest']}</p>"
        )

    def _calculate_lawyer_fee(self) -> None:
        result = calculate_lawyer_fee(
            str(self._lawyer_mode.currentData()),
            fixed_fee=self._lawyer_fixed.value(),
            hourly_rate=self._lawyer_hour_rate.value(),
            hours=self._lawyer_hours.value(),
            claim_amount=self._lawyer_claim_amount.value(),
            rate_pct=self._lawyer_rate.value(),
        )
        self._lawyer_result.setHtml(
            self._format_money_html("律师费参考", result)
            + "<p>提示：律师费受地区、案件类型、委托阶段和风险代理监管规则影响较大，应以正式委托协议和当地规范为准。</p>"
        )

    def _calculate_labor_compensation(self) -> None:
        result = calculate_labor_compensation(
            self._labor_monthly_wage.value(),
            self._labor_local_avg.value(),
            self._labor_years.value(),
            self._labor_extra_months.value(),
        )
        self._labor_result.setHtml(
            self._format_money_html("经济补偿金", result["economic_compensation"])
            + f"<p>计发工资基数：¥ {result['monthly_base']}</p>"
            + f"<p>折算补偿月数：{result['compensation_months']}</p>"
            + f"<p>违法解除赔偿金参考：¥ {result['damages']}</p>"
        )

    def _calculate_work_injury_disability(self) -> None:
        result = calculate_work_injury_disability(
            self._injury_wage.value(),
            self._injury_level.value(),
            local_medical_subsidy=self._injury_local_medical.value(),
            local_employment_subsidy=self._injury_local_employment.value(),
        )
        self._injury_disability_result.setHtml(
            self._format_money_html("伤残待遇参考合计", result["total_reference"])
            + f"<p>一次性伤残补助金：¥ {result['one_time_subsidy']}</p>"
            + f"<p>伤残津贴（月）：¥ {result['monthly_allowance']}</p>"
            + f"<p>一次性医疗补助金：¥ {result['medical_subsidy']}</p>"
            + f"<p>一次性就业补助金：¥ {result['employment_subsidy']}</p>"
        )

    def _calculate_work_injury_death(self) -> None:
        result = calculate_work_injury_death(
            self._death_employee_wage.value(),
            self._death_local_avg.value(),
            self._death_national_income.value(),
            spouse_count=self._death_spouse_count.value(),
            other_dependents=self._death_other_count.value(),
            extra_supported_people=self._death_extra_supported.value(),
        )
        self._injury_death_result.setHtml(
            self._format_money_html("工亡待遇参考合计", result["total_reference"])
            + f"<p>丧葬补助金：¥ {result['funeral_grant']}</p>"
            + f"<p>一次性工亡补助金：¥ {result['death_grant']}</p>"
            + f"<p>供养亲属抚恤金（月参考）：¥ {result['dependent_pension_monthly']}</p>"
        )

    def _calculate_traffic_injury(self) -> None:
        years = compensation_years_for_age(self._traffic_age.value())
        result = calculate_traffic_injury_compensation(
            medical_fee=self._traffic_medical.value(),
            rehabilitation_fee=self._traffic_rehab.value(),
            followup_fee=self._traffic_followup.value(),
            hospital_meal_fee=self._traffic_meal.value(),
            nutrition_fee=self._traffic_nutrition.value(),
            nursing_days=self._traffic_nursing_days.value(),
            nursing_daily_fee=self._traffic_nursing_daily.value(),
            lost_income_days=self._traffic_lost_days.value(),
            lost_income_daily=self._traffic_lost_daily.value(),
            transportation_fee=self._traffic_transport.value(),
            accommodation_fee=self._traffic_accommodation.value(),
            disability_percent=self._traffic_disability_percent.value(),
            disability_base_year_amount=self._traffic_disability_base.value(),
            disability_years=years,
            assistive_device_fee=self._traffic_assistive.value(),
            mental_damage_fee=self._traffic_mental.value(),
            funeral_fee=self._traffic_funeral.value(),
            death_compensation_base_year_amount=self._traffic_death_base.value(),
            death_years=years,
            dependent_living_fee=self._traffic_dependents.value(),
        )
        self._traffic_result.setHtml(
            self._format_money_html("人身损害赔偿参考合计", result["total_reference"])
            + f"<p>护理费：¥ {result['nursing_fee']}</p>"
            + f"<p>误工费：¥ {result['lost_income_fee']}</p>"
            + f"<p>残疾赔偿金：¥ {result['disability_compensation']}</p>"
            + f"<p>死亡赔偿金：¥ {result['death_compensation']}</p>"
        )

    def _to_date(self, widget: QDateEdit) -> date:
        qdate = widget.date()
        return date(qdate.year(), qdate.month(), qdate.day())

    def _calculate_procedural_limit(self) -> None:
        result = calculate_procedural_deadline(
            self._to_date(self._limit_start),
            str(self._limit_rule.currentData()),
            exclude_start_day=self._limit_exclude_start.isChecked(),
            move_to_next_workday=self._limit_move_weekend.isChecked(),
        )
        self._limit_result.setHtml(
            f"<h3>{result['label']}</h3>"
            f"<p><b>推算日期：</b>{result['deadline'].strftime('%Y-%m-%d')}</p>"
            f"<p>{result['note']}</p>"
        )

    def _calculate_date_offset(self) -> None:
        result = calculate_date_offset(
            self._to_date(self._date_offset_start),
            days=self._date_offset_days.value(),
            months=self._date_offset_months.value(),
            years=self._date_offset_years.value(),
            exclude_start_day=self._date_offset_exclude.isChecked(),
            move_to_next_workday=self._date_offset_weekend.isChecked(),
        )
        self._date_offset_result.setHtml(
            f"<h3>推算结果</h3><p><b>{result.strftime('%Y-%m-%d')}</b></p>"
            "<p>说明：当前工作日顺延只排除周六、周日；法定节假日和调休请结合当年放假安排复核。</p>"
        )

    # ── 截图合并 Tab ──

    def _build_screenshot_merge_tab(self) -> QWidget:
        return self._wrap_scroll_tab(self._populate_screenshot_merge_tab)

    def _populate_screenshot_merge_tab(self, layout: QVBoxLayout) -> None:
        body = QGridLayout()
        body.setHorizontalSpacing(12)
        body.setVerticalSpacing(12)
        body.addWidget(self._build_screenshot_image_list_card(), 0, 0)

        # 右侧列：布局设置 + 标签设置 + 操作栏 上下排列
        right_col = QVBoxLayout()
        right_col.setSpacing(12)
        right_col.addWidget(self._build_screenshot_settings_card())
        right_col.addWidget(self._build_screenshot_label_card())
        right_col.addWidget(self._build_screenshot_action_card())
        right_col.addStretch()
        body.addLayout(right_col, 0, 1)

        body.setColumnStretch(0, 7)
        body.setColumnStretch(1, 5)
        body.setRowStretch(0, 1)
        layout.addLayout(body)

    def _build_docx_compare_tab(self) -> QWidget:
        """文档对比标签页。"""
        case_path: Optional[Path] = None
        if self._initial_case_id:
            case = self._case_manager.get_case(self._initial_case_id)
            if case:
                case_path = Path(case.get("path", ""))
                if not case_path.exists():
                    case_path = None
        return DocxCompareWidget(self, case_path=case_path)

    def _build_auto_format_tab(self) -> QWidget:
        """自动排版标签页。"""
        return DocxAutoFormatWidget(self)

    def _build_screenshot_image_list_card(self) -> QWidget:
        card, layout = self._create_card(
            "图片列表",
            "将微信截图或其他图片拖放到下方列表中，或点击「选择文件夹」导入。支持在列表中直接拖拽调整顺序。",
        )

        self._screenshot_image_list = ScreenshotImageList()
        self._screenshot_image_list.setMinimumHeight(520)
        self._screenshot_image_list.images_changed.connect(self._on_screenshot_images_changed)
        layout.addWidget(self._screenshot_image_list)

        info_row = QHBoxLayout()
        self._screenshot_count_label = QLabel("共 0 张图片")
        self._screenshot_count_label.setStyleSheet(f"""
            color: {COLORS['text_muted']};
            font-size: 11px;
        """)
        info_row.addWidget(self._screenshot_count_label)
        info_row.addStretch()

        view_btn = self._make_button("切换视图")
        view_btn.clicked.connect(self._screenshot_image_list.toggle_view_mode)
        info_row.addWidget(view_btn)

        select_btn = self._make_button("选择文件夹")
        select_btn.clicked.connect(self._on_screenshot_select_folder)
        info_row.addWidget(select_btn)

        clear_btn = self._make_button("清空")
        clear_btn.clicked.connect(self._screenshot_image_list.clear_all)
        info_row.addWidget(clear_btn)

        layout.addLayout(info_row)
        return card

    def _build_screenshot_settings_card(self) -> QWidget:
        card, layout = self._create_card(
            "布局设置",
            "调整每页摆放张数、排序规则、页面方向及间距。",
        )

        form = self._create_form()

        self._screenshot_per_page_combo = QComboBox()
        self._screenshot_per_page_combo.addItems(["1 张", "2 张", "3 张"])
        self._screenshot_per_page_combo.setCurrentIndex(1)
        form.addRow("每页张数", self._screenshot_per_page_combo)

        self._screenshot_order_combo = QComboBox()
        self._screenshot_order_combo.addItems(["手动排序", "正序（文件名数字）", "倒序（文件名数字）"])
        self._screenshot_order_combo.setCurrentIndex(1)
        self._screenshot_order_combo.currentIndexChanged.connect(self._on_screenshot_order_changed)
        form.addRow("排序规则", self._screenshot_order_combo)

        self._screenshot_orientation_combo = QComboBox()
        self._screenshot_orientation_combo.addItems(["横向 A4", "纵向 A4"])
        form.addRow("页面方向", self._screenshot_orientation_combo)

        self._screenshot_margin_spin = QSpinBox()
        self._screenshot_margin_spin.setRange(5, 50)
        self._screenshot_margin_spin.setValue(15)
        self._screenshot_margin_spin.setSuffix(" mm")
        form.addRow("页边距", self._screenshot_margin_spin)

        self._screenshot_gap_spin = QSpinBox()
        self._screenshot_gap_spin.setRange(0, 30)
        self._screenshot_gap_spin.setValue(8)
        self._screenshot_gap_spin.setSuffix(" mm")
        form.addRow("图片间距", self._screenshot_gap_spin)

        layout.addLayout(form)
        return card

    def _build_screenshot_label_card(self) -> QWidget:
        card, layout = self._create_card(
            "标签设置",
            "为每张图片配置标签文字的位置和内容。",
        )

        form = self._create_form()

        self._screenshot_label_position_combo = QComboBox()
        self._screenshot_label_position_combo.addItems(["上方", "下方", "不显示"])
        form.addRow("标签位置", self._screenshot_label_position_combo)

        self._screenshot_label_mode_combo = QComboBox()
        self._screenshot_label_mode_combo.addItems(["图1 / 图2 …", "自定义前缀 + 编号", "使用文件名", "不显示"])
        self._screenshot_label_mode_combo.currentIndexChanged.connect(self._on_screenshot_label_mode_changed)
        form.addRow("标签模式", self._screenshot_label_mode_combo)

        self._screenshot_label_prefix_edit = QLineEdit()
        self._screenshot_label_prefix_edit.setPlaceholderText("图")
        self._screenshot_label_prefix_edit.setText("图")
        form.addRow("自定义前缀", self._screenshot_label_prefix_edit)

        self._screenshot_label_size_combo = QComboBox()
        self._screenshot_label_size_combo.addItems(["小 (10pt)", "偏小 (12pt)", "中 (14pt)", "偏大 (16pt)", "大 (18pt)"])
        self._screenshot_label_size_combo.setCurrentIndex(1)
        form.addRow("文字大小", self._screenshot_label_size_combo)

        layout.addLayout(form)
        return card

    def _build_screenshot_action_card(self) -> QWidget:
        card, layout = self._create_card("", "")
        # 移除标题占用的空间，让这个卡片更紧凑
        # _create_card 已经添加了标题和 subtitle，这里我们需要一个简洁的操作栏
        # 清空之前添加的 widget，重新布局
        while layout.count():
            item = layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        action_row = QHBoxLayout()
        action_row.setSpacing(8)

        self._screenshot_generate_btn = self._make_button("生成 PDF", accent=True)
        self._screenshot_generate_btn.clicked.connect(self._on_screenshot_generate)
        action_row.addWidget(self._screenshot_generate_btn)

        self._screenshot_save_to_source = QCheckBox("保存到源文件夹")
        self._screenshot_save_to_source.setStyleSheet(f"""
            QCheckBox {{
                color: {COLORS['text_secondary']};
                font-size: 11px;
            }}
            QCheckBox::indicator {{
                width: 16px;
                height: 16px;
                border-radius: 4px;
                border: 1px solid {COLORS['border']};
            }}
            QCheckBox::indicator:checked {{
                background: {COLORS['accent']};
                border-color: {COLORS['accent']};
                image: url({CHECK_ICON_PATH});
            }}
        """)
        action_row.addWidget(self._screenshot_save_to_source)

        action_row.addStretch()

        self._screenshot_status_label = QLabel("就绪")
        self._screenshot_status_label.setStyleSheet(f"""
            color: {COLORS['text_muted']};
            font-size: 11px;
        """)
        action_row.addWidget(self._screenshot_status_label)

        layout.addLayout(action_row)
        return card

    def _on_screenshot_images_changed(self, paths: List[Path]) -> None:
        count = len(paths)
        self._screenshot_count_label.setText(f"共 {count} 张图片")
        self._screenshot_status_label.setText("就绪" if count > 0 else "请导入图片")
        # 记录源文件夹（第一张图片的父目录）
        if paths:
            self._screenshot_source_folder = Path(paths[0]).parent

    def _on_screenshot_select_folder(self) -> None:
        folder = QFileDialog.getExistingDirectory(self, "选择包含截图的文件夹")
        if folder:
            self._screenshot_source_folder = Path(folder)
            self._screenshot_image_list.add_images([Path(folder)])

    def _on_screenshot_order_changed(self, index: int) -> None:
        if self._screenshot_image_list is None:
            return
        paths = self._screenshot_image_list.get_ordered_paths()
        if not paths:
            return

        order_map = {0: "manual", 1: "asc", 2: "desc"}
        order = order_map.get(index, "manual")

        if order == "manual":
            return

        sorted_paths = self._screenshot_merger.sort_images(paths, order=order)
        self._screenshot_image_list.set_images(sorted_paths)

    def _on_screenshot_label_mode_changed(self, index: int) -> None:
        enabled = index == 1  # "自定义前缀 + 编号"
        self._screenshot_label_prefix_edit.setEnabled(enabled)

    def _get_screenshot_merge_settings(self) -> Dict[str, Any]:
        """收集当前所有设置参数"""
        per_page = self._screenshot_per_page_combo.currentIndex() + 1
        orientation = "L" if self._screenshot_orientation_combo.currentIndex() == 0 else "P"
        label_position_map = {0: "top", 1: "bottom", 2: "none"}
        label_mode_map = {0: "auto", 1: "custom", 2: "filename", 3: "none"}

        label_size_map = {0: 10, 1: 12, 2: 14, 3: 16, 4: 18}

        return {
            "per_page": per_page,
            "orientation": orientation,
            "margin_mm": float(self._screenshot_margin_spin.value()),
            "gap_mm": float(self._screenshot_gap_spin.value()),
            "label_position": label_position_map.get(self._screenshot_label_position_combo.currentIndex(), "top"),
            "label_mode": label_mode_map.get(self._screenshot_label_mode_combo.currentIndex(), "auto"),
            "label_prefix": self._screenshot_label_prefix_edit.text().strip() or "图",
            "label_font_size": label_size_map.get(self._screenshot_label_size_combo.currentIndex(), 12),
        }

    def _on_screenshot_generate(self) -> None:
        if self._screenshot_image_list is None:
            return

        paths = self._screenshot_image_list.get_ordered_paths()
        if not paths:
            QMessageBox.information(self, "请导入图片", "请先导入至少一张图片再生成 PDF。")
            return

        save_to_source = self._screenshot_save_to_source.isChecked() if self._screenshot_save_to_source else False

        if save_to_source and self._screenshot_source_folder and self._screenshot_source_folder.exists():
            # 自动保存到源文件夹
            output_path = self._screenshot_source_folder / "截图合并.pdf"
            # 如果文件已存在，添加序号
            counter = 1
            original_path = output_path
            while output_path.exists():
                output_path = original_path.with_stem(f"截图合并 ({counter})")
                counter += 1
        else:
            # 弹出保存对话框，默认文件名
            default_name = "截图合并.pdf"
            if self._screenshot_source_folder:
                default_name = str(self._screenshot_source_folder / default_name)
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "保存 PDF",
                default_name,
                "PDF 文件 (*.pdf)",
            )
            if not file_path:
                return
            output_path = Path(file_path)

        settings = self._get_screenshot_merge_settings()
        self._screenshot_status_label.setText("生成中…")
        self._screenshot_generate_btn.setEnabled(False)
        self._screenshot_save_to_source.setEnabled(False)
        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)

        self._screenshot_merge_worker = _ScreenshotMergeWorker(
            self._screenshot_merger,
            paths,
            output_path,
            settings,
        )
        self._screenshot_merge_worker.finished.connect(
            lambda p: self._on_screenshot_generate_finished(p, save_to_source)
        )
        self._screenshot_merge_worker.error.connect(self._on_screenshot_generate_error)
        self._screenshot_merge_worker.start()

    def _on_screenshot_generate_finished(self, path: str, save_to_source: bool) -> None:
        worker = self._screenshot_merge_worker
        self._screenshot_merge_worker = None
        QApplication.restoreOverrideCursor()
        self._screenshot_generate_btn.setEnabled(True)
        self._screenshot_save_to_source.setEnabled(True)

        if save_to_source:
            self._screenshot_status_label.setText(f"已保存到源文件夹: {Path(path).name}")
        else:
            self._screenshot_status_label.setText(f"已保存: {Path(path).name}")

        # 打开 PDF
        QDesktopServices.openUrl(QUrl.fromLocalFile(path))
        if worker is not None:
            worker.deleteLater()

    def _on_screenshot_generate_error(self, error_msg: str) -> None:
        worker = self._screenshot_merge_worker
        self._screenshot_merge_worker = None
        QApplication.restoreOverrideCursor()
        self._screenshot_generate_btn.setEnabled(True)
        self._screenshot_save_to_source.setEnabled(True)
        self._screenshot_status_label.setText("生成失败")
        if worker is not None:
            worker.deleteLater()
        QMessageBox.warning(self, "生成失败", error_msg)

    def closeEvent(self, event):  # type: ignore[override]
        if self._has_active_background_workers():
            QMessageBox.information(self, "请稍候", "当前仍有后台任务正在执行，请等待完成后再关闭工具中心。")
            event.ignore()
            return
        super().closeEvent(event)


class _ScreenshotMergeWorker(QThread):
    """在后台线程执行截图合并 PDF 生成。"""

    finished = Signal(str)
    error = Signal(str)

    def __init__(
        self,
        merger: ScreenshotPdfMerger,
        images: List[Path],
        output_path: Path,
        settings: Dict[str, Any],
    ):
        super().__init__()
        self.merger = merger
        self.images = images
        self.output_path = output_path
        self.settings = settings

    def run(self):
        try:
            result = self.merger.generate_pdf(
                self.images,
                self.output_path,
                **self.settings,
            )
            self.finished.emit(str(result))
        except Exception as exc:
            self.error.emit(str(exc))
