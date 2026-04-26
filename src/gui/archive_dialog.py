# -*- coding: utf-8 -*-
"""电子化归档主对话框

提供从现有案卷文件夹提取信息、生成标准化归档文档的功能。
"""

from pathlib import Path
from typing import Dict, List, Optional, Any
from datetime import datetime
import os

from PySide6.QtWidgets import (
    QDialog,
    QVBoxLayout,
    QHBoxLayout,
    QSplitter,
    QWidget,
    QLabel,
    QPushButton,
    QFileDialog,
    QMessageBox,
    QSizePolicy,
)
from PySide6.QtCore import Qt, Signal, QObject, QThread

from src.gui.widgets.archive_file_tree import ArchiveFileTree
from src.gui.widgets.archive_variable_panel import ArchiveVariablePanel
from src.gui.widgets.archive_preview import ArchivePreview
from src.core.archive_engine import ArchiveEngine
from src.utils.logger import get_logger
from src.gui.styles import APP_COLORS as COLORS, button_style
from src.gui.window_metrics import APP_SURFACE_DEFAULT_SIZE, APP_SURFACE_MIN_SIZE


class _DirectoryScanWorker(QObject):
    """后台统计目录文件数量。"""

    finished = Signal(int)
    failed = Signal(str)

    def __init__(self, folder_path: Path):
        super().__init__()
        self._folder_path = folder_path

    def run(self) -> None:
        try:
            count = 0
            for _, dirnames, filenames in os.walk(self._folder_path):
                dirnames[:] = [name for name in dirnames if not name.startswith(".") and name != "__pycache__"]
                count += len([name for name in filenames if not name.startswith(".")])
            self.finished.emit(count)
        except Exception as exc:
            self.failed.emit(str(exc))


class ArchiveDialog(QDialog):
    """电子化归档主对话框 - Modern UI v3"""

    change_folder_requested = Signal()

    def __init__(self, folder_path: Path, parent=None, embed_mode: bool = False):
        super().__init__(parent)
        self._embed_mode = embed_mode
        if embed_mode:
            self.setWindowFlags(Qt.Widget)
            self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self._logger = get_logger()
        self._folder_path = Path(folder_path)
        self._archive_engine = ArchiveEngine()

        self._current_template: Optional[Path] = None
        self._status_message_label: Optional[QLabel] = None
        self._file_count_label: Optional[QLabel] = None
        self._tree_expanded = False
        self._step_labels: List[QLabel] = []
        self._scan_thread: Optional[QThread] = None
        self._scan_worker: Optional[_DirectoryScanWorker] = None

        if not embed_mode:
            self.setWindowTitle(f"电子化归档 - {self._folder_path.name}")
            self.setMinimumSize(*APP_SURFACE_MIN_SIZE)
            self.resize(*APP_SURFACE_DEFAULT_SIZE)

        self._setup_ui()
        self._load_folder()

    def _setup_ui(self) -> None:
        """设置界面"""
        c = COLORS
        self.setStyleSheet(f"QDialog {{ background: {c['surface_1']}; }}")
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # 工具栏
        toolbar = self._create_toolbar()
        layout.addWidget(toolbar)

        # 主内容区（三栏）
        self._splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter = self._splitter
        splitter.setChildrenCollapsible(True)
        splitter.setHandleWidth(6)
        splitter.setStyleSheet(f"""
            QSplitter::handle {{
                background: {c['border']};
                width: 1px;
                margin: 0 2px;
            }}
            QSplitter::handle:hover {{
                background: {c['border_strong']};
            }}
        """)

        # 左侧：变量定义面板
        self._variable_panel = ArchiveVariablePanel()
        self._variable_panel.setMinimumWidth(310)
        self._variable_panel.setMaximumWidth(390)
        self._variable_panel.setSizePolicy(
            QSizePolicy.Policy.Preferred,
            QSizePolicy.Policy.Expanding,
        )
        splitter.addWidget(self._variable_panel)

        # 中间：文件预览区
        self._preview = ArchivePreview()
        self._preview.setMinimumWidth(320)
        # 同步变量列表到预览控件
        self._preview.set_variables(self._variable_panel.get_variable_list())
        splitter.addWidget(self._preview)

        # 右侧：文件夹结构树
        self._file_tree_panel = self._create_file_tree_panel()
        self._file_tree_panel.setMinimumWidth(240)
        self._file_tree_panel.setMaximumWidth(460)
        splitter.addWidget(self._file_tree_panel)

        # 设置分割比例
        splitter.setStretchFactor(0, 0)
        splitter.setStretchFactor(1, 1)
        splitter.setStretchFactor(2, 0)
        splitter.setCollapsible(0, True)
        splitter.setCollapsible(1, False)
        splitter.setCollapsible(2, True)
        splitter.setSizes([330, 730, 320])

        layout.addWidget(splitter, 1)

        # 底部状态栏
        footer = self._create_footer()
        layout.addWidget(footer)

        # 连接信号
        self._connect_signals()

    def _create_toolbar(self) -> QWidget:
        """创建工具栏"""
        c = COLORS
        toolbar = QWidget()
        toolbar.setProperty("archiveToolbar", True)
        toolbar.setStyleSheet(f"""
            QWidget[archiveToolbar="true"] {{
                background: {c['surface_0']};
                border-bottom: 1px solid {c['border']};
            }}
        """)

        layout = QHBoxLayout(toolbar)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.setSpacing(10)

        title_wrap = QWidget()
        title_wrap.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Fixed)
        title_wrap.setMaximumWidth(260)
        title_wrap.setStyleSheet("background: transparent; border: none;")
        title_layout = QVBoxLayout(title_wrap)
        title_layout.setContentsMargins(0, 0, 0, 0)
        title_layout.setSpacing(2)

        title_label = QLabel("电子化归档")
        title_label.setStyleSheet(
            f"background: transparent; border: none; font-size: 16px; "
            f"font-weight: 700; color: {c['text_primary']};"
        )
        title_layout.addWidget(title_label)

        subtitle_label = QLabel(f"案卷：{self._folder_path.name}")
        subtitle_label.setStyleSheet(
            f"background: transparent; border: none; font-size: 12px; color: {c['text_tertiary']};"
        )
        subtitle_label.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        title_layout.addWidget(subtitle_label)
        layout.addWidget(title_wrap)

        self._step_bar = self._create_step_bar()
        self._step_bar.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        layout.addWidget(self._step_bar, 1)

        if self._embed_mode:
            change_btn = QPushButton("更换文件夹")
            change_btn.setFixedHeight(32)
            change_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            change_btn.setStyleSheet(button_style(compact=True))
            change_btn.clicked.connect(self.change_folder_requested.emit)
            layout.addWidget(change_btn)

        # 选择归档模板按钮
        template_btn = QPushButton("选择模板")
        template_btn.setFixedHeight(32)
        template_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        template_btn.setStyleSheet(button_style(compact=True))
        template_btn.clicked.connect(self._on_select_template)
        layout.addWidget(template_btn)

        # 替换导出按钮
        export_doc_btn = QPushButton("替换导出")
        export_doc_btn.setFixedHeight(32)
        export_doc_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        export_doc_btn.setStyleSheet(button_style(primary=True, compact=True))
        export_doc_btn.clicked.connect(self._on_export)
        layout.addWidget(export_doc_btn)

        return toolbar

    def _create_step_bar(self) -> QWidget:
        """创建归档步骤状态栏。"""
        bar = QWidget()
        bar.setStyleSheet("""
            QWidget {
                background: transparent;
                border: none;
            }
        """)
        layout = QHBoxLayout(bar)
        layout.setContentsMargins(4, 0, 4, 0)
        layout.setSpacing(6)

        steps = ["选目录", "选文件", "抽取变量", "选择模板", "导出"]
        self._step_labels = []
        for index, text in enumerate(steps):
            label = QLabel(f"{index + 1}. {text}")
            label.setFixedHeight(26)
            label.setMinimumWidth(76)
            label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            layout.addWidget(label, 1)
            self._step_labels.append(label)
        self._set_archive_step(0)
        return bar

    def _set_archive_step(self, active_index: int) -> None:
        """更新步骤状态。"""
        c = COLORS
        for index, label in enumerate(self._step_labels):
            if index < active_index:
                bg = c['accent_subtle']
                fg = c['accent']
                border = c['accent']
                weight = "700"
            elif index == active_index:
                bg = c['surface_2']
                fg = c['text_primary']
                border = c['border_strong']
                weight = "700"
            else:
                bg = c['surface_0']
                fg = c['text_tertiary']
                border = c['border']
                weight = "500"
            label.setStyleSheet(
                f"background: {bg}; color: {fg}; border: 1px solid {border}; "
                f"border-radius: 6px; font-size: 12px; font-weight: {weight};"
            )

    def _create_file_tree_panel(self) -> QWidget:
        """创建案卷结构面板。"""
        c = COLORS
        panel = QWidget()
        panel.setStyleSheet(f"background: {c['surface_0']};")

        layout = QVBoxLayout(panel)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        header = QWidget()
        header.setStyleSheet(f"""
            QWidget {{
                background: {c['surface_0']};
                border-bottom: 1px solid {c['border']};
            }}
        """)
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(12, 10, 12, 10)
        header_layout.setSpacing(8)

        title_wrap = QWidget()
        title_layout = QVBoxLayout(title_wrap)
        title_layout.setContentsMargins(0, 0, 0, 0)
        title_layout.setSpacing(2)

        title = QLabel("案卷结构")
        title.setStyleSheet(f"font-size: 13px; font-weight: 700; color: {c['text_primary']};")
        title_layout.addWidget(title)

        self._file_count_label = QLabel("正在加载...")
        self._file_count_label.setStyleSheet(f"font-size: 11px; color: {c['text_muted']};")
        title_layout.addWidget(self._file_count_label)
        header_layout.addWidget(title_wrap, 1)

        refresh_btn = QPushButton("刷新")
        refresh_btn.setFixedHeight(28)
        refresh_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        refresh_btn.setStyleSheet(button_style(compact=True))
        refresh_btn.clicked.connect(self._refresh_tree)
        header_layout.addWidget(refresh_btn)

        self._expand_btn = QPushButton("展开")
        self._expand_btn.setFixedHeight(28)
        self._expand_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self._expand_btn.setStyleSheet(button_style(compact=True))
        self._expand_btn.setToolTip("展开全部文件夹")
        self._expand_btn.clicked.connect(self._toggle_expand_collapse)
        header_layout.addWidget(self._expand_btn)

        layout.addWidget(header)

        self._file_tree = ArchiveFileTree()
        self._file_tree.set_lazy_mode(True)
        layout.addWidget(self._file_tree, 1)
        return panel

    def _create_footer(self) -> QWidget:
        """创建底部按钮栏"""
        c = COLORS
        footer = QWidget()
        footer.setProperty("archiveFooter", True)
        footer.setStyleSheet(f"""
            QWidget[archiveFooter="true"] {{
                background: {c['surface_1']};
                border-top: 1px solid {c['border']};
            }}
        """)

        layout = QHBoxLayout(footer)
        layout.setContentsMargins(14, 7, 14, 7)
        layout.setSpacing(12)

        # 文件夹路径
        folder_label = QLabel(f"路径：{self._folder_path}")
        folder_label.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Fixed)
        folder_label.setToolTip(str(self._folder_path))
        folder_label.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        folder_label.setStyleSheet(f"""
            font-size: 12px;
            color: {c['text_tertiary']};
        """)
        layout.addWidget(folder_label, 1)

        self._status_message_label = QLabel("就绪")
        self._status_message_label.setStyleSheet(f"font-size: 12px; color: {c['text_secondary']};")
        layout.addWidget(self._status_message_label)

        # 统计信息
        self._stats_label = QLabel("")
        self._stats_label.setStyleSheet(f"font-size: 12px; color: {c['text_muted']};")
        layout.addWidget(self._stats_label)

        return footer

    def _connect_signals(self) -> None:
        """连接信号"""
        # 文件树信号
        self._file_tree.file_clicked.connect(self._on_file_clicked)
        self._file_tree.file_double_clicked.connect(self._on_file_double_clicked)
        self._file_tree.folder_clicked.connect(self._on_folder_clicked)
        self._file_tree.structure_changed.connect(self._on_structure_changed)

        # 阅览信号
        self._preview.variable_set.connect(self._on_variable_set)
        self._preview.field_defined.connect(self._on_field_defined)
        self._preview.variables_detected.connect(self._on_variables_detected)
        self._preview.save_requested.connect(self._on_save_document)
        self._preview.save_as_requested.connect(self._on_save_as_document)

        # 变量面板信号
        self._variable_panel.variables_changed.connect(self._on_variables_changed)

    def _load_folder(self) -> None:
        """加载文件夹结构"""
        self._show_status("正在加载案卷结构...")
        self._file_tree.load_folder(self._folder_path)
        self._tree_expanded = False
        self._expand_btn.setText("展开")
        self._expand_btn.setToolTip("展开全部文件夹")
        self._update_stats()
        self._set_archive_step(1)
        self._start_background_scan()

    def _update_stats(self) -> None:
        """更新统计信息"""
        files = self._file_tree.get_all_files()
        stats_text = f"已加载 {len(files)} 个文件"
        self._stats_label.setText(stats_text)
        if self._file_count_label is not None:
            self._file_count_label.setText(stats_text)

    def _start_background_scan(self) -> None:
        """后台统计完整目录文件数量。"""
        if self._scan_thread is not None and self._scan_thread.isRunning():
            return
        if self._file_count_label is not None:
            self._file_count_label.setText("正在后台扫描...")

        worker = _DirectoryScanWorker(self._folder_path)
        thread = QThread(self)
        worker.moveToThread(thread)
        thread.started.connect(worker.run)
        worker.finished.connect(self._on_background_scan_finished)
        worker.failed.connect(self._on_background_scan_failed)
        worker.finished.connect(thread.quit)
        worker.failed.connect(thread.quit)
        worker.finished.connect(worker.deleteLater)
        worker.failed.connect(worker.deleteLater)
        thread.finished.connect(thread.deleteLater)
        thread.finished.connect(self._clear_background_scan)
        self._scan_worker = worker
        self._scan_thread = thread
        thread.start()

    def _on_background_scan_finished(self, file_count: int) -> None:
        stats_text = f"已加载 {file_count} 个文件"
        self._stats_label.setText(stats_text)
        if self._file_count_label is not None:
            self._file_count_label.setText(stats_text)
        self._show_status("案卷结构扫描完成")

    def _on_background_scan_failed(self, error_message: str) -> None:
        if self._file_count_label is not None:
            self._file_count_label.setText("扫描失败")
        self._logger.warning(f"后台扫描案卷结构失败: {error_message}")

    def _clear_background_scan(self) -> None:
        self._scan_worker = None
        self._scan_thread = None

    def _show_status(self, message: str) -> None:
        """更新归档页底部状态。"""
        if self._status_message_label is not None:
            self._status_message_label.setText(message)

    def _preview_selected_file(self, file_path: Path) -> None:
        """预览选中的文件并维护状态。"""
        if not file_path.exists():
            QMessageBox.warning(self, "文件不存在", f"找不到文件：{file_path}")
            return
        if hasattr(self, '_variable_replacements'):
            self._variable_replacements = []
        self._preview.preview_file(file_path)
        self._show_status(f"已打开：{file_path.name}")
        self._set_archive_step(2)

    def _on_file_clicked(self, file_path: Path) -> None:
        """文件单击事件"""
        self._preview_selected_file(file_path)

    def _on_file_double_clicked(self, file_path: Path) -> None:
        """文件双击事件"""
        self._preview_selected_file(file_path)

    def _on_folder_clicked(self, folder_path: Path) -> None:
        """文件夹单击事件"""
        self._show_status(f"已选中文件夹：{folder_path.name}")

    def _on_structure_changed(self) -> None:
        """文件夹结构改变"""
        self._update_stats()

    def _refresh_tree(self) -> None:
        """刷新文件树。"""
        self._file_tree.refresh()
        self._tree_expanded = False
        self._expand_btn.setText("展开")
        self._expand_btn.setToolTip("展开全部文件夹")
        self._update_stats()
        self._show_status("案卷结构已刷新")

    def _on_import_variables(self) -> None:
        """导入变量"""
        self._variable_panel._on_import()

    def _on_export_variables(self) -> None:
        """导出变量"""
        self._variable_panel._on_export()

    def _on_select_template(self) -> None:
        """选择归档模板"""
        from src.config.path_manager import get_path_manager
        default_dir = get_path_manager().templates_dir

        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "选择归档模板",
            str(default_dir),
            "Word 文档 (*.docx);;所有文件 (*.*)"
        )

        if not file_path:
            return

        self._current_template = Path(file_path)
        self._preview.preview_file(self._current_template)
        self._show_status(f"已选择模板：{self._current_template.name}")
        self._set_archive_step(4)

    def _toggle_expand_collapse(self) -> None:
        """切换文件夹展开/收缩"""
        if self._tree_expanded:
            self._file_tree.collapse_all()
            self._expand_btn.setText("展开")
            self._expand_btn.setToolTip("展开全部文件夹")
            self._tree_expanded = False
        else:
            self._file_tree.expand_all()
            self._expand_btn.setText("收起")
            self._expand_btn.setToolTip("收起全部文件夹")
            self._tree_expanded = True
        self._update_stats()

    def _on_export(self) -> None:
        """替换导出"""
        if not self._current_template:
            QMessageBox.warning(self, "警告", "请先选择归档模板")
            return

        # 获取变量值
        values = self._variable_panel.get_all_values()

        if not values:
            reply = QMessageBox.question(
                self,
                "确认",
                "变量值为空，是否继续导出？\n\n导出的文档将保留原始变量占位符。",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.No:
                return

        # 选择保存位置
        output_path, _ = QFileDialog.getSaveFileName(
            self,
            "保存归档文档",
            str(self._folder_path.parent),
            "Word 文档 (*.docx)"
        )

        if not output_path:
            return

        try:
            # 执行导出
            output_path = self._archive_engine.export(
                self._current_template,
                values,
                Path(output_path)
            )

            # 保存历史记录
            self._archive_engine.save_history({
                "folder_path": str(self._folder_path),
                "variables": values,
                "template_path": str(self._current_template),
                "output_path": str(output_path),
                "created_at": datetime.now().isoformat()
            })

            QMessageBox.information(
                self,
                "导出成功",
                f"文档已保存到:\n{output_path}"
            )
            self._show_status("导出成功")
            self._set_archive_step(4)

        except Exception as e:
            self._logger.error(f"导出失败: {e}")
            QMessageBox.critical(self, "导出失败", str(e))

    def _on_variable_set(self, var_key: str, text: str, selection_info: Dict[str, Any]) -> None:
        """设置为变量 - 将选中文本替换为变量占位符"""
        # 记录替换操作（用于保存时保持格式）
        if not hasattr(self, '_variable_replacements'):
            self._variable_replacements = []
        self._variable_replacements.append({
            'original': text,
            'variable': f"{{{{{var_key}}}}}",
            'key': var_key,
            'selection_start': selection_info.get('selection_start'),
            'selection_end': selection_info.get('selection_end'),
        })

        self._show_status(f"已设置变量：{var_key}")
        self._set_archive_step(3)

        self._logger.info(f"设置为变量: '{text}' -> {{{{{{var_key}}}}}}")

    def _on_field_defined(self, var_key: str, text: str) -> None:
        """定义为变量字段"""
        # 更新变量面板中的值
        self._variable_panel.set_value(var_key, text)
        self._show_status(f"已定义字段：{var_key}")
        self._set_archive_step(3)

    def _on_variables_detected(self, new_vars: list) -> None:
        """检测到新变量"""
        for var_key in new_vars:
            if not self._variable_panel.has_variable(var_key):
                # 添加新变量
                self._variable_panel.add_variable(var_key, var_key, "")

        if new_vars:
            self._show_status(f"检测到 {len(new_vars)} 个新变量")
            self._set_archive_step(3)

    def _on_variables_changed(self) -> None:
        """变量列表改变"""
        # 更新预览区的变量列表
        self._preview.set_variables(self._variable_panel.get_variable_list())

    def _on_save_document(self) -> None:
        """保存文档 - 将预览区的内容保存回原文件"""
        file_path = self._preview.get_current_file_path()
        if not file_path:
            return

        # 确认是否保存
        reply = QMessageBox.question(
            self,
            "确认保存",
            f"是否将当前内容保存到:\n{file_path}\n\n注意：此操作将覆盖原文件，建议先备份。",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply != QMessageBox.StandardButton.Yes:
            return

        try:
            # 保存文档（基于原始文档应用变量替换）
            self._save_word_document(file_path)

            self._show_status(f"文档已保存：{file_path.name}")
            QMessageBox.information(self, "保存成功", f"文档已保存:\n{file_path}")

        except Exception as e:
            self._logger.error(f"保存文档失败: {e}")
            QMessageBox.critical(self, "保存失败", str(e))

    def _on_save_as_document(self) -> None:
        """另存为文档 - 将预览区的内容保存到新文件"""
        file_path = self._preview.get_current_file_path()
        if not file_path:
            return

        # 选择保存位置
        default_name = f"{file_path.stem}_已设置变量{file_path.suffix}"
        output_path, _ = QFileDialog.getSaveFileName(
            self,
            "另存为",
            str(file_path.parent / default_name),
            "Word 文档 (*.docx)"
        )

        if not output_path:
            return

        try:
            # 保存文档（基于原始文档应用变量替换）
            self._save_word_document(Path(output_path))

            self._show_status(f"文档已另存为：{Path(output_path).name}")
            QMessageBox.information(self, "保存成功", f"文档已保存:\n{output_path}")

        except Exception as e:
            self._logger.error(f"另存为失败: {e}")
            QMessageBox.critical(self, "保存失败", str(e))

    def _save_word_document(self, output_path: Path, content: str = None) -> None:
        """保存Word文档 - 保留原始格式
        
        基于原始文档，应用所有记录的变量替换，保留所有格式。
        
        Args:
            output_path: 输出文件路径
            content: 已弃用，保留参数以保持兼容性
        """
        from docx import Document
        import os
        import shutil
        import tempfile
        
        # 获取原始文件路径
        original_path = self._preview.get_current_file_path()
        if not original_path or not original_path.exists():
            raise ValueError("原始文件不存在")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        same_target = output_path.resolve() == original_path.resolve()
        save_path = output_path
        temp_path: Optional[Path] = None

        try:
            if not same_target:
                temp_fd, temp_name = tempfile.mkstemp(
                    suffix=output_path.suffix or ".docx",
                    prefix=f".{output_path.stem}.",
                    dir=output_path.parent,
                )
                os.close(temp_fd)
                temp_path = Path(temp_name)
                # 复制原始文档到临时文件，全部处理成功后再替换最终目标。
                shutil.copy2(str(original_path), str(temp_path))
                save_path = temp_path
                doc = Document(str(save_path))
            else:
                doc = Document(str(original_path))

            # 如果没有变量替换记录，直接保存/返回
            if not hasattr(self, '_variable_replacements') or not self._variable_replacements:
                if same_target:
                    doc.save(str(output_path))
                else:
                    os.replace(str(save_path), str(output_path))
                self._logger.info(f"Word文档已保存（无变量替换）: {output_path}")
                return

            current_preview_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

            for replacement in self._variable_replacements:
                if not self._apply_recorded_replacement(doc, replacement, current_preview_text):
                    raise ValueError(
                        f"无法应用变量替换: {replacement.get('original', '')} -> {replacement.get('variable', '')}"
                    )

                start = replacement.get('selection_start')
                end = replacement.get('selection_end')
                if isinstance(start, int) and isinstance(end, int):
                    current_preview_text = (
                        current_preview_text[:start] +
                        replacement['variable'] +
                        current_preview_text[end:]
                    )

            doc.save(str(save_path))
            if not same_target:
                os.replace(str(save_path), str(output_path))
            self._logger.info(f"Word文档已保存（含{len(self._variable_replacements)}处变量替换）: {output_path}")
        finally:
            if temp_path is not None and temp_path.exists():
                try:
                    temp_path.unlink()
                except OSError as exc:
                    self._logger.warning(f"清理临时归档文档失败: {temp_path} ({exc})")

    def _apply_recorded_replacement(
        self,
        doc,
        replacement: Dict[str, Any],
        preview_text: str
    ) -> bool:
        """按记录的选区信息应用单次替换。"""
        start = replacement.get('selection_start')
        end = replacement.get('selection_end')
        original = replacement.get('original', '')
        variable = replacement.get('variable', '')

        if not isinstance(start, int) or not isinstance(end, int):
            self._logger.warning("变量替换缺少有效选区信息")
            return False

        if start < 0 or end < start or end > len(preview_text):
            self._logger.warning(f"变量替换选区越界: {start}-{end}")
            return False

        if preview_text[start:end] != original:
            self._logger.warning(
                f"变量替换选区内容不匹配: 期望 '{original}'，实际 '{preview_text[start:end]}'"
            )
            return False

        paragraph_ranges = self._build_preview_paragraph_ranges(doc)
        for item in paragraph_ranges:
            para_start = item['start']
            para_end = item['end']
            if para_start <= start and end <= para_end:
                local_start = start - para_start
                local_end = end - para_start
                paragraph = item['paragraph']
                if paragraph.text[local_start:local_end] != original:
                    self._logger.warning("段落内容与记录的选中文本不一致")
                    return False
                return self._replace_text_in_paragraph_range(
                    paragraph,
                    local_start,
                    local_end,
                    variable
                )

        self._logger.warning("当前替换跨越多个段落，暂不支持保存")
        return False

    def _build_preview_paragraph_ranges(self, doc) -> List[Dict[str, Any]]:
        """构建与预览文本一致的段落位置映射。"""
        ranges = []
        cursor = 0
        paragraphs = list(doc.paragraphs)
        for index, paragraph in enumerate(paragraphs):
            text = paragraph.text
            start = cursor
            end = start + len(text)
            ranges.append({
                'paragraph': paragraph,
                'start': start,
                'end': end,
                'index': index,
            })
            cursor = end
            if index < len(paragraphs) - 1:
                cursor += 1
        return ranges

    def _replace_text_in_paragraph_range(
        self,
        paragraph,
        start: int,
        end: int,
        new_text: str
    ) -> bool:
        """在段落的精确字符区间内替换文本，并尽量保留格式。"""
        if start == end:
            return False

        run_map = []
        char_offset = 0
        for run_index, run in enumerate(paragraph.runs):
            run_text = run.text
            run_length = len(run_text)
            if run_length == 0:
                continue
            run_map.append((run_index, char_offset, char_offset + run_length))
            char_offset += run_length

        start_run_index = None
        end_run_index = None
        for run_position, (run_index, run_start, run_end) in enumerate(run_map):
            if run_start <= start < run_end:
                start_run_index = run_position
            if run_start < end <= run_end:
                end_run_index = run_position
                break

        if start_run_index is None or end_run_index is None:
            self._logger.warning("无法定位替换选区对应的 Word run")
            return False

        if start_run_index == end_run_index:
            run_idx = run_map[start_run_index][0]
            run = paragraph.runs[run_idx]
            offset_start = start - run_map[start_run_index][1]
            offset_end = end - run_map[start_run_index][1]
            run.text = run.text[:offset_start] + new_text + run.text[offset_end:]
            return True

        first_run_idx = run_map[start_run_index][0]
        last_run_idx = run_map[end_run_index][0]

        first_run = paragraph.runs[first_run_idx]
        last_run = paragraph.runs[last_run_idx]

        first_offset = start - run_map[start_run_index][1]
        last_offset = end - run_map[end_run_index][1]

        first_run.text = first_run.text[:first_offset] + new_text

        for middle_position in range(start_run_index + 1, end_run_index):
            paragraph.runs[run_map[middle_position][0]].text = ""

        last_run.text = last_run.text[last_offset:]
        return True

    def closeEvent(self, event) -> None:
        """关闭事件"""
        event.accept()
