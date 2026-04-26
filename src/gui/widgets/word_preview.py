# -*- coding: utf-8 -*-
"""Word 文档预览控件

使用 python-docx 进行文本预览，支持变量高亮显示。
仅支持 Windows 平台。
"""

from pathlib import Path
from typing import Optional, List
import re

from PySide6.QtWidgets import (
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QPushButton,
    QLabel,
    QTextEdit,
    QSplitter,
    QMenu,
)
from PySide6.QtCore import Signal, Qt
from PySide6.QtGui import QTextCharFormat, QBrush, QColor, QFont, QTextCursor

from src.utils.logger import get_logger
from src.gui.styles import APP_COLORS as COLORS, button_style
from src.utils.platform_utils import get_default_ui_font_family


class WordPreviewWidget(QWidget):
    """Word 文档预览控件

    使用 python-docx 进行文本解析和预览。
    支持变量高亮、缩放和文本选择。
    """

    # 信号：文本被选中
    text_selected = Signal(str)
    # 信号：文档加载完成
    document_loaded = Signal(bool)
    # 信号：页面变化
    page_changed = Signal(int, int)  # 当前页, 总页数
    # 信号：请求替换选中文字
    replace_requested = Signal()
    # 信号：请求撤销变量 (变量名, 是否单个撤销)
    undo_variable_requested = Signal(str, bool)

    def __init__(self, parent=None):
        super().__init__(parent)
        self._logger = get_logger()
        self._current_path: Optional[Path] = None
        self._zoom_level: int = 100
        self._document = None
        self._paragraphs: List[str] = []
        self._variables: List[str] = []

        self._setup_ui()

    def _setup_ui(self) -> None:
        """设置界面"""
        c = COLORS
        ui_font = get_default_ui_font_family()
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # 工具栏
        toolbar = QWidget()
        toolbar.setStyleSheet(f"background: {c['surface_0']}; border-bottom: 1px solid {c['border']};")
        toolbar_layout = QHBoxLayout(toolbar)
        toolbar_layout.setContentsMargins(10, 8, 10, 8)
        toolbar_layout.setSpacing(8)

        # 缩放控制
        zoom_out_btn = QPushButton("-")
        zoom_out_btn.setFixedSize(28, 28)
        zoom_out_btn.setStyleSheet(button_style(compact=True))
        zoom_out_btn.clicked.connect(self._zoom_out)
        toolbar_layout.addWidget(zoom_out_btn)

        self._zoom_label = QLabel("100%")
        self._zoom_label.setMinimumWidth(50)
        self._zoom_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        toolbar_layout.addWidget(self._zoom_label)

        zoom_in_btn = QPushButton("+")
        zoom_in_btn.setFixedSize(28, 28)
        zoom_in_btn.setStyleSheet(button_style(compact=True))
        zoom_in_btn.clicked.connect(self._zoom_in)
        toolbar_layout.addWidget(zoom_in_btn)

        toolbar_layout.addSpacing(20)

        # 变量统计
        self._var_count_label = QLabel("变量: 0")
        self._var_count_label.setStyleSheet(f"color: {c['accent']}; font-weight: 700;")
        toolbar_layout.addWidget(self._var_count_label)

        toolbar_layout.addStretch()

        # 文档信息
        self._info_label = QLabel("")
        self._info_label.setStyleSheet(f"color: {c['text_tertiary']};")
        toolbar_layout.addWidget(self._info_label)

        layout.addWidget(toolbar)

        # 预览区域
        self._text_edit = QTextEdit()
        self._text_edit.setReadOnly(True)
        self._text_edit.setAcceptRichText(True)
        self._text_edit.setStyleSheet(f"""
            QTextEdit {{
                border: none;
                padding: 20px;
                font-family: '{ui_font}', 'Microsoft YaHei', 'SimSun';
                font-size: 14px;
                line-height: 1.8;
                background: {c['surface_0']};
                color: {c['text_primary']};
            }}
        """)
        self._text_edit.textChanged.connect(self._on_text_changed)
        self._text_edit.selectionChanged.connect(self._on_selection_changed)
        # 设置自定义右键菜单
        self._text_edit.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self._text_edit.customContextMenuRequested.connect(self._show_context_menu)

        layout.addWidget(self._text_edit, 1)

        # 状态栏
        self._status_label = QLabel("就绪 - 选择 Word 文档进行预览")
        self._status_label.setStyleSheet(f"""
            padding: 8px 12px;
            background: {c['surface_0']};
            color: {c['text_tertiary']};
            border-top: 1px solid {c['border']};
        """)
        layout.addWidget(self._status_label)

    def _detect_variable_at_position(self, pos) -> Optional[str]:
        """检测指定位置是否在变量上

        只有当光标位置确实在 {{xxx}} 变量的范围内时才返回变量名。

        Args:
            pos: 鼠标位置（控件坐标）

        Returns:
            变量名（如 "client_name"），如果不是变量返回 None
        """
        # 获取该位置的文本光标
        cursor = self._text_edit.cursorForPosition(pos)

        # 获取光标在文档中的绝对位置
        current_pos = cursor.position()

        # 获取整个文档的文本
        full_text = self._text_edit.toPlainText()

        # 使用正则表达式找到文档中所有的变量及其位置
        var_pattern = re.compile(r'\{\{([^}]+)\}\}')

        for match in var_pattern.finditer(full_text):
            var_start = match.start()
            var_end = match.end()
            var_text = match.group(0)  # 完整的 {{xxx}}
            var_name = match.group(1).strip()  # 变量名 xxx

            # 检查当前光标位置是否在这个变量的范围内
            if var_start <= current_pos <= var_end:
                self._logger.info(f"检测到变量: {var_name}, 位置: {var_start}-{var_end}, 光标: {current_pos}")
                return var_name

        return None

    def _show_context_menu(self, pos):
        """显示自定义右键菜单"""
        # 创建自定义菜单
        menu = QMenu(self)

        # 首先检测是否点击了变量
        variable_name = self._detect_variable_at_position(pos)

        if variable_name:
            # 在变量上 - 显示撤销菜单
            undo_single = menu.addAction(f"撤销此变量")
            undo_single.triggered.connect(lambda: self._on_undo_single(variable_name))

            menu.addSeparator()

            undo_all = menu.addAction(f"撤销所有 {{{{{variable_name}}}}}")
            undo_all.triggered.connect(lambda: self._on_undo_all(variable_name))

        else:
            # 不是变量，检查是否有选中文本
            cursor = self._text_edit.textCursor()
            has_selection = cursor.hasSelection()

            if has_selection:
                # 添加"替换选中"菜单项
                replace_action = menu.addAction("替换选中...")
                replace_action.triggered.connect(self._on_replace_action)
            else:
                # 没有选中文本时显示提示
                no_select_action = menu.addAction("请先选择文字")
                no_select_action.setEnabled(False)

        # 在鼠标位置显示菜单
        global_pos = self._text_edit.mapToGlobal(pos)
        menu.exec(global_pos)

    def _on_replace_action(self):
        """右键菜单"替换选中"点击处理"""
        self._logger.info("右键菜单触发替换请求")
        self.replace_requested.emit()

    def _on_undo_single(self, variable_name: str):
        """单个撤销变量"""
        self._logger.info(f"请求单个撤销变量: {variable_name}")
        self.undo_variable_requested.emit(variable_name, True)

    def _on_undo_all(self, variable_name: str):
        """整体撤销变量"""
        self._logger.info(f"请求整体撤销变量: {variable_name}")
        self.undo_variable_requested.emit(variable_name, False)

    def load_file(self, path: Path) -> bool:
        """加载 Word 文件

        Args:
            path: 文件路径

        Returns:
            是否加载成功
        """
        if not path.exists():
            self._logger.error(f"文件不存在: {path}")
            self._status_label.setText(f"错误: 文件不存在")
            return False

        self._current_path = path

        # 检查文件格式
        if path.suffix.lower() == '.doc':
            self._status_label.setText("提示: .doc 格式支持有限，建议转换为 .docx")
            # 仍然尝试加载

        return self._load_with_python_docx(path)

    def _load_with_python_docx(self, path: Path) -> bool:
        """使用 python-docx 加载文档"""
        try:
            from docx import Document

            self._document = Document(str(path))
            self._paragraphs = []

            # 提取段落文本
            for para in self._document.paragraphs:
                self._paragraphs.append(para.text)

            # 提取表格文本
            for table in self._document.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        self._paragraphs.append(f"[表格] {row_text}")

            # 显示带高亮的文本
            self._display_with_highlight()

            # 提取变量
            self._variables = self._extract_variables()

            # 更新状态
            self._status_label.setText(f"已加载: {path.name}")
            self._info_label.setText(f"段落: {len(self._paragraphs)}")
            self._var_count_label.setText(f"变量: {len(self._variables)}")

            self.document_loaded.emit(True)
            self._logger.info(f"文档加载成功: {path.name}, 段落数: {len(self._paragraphs)}, 变量数: {len(self._variables)}")
            return True

        except Exception as e:
            self._logger.error(f"python-docx 加载失败: {e}")
            self._status_label.setText(f"加载失败: {e}")
            self.document_loaded.emit(False)
            return False

    def _display_with_highlight(self) -> None:
        """显示文本并高亮变量"""
        self._text_edit.clear()

        cursor = self._text_edit.textCursor()

        for i, para_text in enumerate(self._paragraphs):
            if i > 0:
                cursor.insertBlock()

            self._insert_text_with_highlight(cursor, para_text)

    def _insert_text_with_highlight(self, cursor: QTextCursor, text: str) -> None:
        """插入文本并高亮变量

        Args:
            cursor: 文本光标
            text: 要插入的文本
        """
        # 变量匹配模式
        var_pattern = re.compile(r'(\{\{[^}]+\}\})')

        # 分割文本
        parts = var_pattern.split(text)

        for part in parts:
            if var_pattern.match(part):
                # 这是变量，使用绿色高亮格式
                var_fmt = QTextCharFormat()
                var_fmt.setBackground(QBrush(QColor(212, 237, 218)))  # 浅绿色背景
                var_fmt.setForeground(QBrush(QColor(21, 87, 36)))    # 深绿色文字
                var_fmt.setFontWeight(QFont.Bold)
                cursor.insertText(part, var_fmt)
            else:
                # 普通文本，显式设置默认格式以重置之前的高亮
                normal_fmt = QTextCharFormat()
                normal_fmt.setBackground(QBrush(QColor(250, 250, 250)))  # 背景色与 QTextEdit 背景一致
                normal_fmt.setForeground(QBrush(QColor(0, 0, 0)))        # 黑色文字
                normal_fmt.setFontWeight(QFont.Normal)
                cursor.insertText(part, normal_fmt)

    def _extract_variables(self) -> List[str]:
        """从文档中提取变量"""
        if not self._paragraphs:
            return []

        pattern = re.compile(r'\{\{([^}]+)\}\}')
        variables = []

        for text in self._paragraphs:
            found = pattern.findall(text)
            variables.extend(found)

        # 去重并排序
        return sorted(list(set(v.strip() for v in variables)))

    def _on_text_changed(self) -> None:
        """文本变化"""
        pass

    def _on_selection_changed(self) -> None:
        """选择变化

        注意：QTextCursor.selectedText() 在文本跨越段落时，
        使用 U+2029 (Paragraph Separator) 代替换行符。
        这里将其替换为普通换行符。
        """
        cursor = self._text_edit.textCursor()
        if cursor.hasSelection():
            selected = cursor.selectedText()
            # 替换 Unicode 段落分隔符为普通换行符
            selected = selected.replace('\u2029', '\n')
            self._selected_text = selected
            self.text_selected.emit(selected)
        else:
            self._selected_text = ""
            self.text_selected.emit("")

    def get_selected_text(self) -> str:
        """获取选中的文本"""
        cursor = self._text_edit.textCursor()
        if cursor.hasSelection():
            return cursor.selectedText()
        return ""

    def highlight_variables(self, variables: List[str]) -> None:
        """高亮显示指定变量（重新渲染）

        Args:
            variables: 变量名列表
        """
        # 重新显示带高亮的文本
        self._display_with_highlight()

    def _zoom_in(self) -> None:
        """放大"""
        if self._zoom_level < 200:
            self._zoom_level += 10
            self._apply_zoom()

    def _zoom_out(self) -> None:
        """缩小"""
        if self._zoom_level > 50:
            self._zoom_level -= 10
            self._apply_zoom()

    def _apply_zoom(self) -> None:
        """应用缩放"""
        self._zoom_label.setText(f"{self._zoom_level}%")
        # QTextEdit 不直接支持缩放，通过字体大小模拟
        base_size = 14
        new_size = int(base_size * self._zoom_level / 100)
        font = self._text_edit.font()
        font.setPointSize(new_size)
        self._text_edit.setFont(font)

    def close_document(self) -> None:
        """关闭当前文档"""
        self._document = None
        self._paragraphs = []
        self._variables = []
        self._text_edit.clear()
        self._current_path = None
        self._status_label.setText("就绪 - 选择 Word 文档进行预览")
        self._info_label.setText("")
        self._var_count_label.setText("变量: 0")

    def cleanup(self) -> None:
        """清理资源"""
        self.close_document()

    def get_current_path(self) -> Optional[Path]:
        """获取当前文档路径"""
        return self._current_path

    def is_word_available(self) -> bool:
        """检查 Word 是否可用（始终返回 True，因为我们使用 python-docx）"""
        return True

    def get_variables(self) -> List[str]:
        """获取文档中的变量列表"""
        return self._variables.copy()

    def set_preview_text(self, text: str) -> None:
        """设置预览文本（用于从编辑器更新预览）

        Args:
            text: 要显示的文本（段落用换行分隔）
        """
        # 将文本分割为段落
        self._paragraphs = text.split("\n")

        # 重新显示带高亮的文本
        self._display_with_highlight()

        # 重新提取变量
        self._variables = self._extract_variables()

        # 更新状态
        self._var_count_label.setText(f"变量: {len(self._variables)}")
        self._info_label.setText(f"段落: {len(self._paragraphs)}")
        self._status_label.setText("预览已更新")

    def get_document_text(self) -> str:
        """获取文档全部文本"""
        return "\n\n".join(self._paragraphs)

    def find_text(self, search_text: str) -> bool:
        """查找文本

        Args:
            search_text: 要查找的文本

        Returns:
            是否找到
        """
        if not search_text:
            return False

        # 使用 QTextEdit 的查找功能
        found = self._text_edit.find(search_text)
        if not found:
            # 从头开始查找
            cursor = self._text_edit.textCursor()
            cursor.movePosition(QTextCursor.MoveOperation.Start)
            self._text_edit.setTextCursor(cursor)
            found = self._text_edit.find(search_text)

        return found

    def replace_selection(self, new_text: str) -> bool:
        """替换选中的文本（仅显示层面，不修改原文件）

        Args:
            new_text: 新文本

        Returns:
            是否替换成功
        """
        cursor = self._text_edit.textCursor()
        if cursor.hasSelection():
            cursor.insertText(new_text)
            return True
        return False
