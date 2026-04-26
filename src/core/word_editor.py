# -*- coding: utf-8 -*-
"""Word 文档编辑器核心模块

提供 Word 文档的加载、编辑、保存功能，支持格式保持的文本替换。
"""

import io
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union
import re
import shutil

from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

from src.utils.logger import get_logger


class WordEditor:
    """Word 文档编辑器

    使用 python-docx 进行文档编辑，支持格式保持的文本替换。
    """

    def __init__(self):
        self._logger = get_logger()
        self._document: Optional[DocxDocument] = None
        self._current_path: Optional[Path] = None
        self._initial_snapshot: Optional[bytes] = None
        self._is_modified: bool = False
        self._history: List[Dict[str, Any]] = []  # 操作历史，用于撤销

    def load_document(self, path: Path) -> bool:
        """加载 Word 文档

        Args:
            path: 文档路径

        Returns:
            是否加载成功
        """
        try:
            from docx import Document
        except ImportError:
            self._logger.error("python-docx 未安装")
            return False

        if not path.exists():
            self._logger.error(f"文件不存在: {path}")
            return False

        if path.suffix.lower() not in ['.docx', '.doc']:
            self._logger.error(f"不支持的文件格式: {path.suffix}")
            return False

        try:
            self._document = Document(str(path))
            self._current_path = path
            self._initial_snapshot = self._snapshot_document()
            self._is_modified = False
            self._history.clear()
            self._logger.info(f"文档加载成功: {path}")
            return True
        except Exception as e:
            self._logger.error(f"文档加载失败: {e}")
            return False

    def save_document(self, path: Optional[Path] = None) -> bool:
        """保存文档

        Args:
            path: 保存路径，为 None 则保存到原路径

        Returns:
            是否保存成功
        """
        if self._document is None:
            self._logger.warning("没有打开的文档")
            return False

        save_path = path or self._current_path
        if save_path is None:
            self._logger.error("未指定保存路径")
            return False

        try:
            save_path.parent.mkdir(parents=True, exist_ok=True)
            self._document.save(str(save_path))
            self._current_path = save_path
            self._is_modified = False
            self._logger.info(f"文档保存成功: {save_path}")
            return True
        except Exception as e:
            self._logger.error(f"文档保存失败: {e}")
            return False

    def save_as(self, path: Path) -> bool:
        """另存为

        Args:
            path: 新的保存路径

        Returns:
            是否保存成功
        """
        return self.save_document(path)

    def _snapshot_document(self) -> Optional[bytes]:
        """将当前文档保存为内存快照。"""
        if self._document is None:
            return None

        try:
            buffer = io.BytesIO()
            self._document.save(buffer)
            return buffer.getvalue()
        except Exception as e:
            self._logger.error(f"创建文档快照失败: {e}")
            return None

    def _restore_from_snapshot(self, snapshot: bytes) -> bool:
        """从内存快照恢复文档。"""
        try:
            from docx import Document
            self._document = Document(io.BytesIO(snapshot))
            return True
        except Exception as e:
            self._logger.error(f"从快照恢复文档失败: {e}")
            return False

    def _rebuild_document_from_history(self) -> bool:
        """从初始快照重放历史操作，恢复当前文档状态。"""
        if self._initial_snapshot is None:
            self._logger.error("缺少初始文档快照，无法重建文档")
            return False

        history = list(self._history)
        if not self._restore_from_snapshot(self._initial_snapshot):
            return False

        for entry in history:
            count, _ = self._replace_text_impl(
                entry.get("old_text", ""),
                entry.get("new_text", ""),
                entry.get("replace_all", True),
                record_history=False
            )
            if count == 0:
                self._logger.warning(
                    f"重放历史操作未命中: {entry.get('old_text')} -> {entry.get('new_text')}"
                )

        self._history = history
        self._is_modified = bool(self._history)
        return True

    def replace_text(
        self,
        old_text: str,
        new_text: str,
        replace_all: bool = True
    ) -> Tuple[int, List[Dict]]:
        """替换文本（格式保持）

        在 Run 级别进行文本替换，确保格式不丢失。

        Args:
            old_text: 要替换的文本
            new_text: 替换后的文本
            replace_all: 是否替换所有匹配项

        Returns:
            (替换次数, 替换位置列表)
        """
        return self._replace_text_impl(
            old_text,
            new_text,
            replace_all,
            record_history=True
        )

    def _replace_text_impl(
        self,
        old_text: str,
        new_text: str,
        replace_all: bool = True,
        record_history: bool = True
    ) -> Tuple[int, List[Dict]]:
        """执行文本替换，可选择是否记录历史。"""
        if self._document is None:
            self._logger.warning("没有打开的文档")
            return 0, []

        if not old_text:
            self._logger.warning("替换文本不能为空")
            return 0, []

        count = 0
        replacements = []

        history_entry = {
            "type": "replace",
            "old_text": old_text,
            "new_text": new_text,
            "replace_all": replace_all,
            "replacements": []
        }

        # 替换段落中的文本
        for para_idx, paragraph in enumerate(self._document.paragraphs):
            para_count, para_replacements = self._replace_in_paragraph(
                paragraph, old_text, new_text, replace_all, para_idx
            )
            count += para_count
            replacements.extend(para_replacements)
            history_entry["replacements"].extend(para_replacements)

            if not replace_all and count > 0:
                break

        # 替换表格中的文本
        if replace_all or count == 0:
            for table_idx, table in enumerate(self._document.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            para_count, para_replacements = self._replace_in_paragraph(
                                paragraph, old_text, new_text, replace_all,
                                f"table_{table_idx}_{row_idx}_{cell_idx}_{para_idx}"
                            )
                            count += para_count
                            replacements.extend(para_replacements)
                            history_entry["replacements"].extend(para_replacements)

                            if not replace_all and count > 0:
                                break
                        if not replace_all and count > 0:
                            break
                    if not replace_all and count > 0:
                        break
                if not replace_all and count > 0:
                    break

        if count > 0:
            self._is_modified = True
            if record_history:
                self._history.append(history_entry)
            self._logger.info(f"替换完成: '{old_text}' -> '{new_text}', 共 {count} 处")

        return count, replacements

    def _replace_in_paragraph(
        self,
        paragraph: Paragraph,
        old_text: str,
        new_text: str,
        replace_all: bool,
        location_id: Union[int, str]
    ) -> Tuple[int, List[Dict[str, Any]]]:
        """在段落中替换文本

        支持跨Run的文本替换，处理Word文档中文本被分割到多个Run的情况。

        Args:
            paragraph: 段落对象
            old_text: 要替换的文本
            new_text: 替换后的文本
            replace_all: 是否替换所有
            location_id: 位置标识

        Returns:
            (替换次数, 替换位置列表)
        """
        count = 0
        replacements = []

        # 首先检查段落中是否包含要替换的文本
        if old_text not in paragraph.text:
            return count, replacements

        # 记录需要处理的Run索引和替换位置
        # 使用段落级别的文本处理，然后映射回Run
        paragraph_text = paragraph.text

        # 找到所有匹配位置
        search_len = len(old_text)
        start_pos = 0
        match_positions = []

        while True:
            pos = paragraph_text.find(old_text, start_pos)
            if pos == -1:
                break
            match_positions.append(pos)
            if not replace_all:
                break
            start_pos = pos + search_len

        if not match_positions:
            return count, replacements

        # 构建Run映射表：记录每个字符位置属于哪个Run
        run_map = []  # [(run_idx, run_start, run_end, run_text), ...]
        char_offset = 0

        for run_idx, run in enumerate(paragraph.runs):
            run_text = run.text
            run_len = len(run_text)
            if run_len > 0:
                run_map.append((run_idx, char_offset, char_offset + run_len, run_text))
                char_offset += run_len

        # 从后向前处理每个匹配位置（避免位置偏移问题）
        for match_start in reversed(match_positions):
            match_end = match_start + search_len

            # 找到包含匹配文本起始和结束的Run
            start_run_idx = None
            end_run_idx = None

            for i, (run_idx, run_start, run_end, _) in enumerate(run_map):
                if run_start <= match_start < run_end:
                    start_run_idx = i
                if run_start < match_end <= run_end:
                    end_run_idx = i
                    break

            if start_run_idx is None or end_run_idx is None:
                self._logger.warning(f"无法定位匹配位置: {match_start}-{match_end}")
                continue

            # 获取第一个Run的格式（用于新文本）
            first_run = paragraph.runs[run_map[start_run_idx][0]]

            # 处理跨Run替换
            if start_run_idx == end_run_idx:
                # 简单情况：替换在同一个Run内
                run_idx = run_map[start_run_idx][0]
                run = paragraph.runs[run_idx]
                original_text = run.text
                relative_start = match_start - run_map[start_run_idx][1]
                new_run_text = original_text[:relative_start] + new_text + original_text[relative_start + search_len:]
                run.text = new_run_text

                replacements.append({
                    "location": location_id,
                    "run_idx": run_idx,
                    "original": original_text,
                    "replaced": new_run_text
                })
            else:
                # 复杂情况：替换跨多个Run
                # 策略：保留第一个Run，清空中间的Run，删除后面的Run

                # 第一个Run：保留匹配前的内容 + 新文本
                first_run_idx = run_map[start_run_idx][0]
                first_run_obj = paragraph.runs[first_run_idx]
                first_run_text = first_run_obj.text
                first_relative_start = match_start - run_map[start_run_idx][1]
                first_run_obj.text = first_run_text[:first_relative_start] + new_text

                # 中间的Run：清空
                for i in range(start_run_idx + 1, end_run_idx):
                    mid_run_idx = run_map[i][0]
                    if mid_run_idx < len(paragraph.runs):
                        paragraph.runs[mid_run_idx].text = ""

                # 最后一个Run：保留匹配后的内容
                last_run_idx = run_map[end_run_idx][0]
                last_run_obj = paragraph.runs[last_run_idx]
                last_run_text = last_run_obj.text
                last_relative_end = match_end - run_map[end_run_idx][1]
                last_run_obj.text = last_run_text[last_relative_end:]

                replacements.append({
                    "location": location_id,
                    "run_idx": first_run_idx,
                    "original": paragraph.text,  # 简化记录
                    "replaced": new_text
                })

            count += 1

            # 更新run_map以反映已做的修改（简化处理：重新构建）
            run_map = []
            char_offset = 0
            for run_idx, run in enumerate(paragraph.runs):
                run_text = run.text
                run_len = len(run_text)
                if run_len > 0:
                    run_map.append((run_idx, char_offset, char_offset + run_len, run_text))
                    char_offset += run_len

        return count, replacements

    def find_text(self, search_text: str) -> List[Dict]:
        """查找文本

        Args:
            search_text: 要查找的文本

        Returns:
            匹配位置列表
        """
        if self._document is None:
            return []

        matches = []

        # 在段落中查找
        for para_idx, paragraph in enumerate(self._document.paragraphs):
            if search_text in paragraph.text:
                matches.append({
                    "type": "paragraph",
                    "index": para_idx,
                    "text": paragraph.text,
                    "preview": self._get_text_preview(paragraph.text, search_text)
                })

        # 在表格中查找
        for table_idx, table in enumerate(self._document.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        if search_text in paragraph.text:
                            matches.append({
                                "type": "table",
                                "table_index": table_idx,
                                "row_index": row_idx,
                                "cell_index": cell_idx,
                                "para_index": para_idx,
                                "text": paragraph.text,
                                "preview": self._get_text_preview(paragraph.text, search_text)
                            })

        return matches

    def _get_text_preview(self, text: str, search_text: str, context_len: int = 20) -> str:
        """获取文本预览

        Args:
            text: 完整文本
            search_text: 搜索文本
            context_len: 上下文长度

        Returns:
            预览文本
        """
        idx = text.find(search_text)
        if idx == -1:
            return text[:50] + "..." if len(text) > 50 else text

        start = max(0, idx - context_len)
        end = min(len(text), idx + len(search_text) + context_len)

        prefix = "..." if start > 0 else ""
        suffix = "..." if end < len(text) else ""

        return prefix + text[start:end] + suffix

    def extract_all_text(self) -> str:
        """提取文档所有文本

        Returns:
            文档文本内容
        """
        if self._document is None:
            return ""

        text_parts = []

        for paragraph in self._document.paragraphs:
            text_parts.append(paragraph.text)

        for table in self._document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text_parts.append(paragraph.text)

        return "\n".join(text_parts)

    def extract_variables(self) -> List[str]:
        """提取文档中的变量

        Returns:
            变量名列表
        """
        text = self.extract_all_text()
        pattern = re.compile(r'\{\{(\w+)\}\}')
        variables = pattern.findall(text)
        # 去重并排序
        return sorted(list(set(v.strip() for v in variables)))

    def get_variable_count(self, variable_name: str) -> int:
        """获取变量出现次数

        Args:
            variable_name: 变量名（不含 {{}}）

        Returns:
            出现次数
        """
        text = self.extract_all_text()
        return text.count(f"{{{{{variable_name}}}}}")

    def is_modified(self) -> bool:
        """检查文档是否已修改

        Returns:
            是否已修改
        """
        return self._is_modified

    def can_undo(self) -> bool:
        """检查是否可以撤销

        Returns:
            是否可以撤销
        """
        return len(self._history) > 0

    def undo(self) -> bool:
        """撤销上一步操作

        Returns:
            是否撤销成功
        """
        if not self._history:
            self._logger.info("没有可撤销的操作")
            return False

        removed_entry = self._history.pop()
        success = self._rebuild_document_from_history()
        if not success:
            self._history.append(removed_entry)
            self._rebuild_document_from_history()
        if success:
            self._logger.info("已撤销上一步操作")
        return success

    def undo_variable(self, variable_name: str, single: bool = False) -> int:
        """撤销变量替换

        将变量替换回原始文字。支持单个撤销和整体撤销。

        Args:
            variable_name: 变量名（如 "client_name" 或 "{{client_name}}"）
            single: True=只撤销一个，False=撤销所有同名变量

        Returns:
            撤销的数量
        """
        if not self._document:
            self._logger.warning("没有打开的文档")
            return 0

        # 确保变量名格式正确
        if not variable_name.startswith("{{"):
            variable_pattern = f"{{{{{variable_name}}}}}"
        else:
            variable_pattern = variable_name

        # 从历史记录中找到原始文字
        original_text = None
        for entry in reversed(self._history):
            if entry.get("new_text") == variable_pattern:
                original_text = entry.get("old_text")
                break

        if not original_text:
            self._logger.warning(f"未找到变量 {variable_pattern} 的原始文字")
            return 0

        self._logger.info(f"撤销变量: {variable_pattern}, single={single}")

        matched_indexes = [
            idx for idx, entry in enumerate(self._history)
            if entry.get("new_text") == variable_pattern
        ]
        if not matched_indexes:
            return 0

        if single:
            matched_indexes = [matched_indexes[-1]]

        removed_indexes = set(matched_indexes)
        removed_entries = [self._history[idx] for idx in matched_indexes]
        count = sum(len(entry.get("replacements", [])) for entry in removed_entries)

        previous_history = list(self._history)
        self._history = [
            entry for idx, entry in enumerate(self._history)
            if idx not in removed_indexes
        ]

        if not self._rebuild_document_from_history():
            self._logger.error("撤销变量失败：无法重建文档")
            self._history = previous_history
            self._rebuild_document_from_history()
            return 0

        self._logger.info(f"撤销完成，共撤销 {count} 处")
        return count

    def get_variable_original_text(self, variable_name: str) -> Optional[str]:
        """获取变量的原始文字

        Args:
            variable_name: 变量名

        Returns:
            原始文字，未找到返回 None
        """
        if not variable_name.startswith("{{"):
            variable_pattern = f"{{{{{variable_name}}}}}"
        else:
            variable_pattern = variable_name

        for entry in reversed(self._history):
            if entry.get("new_text") == variable_pattern:
                return entry.get("old_text")

        return None

    def get_current_path(self) -> Optional[Path]:
        """获取当前文档路径

        Returns:
            当前文档路径
        """
        return self._current_path

    def is_loaded(self) -> bool:
        """检查是否已加载文档

        Returns:
            是否已加载
        """
        return self._document is not None

    def close(self) -> None:
        """关闭当前文档"""
        self._document = None
        self._current_path = None
        self._initial_snapshot = None
        self._is_modified = False
        self._history.clear()


def upload_to_template_category(source_path: Path, category: str, template_dir: Path) -> Path:
    """上传文件到模板分类目录

    Args:
        source_path: 源文件路径
        category: 分类（civil, criminal, non_litigation）
        template_dir: 模板根目录

    Returns:
        目标文件路径
    """
    target_dir = template_dir / category
    target_dir.mkdir(parents=True, exist_ok=True)

    target_path = target_dir / source_path.name

    # 如果文件已存在，添加序号
    if target_path.exists():
        counter = 1
        while True:
            new_name = f"{source_path.stem}_{counter}{source_path.suffix}"
            target_path = target_dir / new_name
            if not target_path.exists():
                break
            counter += 1

    shutil.copy2(source_path, target_path)
    return target_path
