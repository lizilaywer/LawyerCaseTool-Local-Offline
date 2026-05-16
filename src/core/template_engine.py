# -*- coding: utf-8 -*-
"""Word模板处理引擎模块"""

import hashlib
import io
import os
import re
import shutil
import threading
import time
from collections import OrderedDict
from pathlib import Path
from typing import Any, Dict, Optional, Tuple

from src.utils.logger import get_logger
from src.utils.exceptions import TemplateError, TemplateFileError


class TemplateEngine:
    """Word模板处理引擎

    支持模板缓存以提高性能，避免重复加载同一模板。
    缓存会跟踪文件的修改时间，自动失效过期的缓存。
    """

    # 缓存过期时间（秒）
    CACHE_TTL = 300  # 5分钟
    # 缓存最大条目数
    CACHE_MAXSIZE = 50

    def __init__(self):
        self._logger = get_logger()
        # 模板字节缓存: {path_str: (bytes_data, file_mtime, cache_time)}，LRU 淘汰
        self._cache: OrderedDict[str, Tuple[bytes, float, float]] = OrderedDict()
        self._cache_lock = threading.Lock()

    def _get_cache_key(self, template_path: Path) -> str:
        """生成缓存键"""
        return str(template_path.resolve())

    def _get_file_mtime(self, template_path: Path) -> float:
        """获取文件修改时间"""
        try:
            return os.path.getmtime(template_path)
        except OSError:
            return 0

    def _is_cache_valid(self, cache_key: str, current_mtime: float) -> bool:
        """检查缓存是否有效"""
        if cache_key not in self._cache:
            return False

        cached_bytes, cached_mtime, cache_time = self._cache[cache_key]

        # 检查文件是否被修改
        if cached_mtime != current_mtime:
            return False

        # 检查缓存是否过期
        if time.time() - cache_time > self.CACHE_TTL:
            return False

        # LRU: 访问时移到末尾
        self._cache.move_to_end(cache_key)
        return True

    def _get_cached_template(self, template_path: Path) -> Optional[bytes]:
        """从缓存获取模板字节"""
        cache_key = self._get_cache_key(template_path)
        current_mtime = self._get_file_mtime(template_path)

        with self._cache_lock:
            if self._is_cache_valid(cache_key, current_mtime):
                self._logger.debug(f"从缓存加载模板: {template_path.name}")
                return self._cache[cache_key][0]
        return None

    def _cache_template(self, template_path: Path, template_bytes: bytes) -> None:
        """缓存模板字节，LRU 淘汰超限条目"""
        cache_key = self._get_cache_key(template_path)
        current_mtime = self._get_file_mtime(template_path)

        with self._cache_lock:
            self._cache[cache_key] = (template_bytes, current_mtime, time.time())
            # LRU: 移到末尾（最近使用）
            self._cache.move_to_end(cache_key)
            # 淘汰最久未使用的条目
            while len(self._cache) > self.CACHE_MAXSIZE:
                self._cache.popitem(last=False)
            self._logger.debug(f"模板已缓存: {template_path.name}")

    def clear_cache(self) -> None:
        """清除所有缓存"""
        with self._cache_lock:
            self._cache.clear()
            self._logger.info("模板缓存已清除")

    def process_template(
        self,
        template_path: Path,
        output_path: Path,
        values: Dict[str, Any]
    ) -> Path:
        """
        处理 Word 模板，替换变量并保存

        Args:
            template_path: 模板文件路径
            output_path: 输出文件路径
            values: 变量值字典

        Returns:
            输出文件路径

        Raises:
            TemplateFileError: 模板文件错误
        """
        try:
            from docxtpl import DocxTemplate
        except ImportError:
            self._logger.error("docxtpl 未安装")
            raise TemplateError("请安装 docxtpl: pip install docxtpl")

        # 检查模板文件
        if not template_path.exists():
            raise TemplateFileError(str(template_path), "文件不存在")

        if not template_path.suffix.lower() == '.docx':
            raise TemplateFileError(str(template_path), "必须是 .docx 文件")

        try:
            # 尝试从缓存获取模板字节
            template_bytes = self._get_cached_template(template_path)

            if template_bytes is not None:
                # 从缓存的字节创建模板
                doc = DocxTemplate(io.BytesIO(template_bytes))
            else:
                # 统一先读取字节，避免重复 I/O
                with open(template_path, 'rb') as f:
                    template_bytes = f.read()
                doc = DocxTemplate(io.BytesIO(template_bytes))
                self._cache_template(template_path, template_bytes)

        except FileNotFoundError as e:
            self._logger.error(f"模板文件不存在: {e}")
            raise TemplateFileError(str(template_path), "文件不存在或无法访问")
        except PermissionError as e:
            self._logger.error(f"无权限访问模板文件: {e}")
            raise TemplateFileError(str(template_path), "无权限访问文件")
        except KeyError as e:
            self._logger.error(f"模板格式错误: {e}")
            raise TemplateFileError(str(template_path), f"模板格式错误: 缺少键 {e}")
        except Exception as e:
            self._logger.error(f"模板加载失败: {e}")
            raise TemplateFileError(str(template_path), f"模板加载失败: {e}")

        try:
            # 替换变量
            context = self._prepare_context(values)
            doc.render(context)

            # 保存
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))

            self._logger.info(f"模板处理完成: {output_path}")
            return output_path

        except PermissionError as e:
            self._logger.error(f"无权限保存文件: {e}")
            raise TemplateFileError(str(output_path), "无权限保存文件")
        except OSError as e:
            self._logger.error(f"文件系统错误: {e}")
            raise TemplateFileError(str(output_path), f"文件系统错误: {e}")
        except Exception as e:
            self._logger.error(f"模板处理失败: {e}")
            raise TemplateFileError(str(template_path), str(e))

    def _prepare_context(self, values: Dict[str, Any]) -> Dict[str, Any]:
        """
        准备模板上下文

        Args:
            values: 变量值字典

        Returns:
            处理后的上下文字典
        """
        context = {}

        for key, value in values.items():
            placeholder = f"{{{{{key}}}}}"
            # docxtpl 基于 Jinja，缺失变量会被渲染为空串。
            # 这里显式回填原始占位符，确保“待后补”的字段不会被静默抹掉。
            if value is None:
                self._logger.debug(f"变量 {key} 值为 None，保留模板中的 {placeholder} 占位符")
                context[key] = placeholder
            # 处理空字符串 - 同样保留原始变量占位符
            elif isinstance(value, str) and not value.strip():
                self._logger.debug(f"变量 {key} 值为空字符串，保留模板中的 {placeholder} 占位符")
                context[key] = placeholder
            # 处理日期
            elif hasattr(value, 'strftime'):
                context[key] = value.strftime("%Y年%m月%d日")
            else:
                context[key] = str(value)

        return context

    def extract_variables(self, template_path: Path) -> list:
        """
        从 Word 模板中提取变量

        Args:
            template_path: 模板文件路径

        Returns:
            变量名列表
        """
        # 检查文件扩展名，.doc 文件需要先转换
        if template_path.suffix.lower() == '.doc':
            self._logger.warning(f".doc 格式文件暂不支持变量提取: {template_path}")
            return []

        # 使用 python-docx 直接读取文档内容
        try:
            from docx import Document
        except ImportError:
            self._logger.error("python-docx 未安装")
            return []

        try:
            doc = Document(str(template_path))
        except FileNotFoundError:
            self._logger.error(f"模板文件不存在: {template_path}")
            return []
        except PermissionError:
            self._logger.error(f"无权限访问模板文件: {template_path}")
            return []
        except Exception as e:
            self._logger.error(f"加载模板失败: {e}")
            return []

        # 匹配 {{variable_name}} 格式，变量名为字母、数字、下划线
        pattern = re.compile(r'\{\{(\w+)\}\}')

        variables = []

        try:
            # 从段落中提取变量
            for para in doc.paragraphs:
                found = pattern.findall(para.text)
                variables.extend(found)

            # 从表格中提取变量
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        found = pattern.findall(cell.text)
                        variables.extend(found)
        except Exception as e:
            self._logger.warning(f"提取变量时发生错误: {e}")

        # 清理变量名（去除空格）
        variables = [v.strip() for v in variables]

        # 去重并排序
        unique_variables = sorted(list(set(variables)))

        self._logger.debug(f"从模板 {template_path.name} 中提取到 {len(unique_variables)} 个变量: {unique_variables}")
        return unique_variables

    def validate_template(self, template_path: Path) -> tuple:
        """
        验证模板文件

        Args:
            template_path: 模板文件路径

        Returns:
            (是否有效, 错误消息)
        """
        try:
            from docxtpl import DocxTemplate
        except ImportError:
            return False, "docxtpl 未安装"

        if not template_path.exists():
            return False, "模板文件不存在"

        if not template_path.suffix.lower() == '.docx':
            return False, "模板必须是 .docx 格式"

        try:
            DocxTemplate(str(template_path))
            return True, ""
        except Exception as e:
            return False, str(e)

    def copy_template(
        self,
        template_path: Path,
        output_path: Path
    ) -> Path:
        """
        直接复制模板文件（不替换变量）

        Args:
            template_path: 模板文件路径
            output_path: 输出文件路径

        Returns:
            输出文件路径
        """
        output_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(template_path, output_path)

        self._logger.info(f"模板复制完成: {output_path}")
        return output_path
