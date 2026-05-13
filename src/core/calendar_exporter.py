# -*- coding: utf-8 -*-
"""期限日历导出服务。"""

from __future__ import annotations

import base64
from collections import defaultdict
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

from src.utils.logger import get_logger


FORMAT_OPTIONS = [
    ("pdf", "PDF 文档"),
    ("docx", "Word 文档"),
    ("html", "HTML 网页"),
    ("md", "Markdown 文档"),
    ("png", "图片 PNG"),
]

THEME_OPTIONS = [
    ("stream", "信息流"),
    ("calendar", "日历卡片"),
    ("compact", "紧凑清单"),
    ("timeline", "时间线"),
    ("board", "看板分栏"),
    ("briefing", "晨会简报"),
    ("agenda", "议程分组"),
]


def _deadline_target_datetime(deadline: Dict[str, Any]) -> Optional[datetime]:
    """构建期限对应时间。"""
    date_text = str(deadline.get("date", "")).strip()
    if not date_text:
        return None
    time_text = str(deadline.get("time", "")).strip()
    try:
        if deadline.get("all_day", not time_text):
            return datetime.strptime(date_text, "%Y-%m-%d")
        return datetime.strptime(f"{date_text} {time_text[:5]}", "%Y-%m-%d %H:%M")
    except ValueError:
        return None


def _deadline_is_completed(deadline: Dict[str, Any]) -> bool:
    return str(deadline.get("status", "pending")).strip() == "completed"


def _deadline_is_overdue(deadline: Dict[str, Any], now: Optional[datetime] = None) -> bool:
    if _deadline_is_completed(deadline):
        return False
    target = _deadline_target_datetime(deadline)
    if target is None:
        return False
    current = now or datetime.now()
    if deadline.get("all_day", True):
        return target.date() < current.date()
    return target < current


def _deadline_days_until(deadline: Dict[str, Any], now: Optional[datetime] = None) -> Optional[int]:
    target = _deadline_target_datetime(deadline)
    if target is None:
        return None
    current = now or datetime.now()
    return (target.date() - current.date()).days


def _normalize_tags(value: Any) -> List[str]:
    if not isinstance(value, (list, tuple, set)):
        return []
    result: List[str] = []
    for item in value:
        text = str(item or "").strip()
        if text:
            result.append(text)
    return result


class CalendarExportService:
    """将筛选后的期限事项导出为多种格式。"""

    def __init__(self) -> None:
        self._logger = get_logger()
        self._logo_path = Path(__file__).resolve().parents[2] / "resources" / "branding" / "lexora_export_logo.png"
        self._logo_data_url: Optional[str] = None

    def export(
        self,
        deadlines: List[Dict[str, Any]],
        output_path: Path,
        *,
        export_format: str,
        theme: str,
        title: str,
        period_label: str,
        summary_label: str,
        filter_notes: Optional[List[str]] = None,
    ) -> Path:
        """导出期限事项。"""
        export_format = str(export_format or "").strip().lower()
        theme = str(theme or "stream").strip().lower()
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        payload = self._build_payload(
            deadlines=deadlines,
            title=title,
            period_label=period_label,
            summary_label=summary_label,
            filter_notes=filter_notes or [],
        )

        if export_format == "html":
            html = self.render_html(payload, theme=theme)
            output_path.write_text(html, encoding="utf-8")
            return output_path

        if export_format == "docx":
            self._write_docx(payload, output_path, theme=theme)
            return output_path

        html = self.render_print_html(payload, theme=theme)
        if export_format == "pdf":
            self._write_pdf_from_html(html, output_path)
            return output_path

        if export_format == "png":
            self._write_png_from_html(html, output_path)
            return output_path

        if export_format == "md":
            self._write_md(payload, output_path, theme=theme)
            return output_path

        raise ValueError(f"不支持的导出格式: {export_format}")

    def render_html(self, payload: Dict[str, Any], *, theme: str = "stream") -> str:
        """渲染导出 HTML。"""
        theme = theme if theme in {key for key, _ in THEME_OPTIONS} else "stream"
        c = self._theme_colors(theme)
        logo_html = ""
        if self._logo_data_url_or_empty():
            logo_html = f'<img class="brand-mark" src="{self._logo_data_url_or_empty()}" alt="LEXORA" width="104">'
        meta_html = "".join(
            f'<span class="meta-chip">{escape(note)}</span>'
            for note in payload["filter_notes"]
        ) or '<span class="meta-chip">按当前筛选结果导出</span>'
        stats_html = "".join(
            f"""
            <div class="stat-card">
                <div class="stat-label">{escape(item['label'])}</div>
                <div class="stat-value">{item['value']}</div>
            </div>
            """
            for item in payload["stats"]
        )

        if theme == "calendar":
            body_html = self._render_calendar_theme(payload)
        elif theme == "compact":
            body_html = self._render_compact_theme(payload)
        elif theme == "timeline":
            body_html = self._render_timeline_theme(payload)
        elif theme == "board":
            body_html = self._render_board_theme(payload)
        elif theme == "briefing":
            body_html = self._render_briefing_theme(payload)
        elif theme == "agenda":
            body_html = self._render_agenda_theme(payload)
        else:
            body_html = self._render_stream_theme(payload)

        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>{escape(payload['title'])}</title>
  <style>
    :root {{
      --bg: {c['bg']};
      --panel: {c['panel']};
      --panel-soft: {c['panel_soft']};
      --line: {c['line']};
      --text: {c['text']};
      --muted: {c['muted']};
      --accent: {c['accent']};
      --accent-soft: {c['accent_soft']};
      --danger: {c['danger']};
      --warning: {c['warning']};
      --success: {c['success']};
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      padding: 32px;
      background: linear-gradient(180deg, #f8fbff 0%, var(--bg) 100%);
      color: var(--text);
      font-family: "PingFang SC", "Microsoft YaHei", "Noto Sans SC", sans-serif;
    }}
    .sheet {{
      max-width: 1180px;
      margin: 0 auto;
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: 24px;
      box-shadow: 0 20px 50px rgba(15, 23, 42, 0.08);
      overflow: hidden;
    }}
    .hero {{
      padding: 28px 30px 18px;
      background:
        radial-gradient(circle at top right, rgba(59,130,246,0.14), transparent 36%),
        linear-gradient(180deg, #ffffff 0%, #f8fbff 100%);
      border-bottom: 1px solid var(--line);
    }}
    .hero-brandbar {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 18px;
      margin-bottom: 18px;
    }}
    .brand-cluster {{
      display: flex;
      align-items: center;
      gap: 16px;
      min-width: 0;
    }}
    .brand-mark {{
      display: block;
      width: 108px;
      height: auto;
      object-fit: contain;
      filter: drop-shadow(0 10px 22px rgba(15, 23, 42, 0.10));
    }}
    .brand-copy {{
      min-width: 0;
    }}
    .brand-label {{
      font-size: 12px;
      letter-spacing: 0.12em;
      font-weight: 800;
      color: var(--accent);
      text-transform: uppercase;
    }}
    .brand-system {{
      margin-top: 6px;
      font-size: 15px;
      font-weight: 800;
      color: var(--text);
      line-height: 1.45;
    }}
    .brand-stamp {{
      flex-shrink: 0;
      text-align: right;
      padding: 10px 12px;
      border-radius: 16px;
      border: 1px solid var(--line);
      background: rgba(255,255,255,0.82);
      box-shadow: 0 10px 22px rgba(15,23,42,0.04);
    }}
    .brand-stamp-label {{
      color: var(--muted);
      font-size: 11px;
      font-weight: 700;
    }}
    .brand-stamp-value {{
      margin-top: 6px;
      color: var(--accent);
      font-size: 18px;
      font-weight: 900;
      letter-spacing: -0.02em;
    }}
    .eyebrow {{
      color: var(--accent);
      font-size: 12px;
      font-weight: 700;
      letter-spacing: 0.08em;
      text-transform: uppercase;
      margin-bottom: 10px;
    }}
    .hero h1 {{
      margin: 0;
      font-size: 30px;
      line-height: 1.2;
    }}
    .hero-sub {{
      margin-top: 8px;
      color: var(--muted);
      font-size: 14px;
      line-height: 1.6;
    }}
    .hero-meta {{
      margin-top: 14px;
      display: flex;
      flex-wrap: wrap;
      gap: 8px;
    }}
    .meta-chip {{
      display: inline-flex;
      align-items: center;
      gap: 6px;
      padding: 6px 12px;
      border-radius: 999px;
      border: 1px solid var(--line);
      background: var(--panel-soft);
      color: var(--muted);
      font-size: 12px;
      font-weight: 600;
    }}
    .stats {{
      display: grid;
      grid-template-columns: repeat(4, minmax(0, 1fr));
      gap: 12px;
      padding: 20px 30px 8px;
    }}
    .stat-card {{
      border: 1px solid var(--line);
      background: var(--panel-soft);
      border-radius: 18px;
      padding: 16px 18px;
    }}
    .stat-label {{
      color: var(--muted);
      font-size: 12px;
      font-weight: 700;
      margin-bottom: 8px;
    }}
    .stat-value {{
      font-size: 28px;
      font-weight: 800;
      letter-spacing: -0.02em;
    }}
    .content {{
      padding: 18px 30px 30px;
    }}
    .empty {{
      border: 1px dashed var(--line);
      border-radius: 18px;
      background: var(--panel-soft);
      text-align: center;
      padding: 30px 18px;
      color: var(--muted);
      font-size: 14px;
      font-weight: 600;
    }}
    .stream-list,
    .timeline-list,
    .calendar-list {{
      display: flex;
      flex-direction: column;
      gap: 14px;
    }}
    .item-shell {{
      display: grid;
      grid-template-columns: 126px minmax(0, 1fr);
      gap: 14px;
      align-items: stretch;
    }}
    .time-panel {{
      border-radius: 20px;
      background: linear-gradient(180deg, var(--accent_soft) 0%, #ffffff 100%);
      border: 1px solid rgba(37, 99, 235, 0.18);
      padding: 16px 14px;
      display: flex;
      flex-direction: column;
      justify-content: center;
      align-items: center;
      text-align: center;
      min-height: 124px;
      box-shadow: inset 0 1px 0 rgba(255,255,255,0.95);
    }}
    .time-date-main {{
      font-size: 30px;
      line-height: 1;
      font-weight: 900;
      letter-spacing: -0.03em;
      color: var(--accent);
    }}
    .time-date-sub {{
      margin-top: 8px;
      color: var(--muted);
      font-size: 12px;
      font-weight: 700;
      line-height: 1.5;
    }}
    .time-clock {{
      margin-top: 10px;
      padding: 6px 12px;
      border-radius: 999px;
      background: #ffffff;
      border: 1px solid rgba(37,99,235,0.16);
      color: var(--text);
      font-size: 16px;
      font-weight: 900;
      letter-spacing: 0.01em;
      box-shadow: 0 8px 18px rgba(37,99,235,0.08);
    }}
    .time-clock.all-day {{
      color: var(--accent);
    }}
    .item-card {{
      border: 1px solid var(--line);
      background: #fff;
      border-radius: 18px;
      padding: 16px 18px;
      box-shadow: 0 10px 24px rgba(15, 23, 42, 0.04);
    }}
    .item-title {{
      margin: 0;
      font-size: 17px;
      font-weight: 800;
      line-height: 1.35;
    }}
    .item-support {{
      margin-top: 8px;
      color: var(--muted);
      font-size: 13px;
      line-height: 1.7;
    }}
    .item-time-highlight {{
      display: inline-flex;
      align-items: center;
      gap: 8px;
      margin-top: 10px;
      padding: 8px 12px;
      border-radius: 12px;
      background: var(--panel-soft);
      border: 1px solid var(--line);
      color: var(--text);
      font-size: 13px;
      font-weight: 800;
    }}
    .badges {{
      display: flex;
      flex-wrap: wrap;
      gap: 8px;
      margin-top: 12px;
    }}
    .badge {{
      display: inline-flex;
      align-items: center;
      padding: 4px 10px;
      border-radius: 999px;
      font-size: 12px;
      font-weight: 700;
      border: 1px solid var(--line);
      background: var(--panel-soft);
      color: var(--text);
    }}
    .badge.danger {{ color: var(--danger); border-color: rgba(220,38,38,0.25); background: rgba(254,226,226,0.9); }}
    .badge.warning {{ color: var(--warning); border-color: rgba(217,119,6,0.24); background: rgba(255,247,237,0.95); }}
    .badge.success {{ color: var(--success); border-color: rgba(22,163,74,0.24); background: rgba(240,253,244,0.95); }}
    .badge.accent {{ color: var(--accent); border-color: rgba(37,99,235,0.24); background: var(--accent_soft); }}
    .desc {{
      margin-top: 12px;
      padding: 12px 14px;
      border-radius: 14px;
      background: var(--panel-soft);
      color: var(--text);
      font-size: 13px;
      line-height: 1.7;
      white-space: pre-wrap;
    }}
    .stream-card .item-card {{
      border-radius: 22px;
      background:
        radial-gradient(circle at top right, rgba(37,99,235,0.06), transparent 28%),
        #ffffff;
    }}
    .theme-stream .time-panel {{
      min-height: 136px;
    }}
    .theme-stream .item-title {{
      font-size: 20px;
    }}
    .theme-stream .item-time-highlight {{
      font-size: 14px;
    }}
    .calendar-grid {{
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(260px, 1fr));
      gap: 14px;
    }}
    .calendar-day-shell {{
      border-radius: 18px;
      border: 1px solid var(--line);
      background:
        linear-gradient(180deg, rgba(239,246,255,0.96) 0%, rgba(255,255,255,1) 100%);
      box-shadow: 0 14px 28px rgba(37,99,235,0.08);
      overflow: hidden;
    }}
    .calendar-day-head {{
      display: grid;
      grid-template-columns: 86px minmax(0, 1fr);
      gap: 14px;
      padding: 16px;
      border-bottom: 1px solid var(--line);
    }}
    .calendar-date-tile {{
      border-radius: 16px;
      background: #ffffff;
      border: 1px solid rgba(37,99,235,0.14);
      padding: 12px 8px;
      text-align: center;
      box-shadow: 0 10px 20px rgba(37,99,235,0.08);
    }}
    .calendar-date-number {{
      font-size: 30px;
      font-weight: 900;
      color: var(--accent);
      line-height: 1;
    }}
    .calendar-date-month {{
      margin-top: 6px;
      color: var(--muted);
      font-size: 11px;
      font-weight: 800;
      letter-spacing: 0.05em;
    }}
    .calendar-date-summary {{
      display: flex;
      flex-direction: column;
      justify-content: center;
      gap: 8px;
    }}
    .calendar-date-weekday {{
      font-size: 16px;
      font-weight: 800;
      color: var(--text);
    }}
    .calendar-date-meta {{
      color: var(--muted);
      font-size: 12px;
      font-weight: 700;
      line-height: 1.5;
    }}
    .calendar-stack {{
      display: flex;
      flex-direction: column;
      gap: 10px;
      padding: 14px 16px 16px;
    }}
    .calendar-entry {{
      border-radius: 14px;
      background: #ffffff;
      border: 1px solid rgba(37,99,235,0.12);
      padding: 12px 14px;
      box-shadow: 0 10px 20px rgba(15, 23, 42, 0.04);
    }}
    .calendar-entry-top {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 8px;
    }}
    .calendar-entry-title {{
      font-size: 14px;
      font-weight: 800;
      line-height: 1.35;
      color: var(--text);
    }}
    .calendar-entry-time {{
      flex-shrink: 0;
      padding: 5px 12px;
      border-radius: 999px;
      background: var(--accent_soft);
      color: var(--accent);
      font-size: 13px;
      font-weight: 900;
    }}
    .calendar-entry-meta {{
      margin-top: 8px;
      color: var(--muted);
      font-size: 12px;
      line-height: 1.6;
    }}
    .compact-table {{
      width: 100%;
      border-collapse: collapse;
      overflow: hidden;
      border-radius: 18px;
      border: 1px solid var(--line);
    }}
    .compact-table th,
    .compact-table td {{
      padding: 12px 14px;
      border-bottom: 1px solid var(--line);
      text-align: left;
      font-size: 13px;
      vertical-align: top;
    }}
    .compact-table th {{
      background: var(--panel-soft);
      color: var(--muted);
      font-size: 12px;
      font-weight: 800;
    }}
    .compact-table tr:last-child td {{
      border-bottom: none;
    }}
    .compact-time {{
      min-width: 132px;
    }}
    .compact-time-date {{
      font-size: 18px;
      font-weight: 900;
      color: var(--text);
      line-height: 1.2;
    }}
    .compact-time-clock {{
      margin-top: 4px;
      display: inline-flex;
      padding: 4px 8px;
      border-radius: 999px;
      background: var(--accent_soft);
      color: var(--accent);
      font-size: 12px;
      font-weight: 800;
    }}
    .timeline {{
      position: relative;
      padding-left: 22px;
    }}
    .timeline::before {{
      content: "";
      position: absolute;
      left: 9px;
      top: 4px;
      bottom: 4px;
      width: 4px;
      border-radius: 999px;
      background: linear-gradient(180deg, rgba(37,99,235,0.25) 0%, rgba(37,99,235,0.05) 100%);
    }}
    .timeline-row {{
      position: relative;
      display: grid;
      grid-template-columns: 146px minmax(0, 1fr);
      gap: 14px;
      margin-bottom: 16px;
    }}
    .timeline-row:last-child {{
      margin-bottom: 0;
    }}
    .timeline-dot {{
      position: absolute;
      left: -22px;
      top: 22px;
      width: 18px;
      height: 18px;
      border-radius: 50%;
      background: #ffffff;
      border: 4px solid var(--accent);
      box-shadow: 0 0 0 6px rgba(37,99,235,0.10);
    }}
    .timeline-date-box {{
      border-radius: 18px;
      padding: 16px 14px;
      background: linear-gradient(180deg, rgba(219,234,254,0.9) 0%, rgba(255,255,255,1) 100%);
      border: 1px solid rgba(37,99,235,0.16);
      text-align: center;
    }}
    .timeline-date-box .time-date-main {{
      font-size: 32px;
    }}
    .timeline-card {{
      border-left: 4px solid var(--accent);
    }}
    .board-grid {{
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(240px, 1fr));
      gap: 14px;
    }}
    .board-column {{
      border: 1px solid var(--line);
      border-radius: 20px;
      background: linear-gradient(180deg, rgba(248,250,252,0.88) 0%, rgba(255,255,255,1) 100%);
      overflow: hidden;
      box-shadow: 0 12px 24px rgba(15,23,42,0.04);
    }}
    .board-column-head {{
      padding: 14px 16px;
      border-bottom: 1px solid var(--line);
      background: var(--panel-soft);
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 8px;
    }}
    .board-column-title {{
      font-size: 14px;
      font-weight: 900;
      color: var(--text);
    }}
    .board-column-count {{
      padding: 4px 8px;
      border-radius: 999px;
      background: #fff;
      border: 1px solid var(--line);
      color: var(--muted);
      font-size: 11px;
      font-weight: 800;
    }}
    .board-column-body {{
      padding: 12px;
      display: flex;
      flex-direction: column;
      gap: 10px;
    }}
    .board-card {{
      border-radius: 16px;
      background: #ffffff;
      border: 1px solid var(--line);
      padding: 12px 12px 14px;
    }}
    .board-time {{
      display: grid;
      grid-template-columns: 56px minmax(0, 1fr);
      gap: 10px;
      margin-bottom: 10px;
      align-items: stretch;
    }}
    .board-time-day {{
      border-radius: 14px;
      background: linear-gradient(180deg, var(--accent) 0%, #1d4ed8 100%);
      color: #ffffff;
      padding: 10px 8px;
      text-align: center;
      box-shadow: 0 10px 20px rgba(37,99,235,0.16);
    }}
    .board-time-day-number {{
      font-size: 24px;
      line-height: 1;
      font-weight: 900;
    }}
    .board-time-day-meta {{
      margin-top: 6px;
      font-size: 11px;
      line-height: 1.4;
      font-weight: 700;
      opacity: 0.95;
    }}
    .board-time-copy {{
      border-radius: 14px;
      background: var(--accent_soft);
      color: var(--accent);
      padding: 10px 12px;
      display: flex;
      flex-direction: column;
      justify-content: center;
      gap: 4px;
      font-size: 12px;
      font-weight: 800;
    }}
    .board-card-title {{
      font-size: 14px;
      font-weight: 800;
      line-height: 1.4;
      color: var(--text);
    }}
    .board-card-meta {{
      margin-top: 8px;
      color: var(--muted);
      font-size: 12px;
      line-height: 1.6;
    }}
    .board-empty {{
      border: 1px dashed var(--line);
      border-radius: 14px;
      padding: 18px 12px;
      text-align: center;
      color: var(--muted);
      font-size: 12px;
      font-weight: 700;
      background: rgba(255,255,255,0.75);
    }}
    .briefing-grid {{
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(320px, 1fr));
      gap: 16px;
    }}
    .briefing-card {{
      border: 1px solid var(--line);
      border-radius: 22px;
      background:
        radial-gradient(circle at top right, rgba(37,99,235,0.10), transparent 26%),
        linear-gradient(180deg, rgba(255,255,255,1) 0%, rgba(248,250,252,0.98) 100%);
      overflow: hidden;
      box-shadow: 0 16px 32px rgba(15,23,42,0.06);
    }}
    .briefing-head {{
      display: flex;
      align-items: stretch;
      gap: 14px;
      padding: 18px 18px 14px;
      border-bottom: 1px solid var(--line);
    }}
    .briefing-time {{
      min-width: 104px;
      border-radius: 18px;
      background: linear-gradient(180deg, var(--accent) 0%, #1d4ed8 100%);
      color: #ffffff;
      padding: 14px 12px;
      text-align: center;
      box-shadow: 0 14px 28px rgba(37,99,235,0.18);
    }}
    .briefing-day {{
      font-size: 30px;
      font-weight: 900;
      line-height: 1;
      letter-spacing: -0.03em;
    }}
    .briefing-datetime {{
      margin-top: 8px;
      font-size: 12px;
      line-height: 1.5;
      font-weight: 700;
      opacity: 0.95;
    }}
    .briefing-title {{
      font-size: 18px;
      line-height: 1.4;
      font-weight: 900;
      color: var(--text);
      margin: 0;
    }}
    .briefing-sub {{
      margin-top: 8px;
      color: var(--muted);
      font-size: 13px;
      line-height: 1.7;
    }}
    .briefing-body {{
      padding: 14px 18px 18px;
    }}
    .briefing-emphasis {{
      display: inline-flex;
      align-items: center;
      gap: 8px;
      padding: 8px 12px;
      border-radius: 12px;
      background: var(--panel_soft);
      border: 1px solid var(--line);
      color: var(--accent);
      font-size: 13px;
      font-weight: 900;
      margin-bottom: 12px;
    }}
    .agenda-groups {{
      display: flex;
      flex-direction: column;
      gap: 16px;
    }}
    .agenda-group {{
      border: 1px solid var(--line);
      border-radius: 20px;
      overflow: hidden;
      background: #ffffff;
      box-shadow: 0 14px 28px rgba(15,23,42,0.05);
    }}
    .agenda-group-head {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 12px;
      padding: 16px 18px;
      background: linear-gradient(180deg, rgba(219,234,254,0.82) 0%, rgba(248,250,252,0.98) 100%);
      border-bottom: 1px solid var(--line);
    }}
    .agenda-group-title {{
      font-size: 18px;
      font-weight: 900;
      color: var(--text);
    }}
    .agenda-group-count {{
      font-size: 12px;
      font-weight: 800;
      color: var(--accent);
      padding: 6px 10px;
      border-radius: 999px;
      background: rgba(255,255,255,0.86);
      border: 1px solid rgba(37,99,235,0.18);
    }}
    .agenda-slots {{
      display: flex;
      flex-direction: column;
    }}
    .agenda-slot {{
      display: grid;
      grid-template-columns: 120px minmax(0, 1fr);
      gap: 0;
      border-bottom: 1px solid var(--line);
    }}
    .agenda-slot:last-child {{
      border-bottom: none;
    }}
    .agenda-slot-time {{
      padding: 18px 14px;
      background: linear-gradient(180deg, rgba(248,250,252,0.95) 0%, rgba(255,255,255,1) 100%);
      border-right: 1px solid var(--line);
      text-align: center;
    }}
    .agenda-slot-clock {{
      font-size: 24px;
      font-weight: 900;
      color: var(--accent);
      line-height: 1.1;
      letter-spacing: -0.02em;
    }}
    .agenda-slot-hint {{
      margin-top: 8px;
      color: var(--muted);
      font-size: 12px;
      font-weight: 700;
      line-height: 1.5;
    }}
    .agenda-slot-body {{
      padding: 16px 18px;
    }}
    .agenda-slot-title {{
      font-size: 16px;
      font-weight: 900;
      line-height: 1.4;
      color: var(--text);
    }}
    .agenda-slot-meta {{
      margin-top: 8px;
      color: var(--muted);
      font-size: 13px;
      line-height: 1.7;
    }}
    .footer {{
      padding: 18px 30px 26px;
      color: var(--muted);
      font-size: 12px;
      border-top: 1px solid var(--line);
      background: linear-gradient(180deg, rgba(248,250,252,0) 0%, rgba(248,250,252,0.9) 100%);
    }}
  </style>
</head>
<body>
    <div class="sheet theme-{theme}">
    <div class="hero">
      <div class="hero-brandbar">
        <div class="brand-cluster">
          {logo_html}
          <div class="brand-copy">
            <div class="brand-label">LEXORA EXPORT</div>
            <div class="brand-system">案件期限导出视图<br>让时间成为案件管理的第一视觉</div>
          </div>
        </div>
        <div class="brand-stamp">
          <div class="brand-stamp-label">导出主题</div>
          <div class="brand-stamp-value">{escape(self._theme_label(theme))}</div>
        </div>
      </div>
      <div class="eyebrow">LEXORA 案件期限导出</div>
      <h1>{escape(payload['title'])}</h1>
      <div class="hero-sub">{escape(payload['summary_label'])} · {escape(payload['period_label'])} · 导出时间 {escape(payload['generated_at'])}</div>
      <div class="hero-meta">{meta_html}</div>
    </div>
    <div class="stats">{stats_html}</div>
    <div class="content">{body_html}</div>
    <div class="footer">本导出内容基于案件文件夹管理系统当前筛选结果自动生成，可用于内部排期、开庭准备和工作汇报。</div>
  </div>
</body>
</html>"""

    def render_print_html(self, payload: Dict[str, Any], *, theme: str = "stream") -> str:
        theme = theme if theme in {key for key, _ in THEME_OPTIONS} else "stream"
        c = self._theme_colors(theme)
        body_renderers = {
            "stream": self._render_print_stream_theme,
            "calendar": self._render_print_calendar_theme,
            "compact": self._render_print_compact_theme,
            "timeline": self._render_print_timeline_theme,
            "board": self._render_print_board_theme,
            "briefing": self._render_print_briefing_theme,
            "agenda": self._render_print_agenda_theme,
        }
        body_html = body_renderers.get(theme, self._render_print_stream_theme)(payload, c)
        notes_html = " / ".join(escape(note) for note in payload["filter_notes"]) or "按当前筛选结果导出"
        stats_html = "".join(
            (
                '<td style="border:1px solid {line}; background:{soft}; padding:8px 10px;">'
                '<div style="font-size:10px; color:{muted}; font-weight:700;">{label}</div>'
                '<div style="font-size:20px; color:{accent}; font-weight:900; margin-top:4px;">{value}</div>'
                '</td>'
            ).format(
                line=c["line"],
                soft=c["panel_soft"],
                muted=c["muted"],
                accent=c["accent"],
                label=escape(item["label"]),
                value=item["value"],
            )
            for item in payload["stats"]
        )
        logo_html = self._print_logo_html(theme)
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <title>{escape(payload['title'])}</title>
</head>
<body style="margin:0; padding:16px; background:{c['bg']}; color:{c['text']}; font-family:'Microsoft YaHei','PingFang SC',sans-serif; font-size:12px;">
  <table cellspacing="0" cellpadding="0" width="100%" style="background:#ffffff; border:1px solid {c['line']}; border-collapse:collapse;">
    <tr>
      <td style="padding:14px 16px 10px; border-bottom:1px solid {c['line']}; background:#f8fbff;">
        <table cellspacing="0" cellpadding="0" width="100%">
          <tr>
            <td valign="top" style="width:100px;">{logo_html}</td>
            <td valign="top" style="padding-left:10px;">
              <div style="font-size:11px; color:{c['accent']}; font-weight:800;">LEXORA EXPORT</div>
              <div style="font-size:20px; font-weight:900; color:{c['text']}; margin-top:4px;">{escape(payload['title'])}</div>
              <div style="font-size:12px; color:{c['muted']}; margin-top:6px;">{escape(payload['summary_label'])} · {escape(payload['period_label'])} · {escape(payload['generated_at'])}</div>
              <div style="font-size:11px; color:{c['muted']}; margin-top:4px;">{notes_html}</div>
            </td>
          </tr>
        </table>
      </td>
    </tr>
    <tr>
      <td style="padding:10px 16px 0;">
        <table cellspacing="6" cellpadding="0" width="100%"><tr>{stats_html}</tr></table>
      </td>
    </tr>
    <tr>
      <td style="padding:12px 16px 16px;">{body_html}</td>
    </tr>
  </table>
</body>
</html>"""

    def _build_payload(
        self,
        *,
        deadlines: Iterable[Dict[str, Any]],
        title: str,
        period_label: str,
        summary_label: str,
        filter_notes: List[str],
    ) -> Dict[str, Any]:
        items = [self._normalize_deadline(item) for item in deadlines]
        items.sort(key=lambda item: (item["date_sort"], item["time_sort"], item["title"]))
        stats = self._build_stats(items)
        return {
            "title": title,
            "period_label": period_label,
            "summary_label": summary_label,
            "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "filter_notes": filter_notes,
            "items": items,
            "stats": stats,
        }

    def _build_stats(self, items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        return [
            {"label": "导出事项", "value": len(items)},
            {"label": "待处理", "value": sum(1 for item in items if item["status_key"] == "pending")},
            {"label": "已逾期", "value": sum(1 for item in items if item["overdue"])},
            {"label": "待开庭", "value": sum(1 for item in items if item["type_key"] == "hearing" and item["status_key"] != "completed")},
        ]

    def _normalize_deadline(self, deadline: Dict[str, Any]) -> Dict[str, Any]:
        target = _deadline_target_datetime(deadline)
        overdue = _deadline_is_overdue(deadline)
        days_until = _deadline_days_until(deadline)
        type_key = str(deadline.get("type", "deadline")).strip() or "deadline"
        status_key = str(deadline.get("status", "pending")).strip() or "pending"
        date_text = str(deadline.get("date", "")).strip() or "未设置日期"
        time_text = "全天" if deadline.get("all_day", True) else (str(deadline.get("time", "")).strip() or "09:00")
        date_sort = target.strftime("%Y-%m-%d") if target else "9999-99-99"
        time_sort = "00:00" if deadline.get("all_day", True) else (str(deadline.get("time", "")).strip() or "99:99")
        case_name = str(deadline.get("case_name", "") or "").strip() or "未关联案件"
        description = str(deadline.get("description", "") or "").strip()
        priority_key = str(deadline.get("priority", "medium")).strip() or "medium"
        weekday_label = self._weekday_label(date_text)
        return {
            "title": str(deadline.get("title", "") or "未命名事项").strip() or "未命名事项",
            "date_label": date_text,
            "date_day": date_text[-2:] if len(date_text) >= 10 else date_text,
            "date_month_label": date_text[5:7] + "月" if len(date_text) >= 7 else date_text,
            "date_brief_label": date_text[5:] if len(date_text) >= 10 else date_text,
            "weekday_label": weekday_label,
            "time_label": time_text,
            "date_sort": date_sort,
            "time_sort": time_sort,
            "type_key": type_key,
            "type_label": self._type_label(type_key),
            "status_key": status_key,
            "status_label": self._status_label(status_key, overdue),
            "priority_key": priority_key,
            "priority_label": self._priority_label(priority_key),
            "case_name": case_name,
            "case_tags": _normalize_tags(deadline.get("case_tags", [])),
            "overdue": overdue,
            "all_day": bool(deadline.get("all_day", True)),
            "description": description,
            "days_hint": self._days_hint(days_until, overdue, status_key),
        }

    def _render_stream_theme(self, payload: Dict[str, Any]) -> str:
        items = payload["items"]
        if not items:
            return '<div class="empty">当前筛选结果下暂无期限事项。</div>'
        cards = [
            f"""
            <div class="item-shell stream-card">
              {self._render_time_panel(item)}
              {self._render_item_card(item, emphasize_time=True, show_absolute_time=False)}
            </div>
            """
            for item in items
        ]
        return '<div class="stream-list">' + "".join(cards) + "</div>"

    def _render_calendar_theme(self, payload: Dict[str, Any]) -> str:
        items = payload["items"]
        if not items:
            return '<div class="empty">当前筛选结果下暂无期限事项。</div>'
        groups: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
        for item in items:
            groups[item["date_label"]].append(item)
        sections: List[str] = []
        for date_label, group_items in groups.items():
            first = group_items[0]
            entries = []
            for item in group_items:
                entries.append(
                    f"""
                    <div class="calendar-entry">
                      <div class="calendar-entry-top">
                        <div class="calendar-entry-title">{escape(item['title'])}</div>
                        <div class="calendar-entry-time">{escape(item['time_label'])}</div>
                      </div>
                      <div class="calendar-entry-meta">
                        {escape(item['case_name'])} · {escape(item['days_hint'])}
                      </div>
                    </div>
                    """
                )
            sections.append(
                f"""
                <div class="calendar-day-shell">
                  <div class="calendar-day-head">
                    <div class="calendar-date-tile">
                      <div class="calendar-date-number">{escape(first['date_day'])}</div>
                      <div class="calendar-date-month">{escape(first['date_month_label'])}</div>
                    </div>
                    <div class="calendar-date-summary">
                      <div class="calendar-date-weekday">{escape(first['weekday_label'])}</div>
                      <div class="calendar-date-meta">{escape(date_label)} · 共 {len(group_items)} 项</div>
                    </div>
                  </div>
                  <div class="calendar-stack">{''.join(entries)}</div>
                </div>
                """
            )
        return '<div class="calendar-grid">' + "".join(sections) + "</div>"

    def _render_compact_theme(self, payload: Dict[str, Any]) -> str:
        items = payload["items"]
        if not items:
            return '<div class="empty">当前筛选结果下暂无期限事项。</div>'
        rows = []
        for item in items:
            tags = " / ".join(item["case_tags"]) or "无标签"
            rows.append(
                f"""
                <tr>
                  <td class="compact-time"><div class="compact-time-date">{escape(item['date_brief_label'])}</div><div class="compact-time-clock">{escape(item['time_label'])}</div><div style="margin-top:4px;color:var(--muted);font-size:11px;">{escape(item['weekday_label'])}</div></td>
                  <td><strong>{escape(item['title'])}</strong><br><span style="color:var(--accent);font-weight:800;">{escape(item['days_hint'])}</span></td>
                  <td>{escape(item['case_name'])}</td>
                  <td>{escape(item['type_label'])}<br><span style="color:var(--muted)">{escape(item['status_label'])}</span></td>
                </tr>
                """
            )
        return f"""
        <table class="compact-table">
          <thead>
            <tr>
              <th>日期时间</th>
              <th>事项</th>
              <th>案件</th>
              <th>类型 / 状态</th>
            </tr>
          </thead>
          <tbody>{''.join(rows)}</tbody>
        </table>
        """

    def _render_timeline_theme(self, payload: Dict[str, Any]) -> str:
        items = payload["items"]
        if not items:
            return '<div class="empty">当前筛选结果下暂无期限事项。</div>'
        rows = []
        for item in items:
            rows.append(
                f"""
                <div class="timeline-row">
                  <div class="timeline-dot"></div>
                  <div class="timeline-date-box">
                    <div class="time-date-main">{escape(item['date_day'])}</div>
                    <div class="time-date-sub">{escape(item['date_month_label'])}<br>{escape(item['weekday_label'])}</div>
                    <div class="time-clock {'all-day' if item['all_day'] else ''}">{escape(item['time_label'])}</div>
                  </div>
                  <div class="item-card timeline-card">
                    <h3 class="item-title">{escape(item['title'])}</h3>
                    <div class="item-time-highlight">{escape(item['date_label'])} · {escape(item['time_label'])} · {escape(item['days_hint'])}</div>
                    <div class="item-support">{escape(item['case_name'])}</div>
                    <div class="badges">{''.join(self._build_badges(item))}</div>
                  </div>
                </div>
                """
            )
        return '<div class="timeline"><div class="timeline-list">' + "".join(rows) + "</div></div>"

    def _render_board_theme(self, payload: Dict[str, Any]) -> str:
        items = payload["items"]
        if not items:
            return '<div class="empty">当前筛选结果下暂无期限事项。</div>'
        groups = [
            ("待处理", [item for item in items if item["status_key"] == "pending" and not item["overdue"] and item["type_key"] != "hearing"]),
            ("开庭", [item for item in items if item["status_key"] != "completed" and item["type_key"] == "hearing"]),
            ("已逾期", [item for item in items if item["overdue"]]),
            ("已完成", [item for item in items if item["status_key"] == "completed"]),
        ]
        columns = []
        for title, group_items in groups:
            if group_items:
                body = "".join(
                    f"""
                    <div class="board-card">
                      <div class="board-time">
                        <div class="board-time-day">
                          <div class="board-time-day-number">{escape(item['date_day'])}</div>
                          <div class="board-time-day-meta">{escape(item['date_month_label'])}<br>{escape(item['weekday_label'])}</div>
                        </div>
                        <div class="board-time-copy">
                          <div>{escape(item['time_label'])}</div>
                          <div>{escape(item['days_hint'])}</div>
                        </div>
                      </div>
                      <div class="board-card-title">{escape(item['title'])}</div>
                      <div class="board-card-meta">{escape(item['case_name'])} · {escape(item['days_hint'])}</div>
                    </div>
                    """
                    for item in group_items
                )
            else:
                body = '<div class="board-empty">当前分栏暂无事项</div>'
            columns.append(
                f"""
                <div class="board-column">
                  <div class="board-column-head">
                    <div class="board-column-title">{escape(title)}</div>
                    <div class="board-column-count">{len(group_items)} 项</div>
                  </div>
                  <div class="board-column-body">{body}</div>
                </div>
                """
            )
        return '<div class="board-grid">' + "".join(columns) + "</div>"

    def _render_briefing_theme(self, payload: Dict[str, Any]) -> str:
        items = payload["items"]
        if not items:
            return '<div class="empty">当前筛选结果下暂无期限事项。</div>'
        cards = []
        for item in items:
            cards.append(
                f"""
                <div class="briefing-card">
                  <div class="briefing-head">
                    <div class="briefing-time">
                      <div class="briefing-day">{escape(item['date_day'])}</div>
                      <div class="briefing-datetime">{escape(item['date_month_label'])} {escape(item['weekday_label'])}<br>{escape(item['time_label'])}</div>
                    </div>
                    <div>
                      <h3 class="briefing-title">{escape(item['title'])}</h3>
                      <div class="briefing-sub">{escape(item['case_name'])} · {escape(item['type_label'])}</div>
                    </div>
                  </div>
                  <div class="briefing-body">
                    <div class="briefing-emphasis">{escape(item['date_label'])} · {escape(item['days_hint'])}</div>
                    <div class="badges">{''.join(self._build_badges(item))}</div>
                  </div>
                </div>
                """
            )
        return '<div class="briefing-grid">' + "".join(cards) + "</div>"

    def _render_agenda_theme(self, payload: Dict[str, Any]) -> str:
        items = payload["items"]
        if not items:
            return '<div class="empty">当前筛选结果下暂无期限事项。</div>'
        groups = self._group_items_by_date(items)
        sections = []
        for _, group_items in groups:
            first = group_items[0]
            rows = []
            for item in group_items:
                rows.append(
                    f"""
                    <div class="agenda-slot">
                      <div class="agenda-slot-time">
                        <div class="agenda-slot-clock">{escape(item['time_label'])}</div>
                        <div class="agenda-slot-hint">{escape(item['days_hint'])}</div>
                      </div>
                      <div class="agenda-slot-body">
                        <div class="agenda-slot-title">{escape(item['title'])}</div>
                        <div class="agenda-slot-meta">{escape(item['case_name'])} · {escape(item['days_hint'])}</div>
                        <div class="badges">{''.join(self._build_badges(item))}</div>
                      </div>
                    </div>
                    """
                )
            sections.append(
                f"""
                <div class="agenda-group">
                  <div class="agenda-group-head">
                    <div class="agenda-group-title">{escape(first['date_label'])} · {escape(first['weekday_label'])}</div>
                    <div class="agenda-group-count">{len(group_items)} 项</div>
                  </div>
                  <div class="agenda-slots">{''.join(rows)}</div>
                </div>
                """
            )
        return '<div class="agenda-groups">' + "".join(sections) + "</div>"

    def _render_time_panel(self, item: Dict[str, Any]) -> str:
        return f"""
        <div class="time-panel">
          <div class="time-date-main">{escape(item['date_day'])}</div>
          <div class="time-date-sub">{escape(item['date_month_label'])}<br>{escape(item['weekday_label'])}</div>
          <div class="time-clock {'all-day' if item['all_day'] else ''}">{escape(item['time_label'])}</div>
        </div>
        """

    def _render_item_card(
        self,
        item: Dict[str, Any],
        *,
        emphasize_time: bool = True,
        show_absolute_time: bool = True,
    ) -> str:
        time_parts = []
        if show_absolute_time:
            time_parts.append(item["date_label"])
            time_parts.append(item["time_label"])
        time_parts.append(item["days_hint"])
        time_html = (
            f'<div class="item-time-highlight">{" · ".join(escape(part) for part in time_parts if part)}</div>'
            if emphasize_time
            else ""
        )
        return f"""
        <div class="item-card">
          <h3 class="item-title">{escape(item['title'])}</h3>
          {time_html}
          <div class="item-support">
            {escape(item['case_name'])}<br>
            {escape(item['type_label'])} / {escape(item['status_label'])}
          </div>
          <div class="badges">{''.join(self._build_badges(item))}</div>
        </div>
        """

    def _build_badges(self, item: Dict[str, Any]) -> List[str]:
        return [
            self._render_badge(item["type_label"], self._badge_tone(item["type_key"])),
            self._render_badge(item["priority_label"], "warning" if item["priority_key"] == "high" else "accent"),
            self._render_badge(item["status_label"], "danger" if item["overdue"] else ("success" if item["status_key"] == "completed" else "accent")),
        ]

    def _description_html(self, item: Dict[str, Any]) -> str:
        if not item["description"]:
            return ""
        return f'<div class="desc">{escape(item["description"])}</div>'

    def _render_badge(self, text: str, tone: str) -> str:
        return f'<span class="badge {tone}">{escape(text)}</span>'

    def _print_logo_html(self, theme: str) -> str:
        data_url = self._logo_data_url_or_empty()
        if not data_url:
            return (
                f'<div style="font-size:20px; font-weight:900; color:{self._theme_colors(theme)["accent"]};">'
                "LEXORA</div>"
            )
        return f'<img src="{data_url}" alt="LEXORA" width="92" style="width:92px; height:auto; display:block;">'

    def _print_badges_html(self, item: Dict[str, Any], c: Dict[str, str]) -> str:
        tone_map = {
            "danger": ("#fee2e2", c["danger"]),
            "warning": ("#ffedd5", c["warning"]),
            "success": ("#dcfce7", c["success"]),
            "accent": (c["accent_soft"], c["accent"]),
        }
        chunks = []
        for tone_text, tone in (
            (item["type_label"], self._badge_tone(item["type_key"])),
            (item["priority_label"], "warning" if item["priority_key"] == "high" else "accent"),
            (item["status_label"], "danger" if item["overdue"] else ("success" if item["status_key"] == "completed" else "accent")),
        ):
            bg, fg = tone_map[tone]
            chunks.append(
                f'<span style="display:inline-block; margin:0 4px 4px 0; padding:3px 6px; background:{bg}; color:{fg}; '
                'font-size:10px; font-weight:700; border:1px solid transparent;">'
                f'{escape(tone_text)}</span>'
            )
        return "".join(chunks)

    def _print_time_box_html(
        self,
        item: Dict[str, Any],
        c: Dict[str, str],
        *,
        strong: bool = False,
        compact: bool = False,
    ) -> str:
        accent = c["danger"] if item["type_key"] == "hearing" else c["accent"]
        soft = "#fee2e2" if item["type_key"] == "hearing" else c["accent_soft"]
        day_size = "26px" if strong else "22px"
        time_size = "16px" if strong else "14px"
        padding = "12px 10px" if not compact else "10px 8px"
        return (
            f'<table cellspacing="0" cellpadding="0" width="100%" style="border:1px solid {c["line"]}; background:{soft};">'
            f'<tr><td align="center" style="padding:{padding};">'
            f'<div style="font-size:{day_size}; line-height:1; font-weight:900; color:{accent};">{escape(item["date_day"])}</div>'
            f'<div style="margin-top:5px; font-size:10px; font-weight:700; color:{c["muted"]};">{escape(item["date_month_label"])} {escape(item["weekday_label"])}</div>'
            f'<div style="margin-top:8px; font-size:{time_size}; font-weight:900; color:{accent};">{escape(item["time_label"])}</div>'
            f'<div style="margin-top:6px; font-size:10px; font-weight:700; color:{c["text"]};">{escape(item["days_hint"])}</div>'
            '</td></tr></table>'
        )

    def _render_print_stream_theme(self, payload: Dict[str, Any], c: Dict[str, str]) -> str:
        if not payload["items"]:
            return self._print_empty_html(c)
        cards = []
        for item in payload["items"]:
            cards.append(
                f'''
                <table cellspacing="0" cellpadding="0" width="100%" style="margin-bottom:12px; border:1px solid {c["line"]}; background:#ffffff;">
                  <tr>
                    <td valign="top" width="118" style="padding:12px;">{self._print_time_box_html(item, c, strong=True)}</td>
                    <td valign="top" style="padding:14px 16px;">
                      <div style="font-size:17px; line-height:1.35; font-weight:900; color:{c["text"]};">{escape(item["title"])}</div>
                      <div style="margin-top:6px; font-size:12px; line-height:1.6; color:{c["accent"]}; font-weight:800;">{escape(item["date_label"])} · {escape(item["time_label"])} · {escape(item["days_hint"])}</div>
                      <div style="margin-top:5px; font-size:11px; color:{c["muted"]};">{escape(item["case_name"])}</div>
                      <div style="margin-top:6px;">{self._print_badges_html(item, c)}</div>
                    </td>
                  </tr>
                </table>
                '''
            )
        return "".join(cards)

    def _render_print_calendar_theme(self, payload: Dict[str, Any], c: Dict[str, str]) -> str:
        groups = self._group_items_by_date(payload["items"])
        if not groups:
            return self._print_empty_html(c)
        sections = []
        for _date_label, group_items in groups:
            first = group_items[0]
            rows = []
            for item in group_items:
                rows.append(
                    f'''
                    <tr>
                      <td valign="top" width="110" style="padding:10px 12px; border-top:1px solid {c["line"]}; color:{c["accent"]}; font-size:15px; font-weight:900;">{escape(item["time_label"])}</td>
                      <td valign="top" style="padding:10px 12px; border-top:1px solid {c["line"]};">
                        <div style="font-size:14px; font-weight:900; color:{c["text"]};">{escape(item["title"])}</div>
                        <div style="margin-top:4px; font-size:11px; color:{c["muted"]};">{escape(item["case_name"])} · {escape(item["days_hint"])}</div>
                        <div style="margin-top:5px;">{self._print_badges_html(item, c)}</div>
                      </td>
                    </tr>
                    '''
                )
            sections.append(
                f'''
                <table cellspacing="0" cellpadding="0" width="100%" style="margin-bottom:14px; border:1px solid {c["line"]}; background:#ffffff;">
                  <tr>
                    <td valign="top" width="92" style="padding:12px;">{self._print_time_box_html(first, c, strong=True, compact=True)}</td>
                    <td valign="middle" style="padding:12px 14px; background:{c["panel_soft"]}; border-left:1px solid {c["line"]};">
                      <div style="font-size:18px; font-weight:900; color:{c["text"]};">{escape(first["date_label"])}</div>
                      <div style="margin-top:6px; font-size:12px; line-height:1.7; color:{c["muted"]};">{escape(first["weekday_label"])} · 共 {len(group_items)} 项</div>
                    </td>
                  </tr>
                  {''.join(rows)}
                </table>
                '''
            )
        return "".join(sections)

    def _render_print_compact_theme(self, payload: Dict[str, Any], c: Dict[str, str]) -> str:
        if not payload["items"]:
            return self._print_empty_html(c)
        rows = []
        for item in payload["items"]:
            rows.append(
                f'''
                <tr>
                  <td valign="top" style="padding:8px 10px; border-top:1px solid {c["line"]}; background:{c["panel_soft"]};">
                    <div style="font-size:14px; font-weight:900; color:{c["text"]};">{escape(item["date_brief_label"])}</div>
                    <div style="margin-top:3px; font-size:13px; font-weight:900; color:{c["accent"]};">{escape(item["time_label"])}</div>
                    <div style="margin-top:2px; font-size:10px; color:{c["muted"]};">{escape(item["days_hint"])}</div>
                  </td>
                  <td valign="top" style="padding:8px 10px; border-top:1px solid {c["line"]};">
                    <div style="font-size:12px; font-weight:900; color:{c["text"]};">{escape(item["title"])}</div>
                    <div style="margin-top:3px; font-size:10px; color:{c["muted"]};">{escape(item["case_name"])}</div>
                  </td>
                  <td valign="top" style="padding:8px 10px; border-top:1px solid {c["line"]}; font-size:10px; color:{c["text"]};">{escape(item["type_label"])}<br/>{escape(item["status_label"])}</td>
                </tr>
                '''
            )
        return (
            f'<table cellspacing="0" cellpadding="0" width="100%" style="border:1px solid {c["line"]}; background:#ffffff;">'
            f'<tr style="background:{c["panel_soft"]};">'
            f'<td style="padding:8px 10px; font-size:11px; font-weight:800; color:{c["muted"]};">时间</td>'
            f'<td style="padding:8px 10px; font-size:11px; font-weight:800; color:{c["muted"]};">事项 / 案件</td>'
            f'<td style="padding:8px 10px; font-size:11px; font-weight:800; color:{c["muted"]};">类型 / 状态</td></tr>'
            + "".join(rows)
            + "</table>"
        )

    def _render_print_timeline_theme(self, payload: Dict[str, Any], c: Dict[str, str]) -> str:
        if not payload["items"]:
            return self._print_empty_html(c)
        blocks = []
        for item in payload["items"]:
            blocks.append(
                f'''
                <table cellspacing="0" cellpadding="0" width="100%" style="margin-bottom:12px;">
                  <tr>
                    <td valign="top" width="122" style="padding-right:10px;">{self._print_time_box_html(item, c, strong=True)}</td>
                    <td valign="top" style="border-left:4px solid {c["accent"]}; padding:12px 14px; background:#ffffff; border-top:1px solid {c["line"]}; border-right:1px solid {c["line"]}; border-bottom:1px solid {c["line"]};">
                      <div style="font-size:16px; font-weight:900; color:{c["text"]};">{escape(item["title"])}</div>
                      <div style="margin-top:5px; font-size:12px; color:{c["accent"]}; font-weight:800;">{escape(item["date_label"])} · {escape(item["time_label"])} · {escape(item["days_hint"])}</div>
                      <div style="margin-top:4px; font-size:11px; color:{c["muted"]};">{escape(item["case_name"])}</div>
                      <div style="margin-top:5px;">{self._print_badges_html(item, c)}</div>
                    </td>
                  </tr>
                </table>
                '''
            )
        return "".join(blocks)

    def _render_print_board_theme(self, payload: Dict[str, Any], c: Dict[str, str]) -> str:
        items = payload["items"]
        if not items:
            return self._print_empty_html(c)
        groups = [
            ("待处理", [item for item in items if item["status_key"] == "pending" and not item["overdue"] and item["type_key"] != "hearing"]),
            ("开庭", [item for item in items if item["status_key"] != "completed" and item["type_key"] == "hearing"]),
            ("已逾期", [item for item in items if item["overdue"]]),
            ("已完成", [item for item in items if item["status_key"] == "completed"]),
        ]
        rendered = []
        for title, group_items in groups:
            cards = []
            if group_items:
                for item in group_items:
                    cards.append(
                        f'''
                        <table cellspacing="0" cellpadding="0" width="100%" style="margin-bottom:8px; border:1px solid {c["line"]}; background:#ffffff;">
                          <tr>
                            <td valign="top" width="86" style="padding:8px;">{self._print_time_box_html(item, c, compact=True)}</td>
                            <td valign="top" style="padding:10px 12px;">
                              <div style="font-size:12px; font-weight:900; color:{c["text"]};">{escape(item["title"])}</div>
                              <div style="margin-top:3px; font-size:10px; color:{c["muted"]};">{escape(item["case_name"])} · {escape(item["days_hint"])}</div>
                            </td>
                          </tr>
                        </table>
                        '''
                    )
            else:
                cards.append(f'<div style="font-size:11px; color:{c["muted"]}; padding:10px 0;">当前分栏暂无事项</div>')
            rendered.append(
                f'''
                <td valign="top" width="50%" style="padding:0 6px 12px 0;">
                  <table cellspacing="0" cellpadding="0" width="100%" style="border:1px solid {c["line"]}; background:{c["panel_soft"]};">
                    <tr><td style="padding:10px 12px; font-size:13px; font-weight:900; color:{c["text"]};">{escape(title)}（{len(group_items)}）</td></tr>
                    <tr><td style="padding:10px 10px 2px; background:#ffffff;">{''.join(cards)}</td></tr>
                  </table>
                </td>
                '''
            )
        rows = []
        for index in range(0, len(rendered), 2):
            left = rendered[index]
            right = rendered[index + 1] if index + 1 < len(rendered) else '<td width="50%"></td>'
            rows.append(f'<tr>{left}{right}</tr>')
        return f'<table cellspacing="0" cellpadding="0" width="100%">{"".join(rows)}</table>'

    def _render_print_briefing_theme(self, payload: Dict[str, Any], c: Dict[str, str]) -> str:
        if not payload["items"]:
            return self._print_empty_html(c)
        chunks = []
        for item in payload["items"]:
            chunks.append(
                f'''
                <table cellspacing="0" cellpadding="0" width="100%" style="margin-bottom:14px; border:1px solid {c["line"]}; background:#ffffff;">
                  <tr>
                    <td valign="top" width="132" style="padding:12px; background:{c["panel_soft"]};">{self._print_time_box_html(item, c, strong=True)}</td>
                    <td valign="top" style="padding:14px 16px;">
                      <div style="font-size:17px; font-weight:900; line-height:1.3; color:{c["text"]};">{escape(item["title"])}</div>
                      <div style="margin-top:5px; font-size:12px; color:{c["accent"]}; font-weight:800;">{escape(item["time_label"])} · {escape(item["days_hint"])}</div>
                      <div style="margin-top:4px; font-size:11px; color:{c["muted"]};">{escape(item["case_name"])}</div>
                      <div style="margin-top:5px;">{self._print_badges_html(item, c)}</div>
                    </td>
                  </tr>
                </table>
                '''
            )
        return "".join(chunks)

    def _render_print_agenda_theme(self, payload: Dict[str, Any], c: Dict[str, str]) -> str:
        groups = self._group_items_by_date(payload["items"])
        if not groups:
            return self._print_empty_html(c)
        sections = []
        for _date_label, group_items in groups:
            first = group_items[0]
            rows = []
            for item in group_items:
                rows.append(
                    f'''
                    <tr>
                      <td valign="top" width="118" style="padding:10px 12px; border-top:1px solid {c["line"]}; background:{c["panel_soft"]};">
                        <div style="font-size:18px; font-weight:900; color:{c["accent"]};">{escape(item["time_label"])}</div>
                        <div style="margin-top:5px; font-size:10px; font-weight:700; color:{c["muted"]};">{escape(item["days_hint"])}</div>
                      </td>
                      <td valign="top" style="padding:10px 12px; border-top:1px solid {c["line"]};">
                        <div style="font-size:14px; font-weight:900; color:{c["text"]};">{escape(item["title"])}</div>
                        <div style="margin-top:4px; font-size:11px; color:{c["muted"]};">{escape(item["case_name"])} · {escape(item["days_hint"])}</div>
                        <div style="margin-top:5px;">{self._print_badges_html(item, c)}</div>
                      </td>
                    </tr>
                    '''
                )
            sections.append(
                f'''
                <table cellspacing="0" cellpadding="0" width="100%" style="margin-bottom:14px; border:1px solid {c["line"]}; background:#ffffff;">
                  <tr>
                    <td width="106" style="padding:10px 12px;">{self._print_time_box_html(first, c, strong=True)}</td>
                    <td valign="middle" style="padding:12px 14px; background:{c["panel_soft"]}; border-left:1px solid {c["line"]};">
                      <div style="font-size:18px; font-weight:900; color:{c["text"]};">{escape(first["date_label"])}</div>
                      <div style="margin-top:5px; font-size:12px; color:{c["muted"]}; font-weight:700;">{escape(first["weekday_label"])} · 议程 {len(group_items)} 项</div>
                    </td>
                  </tr>
                  {''.join(rows)}
                </table>
                '''
            )
        return "".join(sections)

    def _print_empty_html(self, c: Dict[str, str]) -> str:
        return (
            f'<div style="padding:24px; border:1px dashed {c["line"]}; background:{c["panel_soft"]}; '
            f'font-size:12px; color:{c["muted"]}; text-align:center;">当前筛选结果下暂无期限事项。</div>'
        )

    def _group_items_by_date(self, items: List[Dict[str, Any]]) -> List[Tuple[str, List[Dict[str, Any]]]]:
        groups: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
        for item in items:
            groups[item["date_label"]].append(item)
        return list(groups.items())

    def _logo_data_url_or_empty(self) -> str:
        if self._logo_data_url is not None:
            return self._logo_data_url
        if not self._logo_path.exists():
            self._logo_data_url = ""
            return self._logo_data_url
        encoded = base64.b64encode(self._logo_path.read_bytes()).decode("ascii")
        self._logo_data_url = f"data:image/png;base64,{encoded}"
        return self._logo_data_url

    def _theme_label(self, theme: str) -> str:
        mapping = dict(THEME_OPTIONS)
        return mapping.get(theme, "信息流")

    def _write_docx(self, payload: Dict[str, Any], output_path: Path, *, theme: str) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt

        document = Document()
        style = document.styles["Normal"]
        style.font.name = "Microsoft YaHei"
        r_pr = style._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)

        if self._logo_path.exists():
            brand = document.add_paragraph()
            brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
            brand_run = brand.add_run()
            brand_run.add_picture(str(self._logo_path), width=Inches(1.05))

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(payload["title"])
        run.bold = True
        run.font.size = Pt(18)

        meta = document.add_paragraph()
        meta.add_run(f"{payload['summary_label']} · {payload['period_label']}\n").bold = True
        meta.add_run(f"导出时间：{payload['generated_at']}")

        if payload["filter_notes"]:
            tags_paragraph = document.add_paragraph()
            tags_paragraph.add_run("筛选条件：").bold = True
            tags_paragraph.add_run(" / ".join(payload["filter_notes"]))

        stats_table = document.add_table(rows=1, cols=len(payload["stats"]))
        stats_table.style = "Table Grid"
        for idx, stat in enumerate(payload["stats"]):
            cell = stats_table.rows[0].cells[idx]
            cell.text = f"{stat['label']}\n{stat['value']}"

        document.add_paragraph("")
        if theme == "compact":
            table = document.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            headers = ["事项", "日期时间", "案件", "类型 / 状态"]
            for idx, header in enumerate(headers):
                table.rows[0].cells[idx].text = header
            for item in payload["items"]:
                row = table.add_row().cells
                row[0].text = item["title"]
                row[1].text = f"{item['date_label']} {item['time_label']}\n{item['days_hint']}"
                row[2].text = item["case_name"]
                row[3].text = f"{item['type_label']} / {item['status_label']}"
        elif theme == "board":
            groups = [
                ("待处理", [item for item in payload["items"] if item["status_key"] == "pending" and not item["overdue"] and item["type_key"] != "hearing"]),
                ("开庭", [item for item in payload["items"] if item["status_key"] != "completed" and item["type_key"] == "hearing"]),
                ("已逾期", [item for item in payload["items"] if item["overdue"]]),
                ("已完成", [item for item in payload["items"] if item["status_key"] == "completed"]),
            ]
            for title_text, items in groups:
                if not items:
                    continue
                heading = document.add_paragraph()
                heading.add_run(f"{title_text}（{len(items)}项）").bold = True
                for item in items:
                    paragraph = document.add_paragraph(style=None)
                    paragraph.paragraph_format.space_after = Pt(6)
                    time_run = paragraph.add_run(f"【{item['date_label']} {item['time_label']}】 ")
                    time_run.bold = True
                    time_run.font.size = Pt(11.5)
                    title_run = paragraph.add_run(item["title"])
                    title_run.bold = True
                    paragraph.add_run(f"\n{item['case_name']} · {item['days_hint']}")
        elif theme in {"calendar", "agenda"}:
            current_group = None
            for item in payload["items"]:
                if item["date_label"] != current_group:
                    current_group = item["date_label"]
                    heading = document.add_paragraph()
                    head_run = heading.add_run(f"{item['date_label']} {self._weekday_label(item['date_label'])}")
                    head_run.bold = True
                    head_run.font.size = Pt(13)
                paragraph = document.add_paragraph(style=None)
                paragraph.paragraph_format.space_after = Pt(6)
                time_run = paragraph.add_run(f"{item['time_label']}  ")
                time_run.bold = True
                time_run.font.size = Pt(12)
                title_run = paragraph.add_run(item["title"])
                title_run.bold = True
                paragraph.add_run(f"\n{item['case_name']} · {item['days_hint']}")
        else:
            for item in payload["items"]:
                paragraph = document.add_paragraph(style=None)
                paragraph.paragraph_format.space_after = Pt(6)
                time_run = paragraph.add_run(f"{item['date_label']} {item['time_label']}\n")
                time_run.bold = True
                time_run.font.size = Pt(12)
                first = paragraph.add_run(f"{item['title']}  ")
                first.bold = True
                paragraph.add_run(f"{item['case_name']}\n")
                paragraph.add_run(f"{item['type_label']} / {item['status_label']} · {item['days_hint']}")

        document.add_paragraph("")
        footer = document.add_paragraph()
        footer.add_run("本导出内容基于案件文件夹管理系统当前筛选结果生成。").italic = True
        document.save(str(output_path))

    def _write_pdf_from_html(self, html: str, output_path: Path) -> None:
        from PySide6.QtGui import QPageLayout, QPageSize, QTextDocument
        from PySide6.QtCore import QMarginsF
        from PySide6.QtPrintSupport import QPrinter

        document = QTextDocument()
        document.setHtml(html)

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(str(output_path))
        printer.setPageSize(QPageSize(QPageSize.PageSizeId.A4))
        printer.setPageMargins(QMarginsF(20, 20, 20, 20), QPageLayout.Unit.Millimeter)
        # 关键：QTextDocument 中的 px 单位按逻辑 DPI 换算；默认 HighResolution=1200 DPI
        # 会导致 18px 字体仅有 ~1.3pt，肉眼不可读。固定 96 DPI 使 px≈pt，确保可读。
        printer.setResolution(96)
        document.print_(printer)

    def _write_md(self, payload: Dict[str, Any], output_path: Path, *, theme: str = "stream") -> None:
        """将期限事项导出为 Markdown 文档。"""
        lines: List[str] = []

        # 标题
        lines.append(f"# {payload.get('title', '期限事项导出')}")
        lines.append("")

        # 周期与摘要
        lines.append(f"> **{payload.get('period_label', '')}** ｜ {payload.get('summary_label', '')}")
        lines.append("")

        # 筛选条件
        filter_notes = payload.get("filter_notes", [])
        if filter_notes:
            lines.append(f"📋 筛选条件：{', '.join(filter_notes)}")
            lines.append("")

        # 统计
        stats = payload.get("stats", [])
        if stats:
            lines.append("## 📊 统计概览")
            lines.append("")
            for stat in stats:
                lines.append(f"- **{stat['label']}**：{stat['value']}")
            lines.append("")

        # 事项列表
        items = payload.get("items", [])
        if not items:
            lines.append("> 当前筛选结果下暂无期限事项。")
            lines.append("")
        else:
            lines.append("## 📅 期限事项")
            lines.append("")

            # 按日期分组
            from itertools import groupby
            sorted_items = sorted(items, key=lambda x: x.get("date_label", ""))
            for date_label, group in groupby(sorted_items, key=lambda x: x.get("date_label", "")):
                lines.append(f"### {date_label}")
                lines.append("")
                for item in group:
                    title = item.get("title", "无标题")
                    time_label = item.get("time_label", "")
                    weekday_label = item.get("weekday_label", "")
                    type_label = item.get("type_label", "")
                    status_label = item.get("status_label", "")
                    priority_label = item.get("priority_label", "")
                    case_name = item.get("case_name", "")
                    days_hint = item.get("days_hint", "")
                    description = item.get("description", "")

                    # 状态 emoji
                    status_emoji = "✅" if "已完成" in status_label else "⚠️" if "逾期" in status_label else "⏳"

                    # 头部行
                    header_parts = [status_emoji, f"**{title}**"]
                    if time_label:
                        header_parts.append(f"`{time_label}`")
                    if weekday_label:
                        header_parts.append(f"*{weekday_label}*")
                    lines.append(" | ".join(header_parts))
                    lines.append("")

                    # 属性行
                    attrs = []
                    if type_label:
                        attrs.append(f"类型：{type_label}")
                    if status_label:
                        attrs.append(f"状态：{status_label}")
                    if priority_label:
                        attrs.append(f"优先级：{priority_label}")
                    if case_name:
                        attrs.append(f"案件：{case_name}")
                    if days_hint:
                        attrs.append(f"提醒：{days_hint}")
                    if attrs:
                        lines.append(" · ".join(attrs))
                        lines.append("")

                    # 描述
                    if description and description.strip():
                        lines.append(f"> {description.strip()}")
                        lines.append("")

                lines.append("---")
                lines.append("")

        # 页脚
        lines.append("")
        lines.append(f"---")
        lines.append(f"*由 案件文件夹管理系统 生成于 {payload.get('generated_at', '')}*")
        lines.append("")

        output_path.write_text("\n".join(lines), encoding="utf-8")

    def _write_png_from_html(self, html: str, output_path: Path) -> None:
        from PySide6.QtCore import QSizeF
        from PySide6.QtGui import QColor, QImage, QPainter, QTextDocument

        document = QTextDocument()
        document.setHtml(html)
        document.setPageSize(QSizeF(1000, -1))
        document.setTextWidth(1000)
        size = document.size().toSize()
        width = max(size.width(), 1000)
        height = max(size.height() + 40, 600)

        image = QImage(width, height, QImage.Format.Format_ARGB32_Premultiplied)
        image.fill(QColor("#f8fafc"))
        painter = QPainter(image)
        painter.translate(0, 0)
        document.drawContents(painter)
        painter.end()
        image.save(str(output_path))

    def _type_label(self, value: str) -> str:
        mapping = {
            "deadline": "普通期限",
            "hearing": "开庭事项",
            "custom": "自定义提醒",
        }
        return mapping.get(value, "普通期限")

    def _status_label(self, value: str, overdue: bool) -> str:
        if value == "completed":
            return "已完成"
        if overdue:
            return "已逾期"
        return "待处理"

    def _priority_label(self, value: str) -> str:
        return {
            "high": "高优先",
            "medium": "中优先",
            "low": "低优先",
        }.get(value, "中优先")

    def _days_hint(self, days_until: Optional[int], overdue: bool, status_key: str) -> str:
        if status_key == "completed":
            return "事项已完成"
        if overdue:
            return "已逾期，建议立即处理"
        if days_until is None:
            return "日期待确认"
        if days_until == 0:
            return "今天到期"
        if days_until == 1:
            return "明天到期"
        return f"D-{days_until}" if days_until > 1 else "日期待确认"

    def _weekday_label(self, date_text: str) -> str:
        try:
            value = datetime.strptime(date_text, "%Y-%m-%d").date()
        except ValueError:
            return ""
        return ["周一", "周二", "周三", "周四", "周五", "周六", "周日"][value.weekday()]

    def _badge_tone(self, type_key: str) -> str:
        if type_key == "hearing":
            return "danger"
        if type_key == "custom":
            return "warning"
        return "accent"

    def _theme_colors(self, theme: str) -> Dict[str, str]:
        if theme == "calendar":
            return {
                "bg": "#f3f7fd",
                "panel": "#ffffff",
                "panel_soft": "#f8fbff",
                "line": "#d9e6f7",
                "text": "#10233f",
                "muted": "#64748b",
                "accent": "#2563eb",
                "accent_soft": "#dbeafe",
                "danger": "#dc2626",
                "warning": "#d97706",
                "success": "#16a34a",
            }
        if theme == "timeline":
            return {
                "bg": "#f7fbff",
                "panel": "#ffffff",
                "panel_soft": "#edf5ff",
                "line": "#d7e4f8",
                "text": "#0f1f39",
                "muted": "#64748b",
                "accent": "#0f62fe",
                "accent_soft": "#dbeafe",
                "danger": "#dc2626",
                "warning": "#d97706",
                "success": "#15803d",
            }
        if theme == "board":
            return {
                "bg": "#f8fafc",
                "panel": "#ffffff",
                "panel_soft": "#f1f5f9",
                "line": "#d9e2ec",
                "text": "#10233f",
                "muted": "#5f7087",
                "accent": "#1d4ed8",
                "accent_soft": "#dbeafe",
                "danger": "#dc2626",
                "warning": "#b45309",
                "success": "#15803d",
            }
        if theme == "briefing":
            return {
                "bg": "#f8fbff",
                "panel": "#ffffff",
                "panel_soft": "#eef4ff",
                "line": "#d8e4fb",
                "text": "#10233f",
                "muted": "#64748b",
                "accent": "#2563eb",
                "accent_soft": "#dbeafe",
                "danger": "#dc2626",
                "warning": "#d97706",
                "success": "#15803d",
            }
        if theme == "agenda":
            return {
                "bg": "#f9fbfd",
                "panel": "#ffffff",
                "panel_soft": "#f3f7fb",
                "line": "#d9e3ee",
                "text": "#13243e",
                "muted": "#64748b",
                "accent": "#1d4ed8",
                "accent_soft": "#e0ecff",
                "danger": "#dc2626",
                "warning": "#c2410c",
                "success": "#15803d",
            }
        if theme == "compact":
            return {
                "bg": "#f8fafc",
                "panel": "#ffffff",
                "panel_soft": "#f4f7fb",
                "line": "#d7dee8",
                "text": "#0f172a",
                "muted": "#64748b",
                "accent": "#1d4ed8",
                "accent_soft": "#e0ecff",
                "danger": "#b91c1c",
                "warning": "#b45309",
                "success": "#15803d",
            }
        return {
            "bg": "#f8fafc",
            "panel": "#ffffff",
            "panel_soft": "#f8fbff",
            "line": "#dbe4f0",
            "text": "#10233f",
            "muted": "#64748b",
            "accent": "#2563eb",
            "accent_soft": "#dbeafe",
            "danger": "#dc2626",
            "warning": "#d97706",
            "success": "#16a34a",
        }
