# -*- coding: utf-8 -*-
"""Word 文档自动排版引擎

支持法律文书等文档的智能识别与排版。
参照 GB/T 9704-2012 简化实现。

v2.2 — 增加网页/AI粘贴垃圾字符清洗。
"""

import re
import shutil
from collections import Counter
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class ParagraphType:
    """段落类型常量"""

    EMPTY = "empty"
    MAIN_TITLE = "main_title"
    LEVEL1_TITLE = "level1_title"
    LEVEL2_TITLE = "level2_title"
    LEVEL3_TITLE = "level3_title"
    BODY = "body"
    SIGNATURE = "signature"  # 落款行（此致、具状人、日期、法院名等）


# ---------------------------------------------------------------------------
# 段落特征提取辅助
# ---------------------------------------------------------------------------

def _get_effective_font_size(para) -> Optional[float]:
    """获取段落的有效字号（pt）。"""
    for run in para.runs:
        if run.font.size is not None:
            return run.font.size.pt
    try:
        if para.style and para.style.font and para.style.font.size:
            return para.style.font.size.pt
    except Exception:
        pass
    return None


def _get_effective_font_name(para) -> Optional[str]:
    """获取段落的有效字体名称。"""
    for run in para.runs:
        if run.font.name:
            return run.font.name
        rpr = run._element.find(qn("w:rPr"))
        if rpr is not None:
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is not None:
                for attr in (qn("w:eastAsia"), qn("w:ascii")):
                    val = rfonts.get(attr)
                    if val:
                        return val
    try:
        if para.style and para.style.font and para.style.font.name:
            return para.style.font.name
    except Exception:
        pass
    return None


def _is_run_bold(run) -> bool:
    """判断 run 是否加粗。"""
    if run.bold is True:
        return True
    if run.font.bold is True:
        return True
    return False


def _is_para_bold(para) -> bool:
    """判断段落是否有加粗 run。"""
    for run in para.runs:
        if _is_run_bold(run):
            return True
    return False


def _is_para_centered(para) -> bool:
    """判断段落是否居中对齐。"""
    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return True
    try:
        if para.style and para.style.paragraph_format:
            if para.style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                return True
    except Exception:
        pass
    return False


def _text_ends_with_punctuation(text: str) -> bool:
    """判断文本是否以句末标点结尾。"""
    return text.endswith(("，", "。", "；", "！", "？", "：", "、", ",", ".", ";", "!", "?", ":"))


def _text_looks_like_body(text: str) -> bool:
    """判断文本是否看起来像正文（含句中标点的稍长文本）。"""
    if len(text) > 20 and _text_ends_with_punctuation(text):
        return True
    if len(text) > 18 and any(p in text for p in ("，", "。", "；")):
        return True
    return False


# ---------------------------------------------------------------------------
# 文本清洗：去除网页/AI 粘贴带来的垃圾字符
# ---------------------------------------------------------------------------

# 零宽/不可见字符 → 直接删除
_INVISIBLE_CHARS = set(
    "\u200B\u200C\u200D\uFEFF\u00AD"       # 零宽空间、BOM、软连字符
    "\u202A\u202B\u202C\u202D\u202E"       # 方向控制符
    "\u2060\u2061\u2062\u2063\u2064"       # 不可见运算符
    "\u2066\u2067\u2068\u2069"             # 方向隔离
    "\u200E\u200F"                         # LTR/RTL 标记
    "\u180E"                               # 蒙古语元音分隔符
)

# 各种空白字符 → 替换为普通空格
_SPACE_CHARS = str.maketrans({
    "\u00A0": " ",   # NO-BREAK SPACE (最常见，&nbsp;)
    "\u2002": " ",   # EN SPACE
    "\u2003": " ",   # EM SPACE
    "\u2004": " ",   # THREE-PER-EM SPACE
    "\u2005": " ",   # FOUR-PER-EM SPACE
    "\u2006": " ",   # SIX-PER-EM SPACE
    "\u2007": " ",   # FIGURE SPACE
    "\u2008": " ",   # PUNCTUATION SPACE
    "\u2009": " ",   # THIN SPACE
    "\u200A": " ",   # HAIR SPACE
    "\u3000": " ",   # IDEOGRAPHIC SPACE (全角空格)
    "\u2028": " ",   # LINE SEPARATOR
    "\u2029": " ",   # PARAGRAPH SEPARATOR
    "\t":     " ",   # TAB
    "\r":    "",     # CARRIAGE RETURN (直接删)
})

# 连续空格合并为一个
_RE_MULTI_SPACE = re.compile(r" {2,}")

# ── Markdown 符号清洗 ──
# **加粗** 或 __加粗__
_RE_MD_BOLD = re.compile(r"\*\*(.+?)\*\*|__(.+?)__")
# *斜体* 或 _斜体_（避免匹配中文下划线名如 祥瑞公司_简称）
_RE_MD_ITALIC = re.compile(r"(?<!\w)\*(.+?)\*(?!\w)|(?<![.\w])_(.+?)_(?![.\w])")
# ~~删除线~~
_RE_MD_STRIKE = re.compile(r"~~(.+?)~~")
# `行内代码`
_RE_MD_CODE = re.compile(r"`(.+?)`")
# [链接文字](url)
_RE_MD_LINK = re.compile(r"\[([^\]]+)\]\([^)]+\)")
# ![图片](url)
_RE_MD_IMAGE = re.compile(r"!\[([^\]]*)\]\([^)]+\)")
# 标题前缀 # ## ###（行首，后跟空格）
_RE_MD_HEADING = re.compile(r"^#{1,6}\s+")
# 任意位置的连续2个及以上#号（需删除）
_RE_MD_HASHES = re.compile(r"#{2,}")
# 无序列表前缀 - * + （行首）
_RE_MD_UL_PREFIX = re.compile(r"^[-*+]\s+")
# 有序列表前缀 1. 2.
_RE_MD_OL_PREFIX = re.compile(r"^\d+\.\s+")
# 引用前缀 >
_RE_MD_QUOTE = re.compile(r"^>\s*")
# 分隔线 --- *** ___
_RE_MD_HR = re.compile(r"^[-*_]{3,}\s*$")


def _strip_md_inline(text: str) -> str:
    """去除 Markdown 行内格式标记，保留文字内容。"""
    text = _RE_MD_IMAGE.sub("", text)        # 图片标记整块删除
    text = _RE_MD_LINK.sub(r"\1", text)       # [文字](url) → 文字
    text = _RE_MD_BOLD.sub(r"\1\2", text)     # **粗体** → 粗体
    text = _RE_MD_ITALIC.sub(r"\1\2", text)   # *斜体* → 斜体
    text = _RE_MD_STRIKE.sub(r"\1", text)     # ~~删除~~ → 删除
    text = _RE_MD_CODE.sub(r"\1", text)       # `代码` → 代码
    return text


def clean_text(text: str) -> str:
    """清洗文本中的网页/AI粘贴垃圾字符。

    处理策略：
    1. 删除零宽/不可见字符（不影响任何显示内容的控制符）
    2. 将各类空白字符统一为普通空格
    3. 合并连续多个空格为单个空格
    4. 去除首尾空白

    不处理的字符（保留原样）：
    - 中英文标点符号（，。！？：；、""''（）《》等）
    - 智能引号（"" ''）
    - 全角符号（，。）—— 这些是正常的中文标点
    - 特殊数字/单位符号

    Markdown 格式处理：
    - **粗体**、*斜体*、~~删除线~~、`代码` → 去除标记保留文字
    - [文字](链接) → 保留文字
    - 行首的 # 标题、- 列表、> 引用 → 去除前缀
    - 分隔线 --- → 清除
    - 行末的 - 破折号（可能是段落分隔标记） → 清除
    """
    # 0. Markdown 行级标记（分隔线等需在行内格式之前处理）
    text = _RE_MD_HR.sub("", text)
    text = _RE_MD_HEADING.sub("", text)
    text = _RE_MD_QUOTE.sub("", text)
    text = _RE_MD_UL_PREFIX.sub("", text)
    text = _RE_MD_OL_PREFIX.sub("", text)

    # 0.2 删除任意位置的连续2个及以上#号
    text = _RE_MD_HASHES.sub("", text)

    # 0.3 Markdown 行内格式清洗
    text = _strip_md_inline(text)

    # 0.5 行末破折号（markdown 段落分隔残留）
    text = text.rstrip("-\u2014\u2013")

    # 1. 删除零宽/不可见字符
    text = "".join(ch for ch in text if ch not in _INVISIBLE_CHARS)

    # 2. 统一空白字符
    text = text.translate(_SPACE_CHARS)

    # 3. 合并连续空格
    text = _RE_MULTI_SPACE.sub(" ", text)

    # 4. 首尾空白
    text = text.strip()

    return text


def _split_paragraphs_on_separator(doc) -> None:
    """将含段落分隔符的段落拆分为多个段落。

    处理的分隔符：
    - " - "（空格-空格）：网页/AI复制用做段落分隔
    - "##" 及以上：markdown 标题标记残留

    例如：
      "第一段内容 - 第二段内容" → 两个段落
      "诉讼请求##事实与理由" → 两个段落
    """
    from docx.oxml.ns import qn as _qn
    paragraphs = list(doc.paragraphs)
    for para in paragraphs:
        full_text = para.text.strip()
        if not full_text:
            continue

        need_split = False
        # 检查是否有 " - " 分隔符
        if " - " in full_text:
            need_split = True
        # 检查是否有 ## 分隔符（连续2个及以上#号）
        if re.search(r"#{2,}", full_text):
            need_split = True

        if not need_split:
            continue

        # 统一用正则拆分：按 " - " 或 "##+" 分割
        parts = re.split(r"\s+-\s+|#{2,}", full_text)
        parts = [p.strip() for p in parts if p.strip()]

        if len(parts) <= 1:
            continue

        # 第一段保留在原段落
        for run in list(para.runs):
            run._element.getparent().remove(run._element)
        para.add_run(parts[0])

        # 后续段落插入到原段落后面
        p_elem = para._element
        for part in parts[1:]:
            new_p = p_elem.makeelement(_qn("w:p"), {})
            new_r = new_p.makeelement(_qn("w:r"), {})
            new_t = new_r.makeelement(_qn("w:t"), {})
            new_t.text = part
            new_r.append(new_t)
            new_p.append(new_r)
            p_elem.addnext(new_p)
            p_elem = new_p


def clean_document(doc) -> int:
    """清洗 Word 文档中所有段落的文本，返回修改的 run 数量。

    处理步骤：
    1. 拆分含 " - " 的段落（网页/AI段落分隔符）
    2. 逐 run 清洗文本
    3. 清空全是垃圾字符的段落
    """
    # 1. 段落拆分（按 " - " 和 "##" 分隔符）
    _split_paragraphs_on_separator(doc)

    # 2. 逐 run 清洗
    changed = 0
    for para in doc.paragraphs:
        all_empty = True
        has_content = False
        for run in para.runs:
            original = run.text
            if original.strip():
                has_content = True
            cleaned = clean_text(original)
            if cleaned != original:
                run.text = cleaned
                changed += 1
            if cleaned.strip():
                all_empty = False

        # 如果整段清洗后完全为空（原本全是 Markdown 标记/垃圾字符），
        # 将段落文本设为空字符串
        if has_content and all_empty:
            for run in para.runs:
                run.text = ""
    return changed


class DocxAutoFormatter:
    """Word 文档自动排版器"""

    # ── 大标题关键词 ──
    MAIN_TITLE_KEYWORDS = [
        "起诉状", "答辩状", "判决书", "裁定书", "调解书",
        "合同", "协议", "申请书", "上诉状", "反诉状",
        "委托书", "公证书", "仲裁书", "通知书", "意见书",
        "代理词", "辩护词", "律师函", "法律意见书",
        "催告函", "告知书", "承诺书", "保证书",
        "民事起诉状", "民事答辩状", "行政起诉状", "刑事附带民事起诉状",
        "劳动仲裁申请书", "仲裁申请书", "强制执行申请书",
        "财产保全申请书", "先予执行申请书",
        "授权委托书", "法定代表人身份证明",
        "证据目录", "证据清单", "质证意见", "举证意见",
        "庭审笔录", "询问笔录", "调查笔录",
        "财产清单", "债务清单", "还款计划",
        "撤诉申请书", "再审申请书", "抗诉申请书",
        "管辖异议申请书", "回避申请书",
        "司法鉴定意见书", "评估报告",
    ]

    # ── 一级标题：高置信度模式（纯标题，短文本独占一行）──
    LEVEL1_PATTERNS_ALWAYS = [
        # 法律文书段落标题 —— 独立成行的短标题
        r"^(诉讼请求|仲裁请求)[：:]*\s*$",
        r"^(事实与理由|事实和理由|案由|案情)[：:]*\s*$",
        r"^(证据[及和]?|证据材料|证据目录)[：:]*\s*$",
        r"^(代理意见|辩护意见|答辩意见|答辩理由)[：:]*\s*$",
        r"^(上诉请求|上诉理由|申请事项|复议请求)[：:]*\s*$",
        r"^(裁判结果|裁判理由|本院认为|审理查明|审理经过)[：:]*\s*$",
        r"^(争议焦点|焦点问题)[：:]*\s*$",
        r"^(法律依据|适用法律|相关法律)[：:]*\s*$",
        r"^(诉讼费|案件受理费|费用承担)[：:]*\s*$",
        # "关于XXX" 格式标题（短文本，无句末标点）
        r"^关于.{2,20}$",
        # 案号独立行
        r"^[（(]\d{4}[）)].+[民刑行]初?\d+号$",
    ]

    # ── 落款模式（文档尾部的签名/日期/致词等，不算标题）──
    SIGNATURE_PATTERNS = [
        r"^此致\s*$",
        r"^(具状人|起诉人|申请人|上诉人|答辩人|代理人)[：:]",
        r"^(审判[长员]|代理审判员|人民陪审员|书记员|法官助理)[：:]?\s*\S",
        r"^\d{4}年\d{1,2}月\d{1,2}日\s*$",
    ]

    # ── 一级标题：需验证模式（可能是标题也可能是正文编号）──
    LEVEL1_PATTERNS_AMBIGUOUS = [
        # 当事人信息行（需验证长度——短的是标题，长的是正文）
        r"^(原告|被告|第三人|上诉人|被上诉人|申请人|被申请人|再审申请人|再审被申请人|异议人)[：:]",
        r"^(法定代表人|委托诉讼代理人|委托代理人|诉讼代理人|辩护人)[：:]",
        # 法院文书首部/尾部（需验证长度）
        r"^(如不服本判决|如不服本裁定|本判决[为书]|本裁定[为书])",
        # "一、" "二、" 中文数字编号（需排除以；。结尾的长列举项）
        r"^[一二三四五六七八九十百]+[、．.]",
        # "（一）" "（二）" 带括号中文数字
        r"^[（(][一二三四五六七八九十百]+[）)]",
        # "第X条" 法条编号
        r"^第[一二三四五六七八九十百零\d]+[条章节款]",
        # 阿拉伯数字编号（1. 2.）
        r"^\d{1,2}[、．.]\s*\S",
    ]

    # ── 二级标题 ──
    LEVEL2_PATTERNS = [
        r"^\([\d一二三四五六七八九十]+\)",
        r"^\d+\.\d+[、．.]?\s",
        r"^[（(]\d+[）)]",
        r"^第\d+款",
    ]

    # ── 三级标题 ──
    LEVEL3_PATTERNS = [
        r"^[①②③④⑤⑥⑦⑧⑨⑩]",
        r"^[（(]\d+[）)]\s*\S",
    ]

    # ── 短关键词（容易在正文中误命中，需要更严格匹配）──
    _SHORT_KEYWORDS = {"合同", "协议", "通知书", "意见书", "告知书"}

    def _has_main_title_keyword(self, text: str) -> bool:
        """检查文本是否包含大标题关键词。

        对于短关键词（如"合同"），要求文本整体等于关键词或以关键词结尾，
        避免在长正文中误命中。
        """
        for kw in self.MAIN_TITLE_KEYWORDS:
            if kw in text:
                if kw in self._SHORT_KEYWORDS:
                    # 短关键词：要求文本就是关键词本身，或以关键词结尾
                    if text == kw or text.endswith(kw):
                        return True
                else:
                    return True
        return False

    # 标题常用字体
    TITLE_FONTS = {
        "黑体", "楷体", "楷体_GB2312", "SimHei", "SimKai",
        "方正小标宋简体", "方正小标宋_GBK", "华文中宋", "STZhongsong",
        "宋体", "仿宋", "仿宋_GB2312",  # 有时也用作标题
    }

    def __init__(self):
        self._doc: Optional[Document] = None
        self._paragraph_types: List[str] = []
        self._stats: Dict = {}
        self._original_path: Optional[Path] = None
        self._first_non_empty_idx: int = -1

    # ------------------------------------------------------------------
    # 加载与分析
    # ------------------------------------------------------------------

    def load(self, path: Path) -> None:
        """加载 Word 文档。"""
        self._original_path = path
        self._doc = Document(str(path))
        self._cleaned_count = clean_document(self._doc)
        self._analyze()

    def _analyze(self) -> None:
        """分析文档结构，为每个段落分类。"""
        paragraphs = self._doc.paragraphs
        self._stats = self._compute_stats(paragraphs)

        # 找到第一个非空段落索引
        self._first_non_empty_idx = -1
        for i, p in enumerate(paragraphs):
            if p.text.strip():
                self._first_non_empty_idx = i
                break

        # 分类前先做一轮前后上下文修正——合并相邻同类项
        self._paragraph_types = []
        for i, para in enumerate(paragraphs):
            ptype = self._classify_paragraph(para, i, paragraphs)
            self._paragraph_types.append(ptype)

    def _compute_stats(self, paragraphs) -> Dict:
        """计算全局统计信息。"""
        font_sizes = []
        for para in paragraphs:
            sz = _get_effective_font_size(para)
            if sz is not None:
                font_sizes.append(sz)

        if not font_sizes:
            font_sizes = [12.0]

        size_counter = Counter(font_sizes)
        most_common_size = size_counter.most_common(1)[0][0]
        max_size = max(font_sizes)
        min_size = min(font_sizes)

        bold_count = sum(1 for p in paragraphs if _is_para_bold(p))
        bold_ratio = bold_count / max(len(paragraphs), 1)

        return {
            "most_common_size": most_common_size,
            "max_size": max_size,
            "min_size": min_size,
            "bold_ratio": bold_ratio,
            "total_paragraphs": len(paragraphs),
            "size_counter": size_counter,
        }

    # ------------------------------------------------------------------
    # 分类主逻辑
    # ------------------------------------------------------------------

    def _classify_paragraph(self, para, index: int, all_paras) -> str:
        """对单个段落分类。

        优先级：空段落 > 落款模式 > 首行大标题 > 文本模式 > 字体特征 > 默认正文
        """
        text = para.text.strip()
        if not text:
            return ParagraphType.EMPTY

        # 0. 落款模式检测（优先于标题检测）
        if self._match_signature(text, para, index, all_paras):
            return ParagraphType.SIGNATURE

        # 0.5 首行大标题：文档第一个非空段落，如果足够短且不像正文
        if index == self._first_non_empty_idx:
            if self._looks_like_main_title_by_position(text, para):
                return ParagraphType.MAIN_TITLE

        # 1. 文本模式匹配
        ptype = self._match_by_text_patterns(text, para, index)
        if ptype is not None:
            return ptype

        # 2. 字体特征判断
        ptype = self._match_by_features(text, para, index, all_paras)
        if ptype is not None:
            return ptype

        return ParagraphType.BODY

    def _looks_like_main_title_by_position(self, text: str, para) -> bool:
        """判断文档第一个非空段落是否像大标题。

        规则：短文本（≤30字）+ 不以句末标点结尾 + 不以当事人/编号开头 → 大概率是标题。
        """
        if len(text) > 30:
            return False
        if _text_ends_with_punctuation(text):
            return False

        # 排除当事人/编号开头——这些不应该是大标题
        _exclude_prefixes = [
            r"^(原告|被告|第三人|上诉人|被上诉人|申请人|被申请人)[：:]",
            r"^(法定代表人|委托诉讼代理人|委托代理人|诉讼代理人|辩护人)[：:]",
            r"^[一二三四五六七八九十百]+[、．.]",
            r"^[（(][一二三四五六七八九十百]+[）)]",
            r"^\d{1,2}[、．.]\s",
        ]
        for p in _exclude_prefixes:
            if re.match(p, text):
                return False

        return True

    def _match_signature(self, text: str, para, index: int, all_paras) -> bool:
        """判断段落是否为落款行（此致、具状人、日期、审判员签名等）。

        落款行不应当作标题处理，排版时应保留原有对齐方式（通常右对齐）。
        """
        for pattern in self.SIGNATURE_PATTERNS:
            if re.match(pattern, text):
                return True

        # 右对齐 + 短文本 + 在文档后半部分 → 可能是落款
        is_right = para.alignment == WD_ALIGN_PARAGRAPH.RIGHT
        if is_right and len(text) <= 30 and index > len(all_paras) * 0.5:
            return True

        return False

    def _match_by_text_patterns(self, text: str, para, index: int) -> Optional[str]:
        """通过文本正则模式匹配标题。

        在匹配 LEVEL1 模式后会二次检查：如果该段落有强 MAIN_TITLE 信号
        （关键词 + 居中/靠前 + 短文本），则提升为大标题。
        """
        # 一级标题 —— 高置信度
        for pattern in self.LEVEL1_PATTERNS_ALWAYS:
            if re.match(pattern, text):
                # 检查是否应提升为 MAIN_TITLE
                if self._promote_to_main_title(text, para, index):
                    return ParagraphType.MAIN_TITLE
                return ParagraphType.LEVEL1_TITLE

        # 一级标题 —— 需验证模式
        for pattern in self.LEVEL1_PATTERNS_AMBIGUOUS:
            if re.match(pattern, text):
                if self._looks_like_body_despite_pattern(text):
                    return None
                if self._promote_to_main_title(text, para, index):
                    return ParagraphType.MAIN_TITLE
                return ParagraphType.LEVEL1_TITLE

        # 二级标题
        for pattern in self.LEVEL2_PATTERNS:
            if re.match(pattern, text):
                if self._looks_like_body_despite_pattern(text):
                    return None
                return ParagraphType.LEVEL2_TITLE

        # 三级标题
        for pattern in self.LEVEL3_PATTERNS:
            if re.match(pattern, text):
                return ParagraphType.LEVEL3_TITLE

        return None

    def _promote_to_main_title(self, text: str, para, index: int) -> bool:
        """判断已匹配 LEVEL1 模式的段落是否应提升为 MAIN_TITLE。

        条件：包含大标题关键词 AND 至少满足以下之一：
        - 居中对齐
        - 全文前 3 个非空段落
        - 字号 >= 全文正文字号 + 3pt
        """
        if not self._has_main_title_keyword(text):
            return False
        score = 0
        if _is_para_centered(para):
            score += 1
        if index <= self._first_non_empty_idx + 2:
            score += 1
        common_size = self._stats.get("most_common_size", 0)
        font_size = _get_effective_font_size(para)
        if font_size is not None and font_size >= common_size + 3:
            score += 1
        return score >= 1

    def _looks_like_body_despite_pattern(self, text: str) -> bool:
        """判断匹配了标题模式的文本是否实际是正文。

        典型场景：
        - "一、判令被告支付货款100,000元及利息；" → 正文（编号列举）
        - "原告：张三，男，1990年出生，住……" → 正文（当事人信息长行）
        - "如不服本判决，可在判决书送达之日起十五日内……" → 正文
        - "（一）启动评估给全体债权人的告知函已经准备，待发" → 正文
        """
        # 以句号或分号结尾 → 正文
        if text.endswith(("；", "。", "；", "。", ";", ".")):
            return True

        # 包含逗号且超过 20 字的以逗号结尾的长句 → 正文
        if len(text) > 20 and text.endswith("，"):
            return True
        if len(text) > 20 and _text_ends_with_punctuation(text):
            return True

        # 较长文本（>35字）无论结尾如何都大概率是正文
        if len(text) > 35:
            return True

        return False

    def _match_by_features(self, text: str, para, index: int, all_paras) -> Optional[str]:
        """通过综合特征判断段落类型。"""
        font_size = _get_effective_font_size(para)
        is_bold = _is_para_bold(para)
        is_centered = _is_para_centered(para)
        common_size = self._stats["most_common_size"]
        font_name = _get_effective_font_name(para)
        max_size = self._stats["max_size"]

        # ── 大标题检测 ──
        main_score = 0
        has_keyword = self._has_main_title_keyword(text)
        if has_keyword:
            main_score += 3
        if is_centered:
            main_score += 2
        if font_size is not None and font_size >= common_size + 6:
            main_score += 3
        elif font_size is not None and font_size >= common_size + 3:
            main_score += 2
        elif font_size is not None and font_size >= common_size + 1.5:
            main_score += 1
        if index <= 4:
            main_score += 1
        if len(text) <= 30:
            main_score += 1
        elif len(text) > 55:
            main_score -= 2
        if not _text_ends_with_punctuation(text):
            main_score += 1

        # 如果它是全文最大字号（且不是全部统一字号），单独加分
        if font_size is not None and max_size > common_size and font_size == max_size:
            main_score += 1

        if main_score >= 4:
            return ParagraphType.MAIN_TITLE

        # ── 一级标题特征型 ──
        # 条件：字号明显大于正文 + 加粗 + 短文本
        if (
            font_size is not None
            and font_size >= common_size + 2
            and is_bold
            and len(text) <= 35
            and not _text_ends_with_punctuation(text)
        ):
            return ParagraphType.LEVEL1_TITLE

        # 条件：居中对齐 + 加粗 + 短文本（常见于无样式标题）
        if (
            is_centered
            and is_bold
            and len(text) <= 25
            and not _text_ends_with_punctuation(text)
        ):
            return ParagraphType.LEVEL1_TITLE

        # ── 标题字体特征 ──
        if font_name and font_name in self.TITLE_FONTS:
            if font_size is not None and font_size > common_size and len(text) <= 40:
                if len(text) > 25 and _text_ends_with_punctuation(text):
                    return None
                return ParagraphType.LEVEL1_TITLE

        return None

    # ------------------------------------------------------------------
    # 公共查询接口
    # ------------------------------------------------------------------

    def get_paragraph_types(self) -> List[str]:
        return list(self._paragraph_types)

    def get_paragraph_info(self) -> List[Tuple[str, str, str]]:
        """返回每个段落的 (类型, 文本前60字, 对齐方式)。"""
        result = []
        for para, ptype in zip(self._doc.paragraphs, self._paragraph_types):
            text = para.text.strip()
            preview = text[:60] + ("…" if len(text) > 60 else "")
            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                align = "居中"
            elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                align = "右对齐"
            elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                align = "两端对齐"
            else:
                align = "左对齐"
            result.append((ptype, preview, align))
        return result

    # ------------------------------------------------------------------
    # 格式应用
    # ------------------------------------------------------------------

    def apply_format(self, style_config: Optional[Dict] = None) -> None:
        """应用排版格式。"""
        config = style_config or self._default_config()
        self._apply_page_format(config.get("page", {}))

        for para, ptype in zip(self._doc.paragraphs, self._paragraph_types):
            style = config.get(ptype, config.get("body"))
            if style is None:
                continue
            self._apply_paragraph_format(para, style, ptype)

    def _default_config(self) -> Dict:
        return {
            "page": {
                "top_margin": Cm(3.7),
                "bottom_margin": Cm(3.5),
                "left_margin": Cm(2.8),
                "right_margin": Cm(2.6),
            },
            "main_title": {
                "font_name": "方正小标宋简体",
                "font_size": Pt(22),
                "bold": False,
                "alignment": WD_ALIGN_PARAGRAPH.CENTER,
                "first_line_indent": Cm(0),
                "line_spacing": Pt(36),
            },
            "level1_title": {
                "font_name": "黑体",
                "font_size": Pt(16),
                "bold": True,
                "alignment": WD_ALIGN_PARAGRAPH.LEFT,
                "first_line_indent": Cm(0),
                "line_spacing": Pt(28.8),
            },
            "level2_title": {
                "font_name": "楷体",
                "font_size": Pt(16),
                "bold": True,
                "alignment": WD_ALIGN_PARAGRAPH.LEFT,
                "first_line_indent": Cm(0),
                "line_spacing": Pt(28.8),
            },
            "level3_title": {
                "font_name": "仿宋_GB2312",
                "font_size": Pt(16),
                "bold": True,
                "alignment": WD_ALIGN_PARAGRAPH.LEFT,
                "first_line_indent": Pt(32),
                "line_spacing": Pt(28.8),
            },
            "body": {
                "font_name": "仿宋_GB2312",
                "font_size": Pt(16),
                "bold": False,
                "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
                "first_line_indent": Pt(32),
                "line_spacing": Pt(28.8),
            },
            "signature": {
                "font_name": "仿宋_GB2312",
                "font_size": Pt(16),
                "bold": False,
                # 不设 alignment — 保留原始对齐（通常右对齐）
                "first_line_indent": Cm(0),
                "line_spacing": Pt(28.8),
            },
            "empty": {
                "font_name": "仿宋_GB2312",
                "font_size": Pt(16),
                "bold": False,
                "first_line_indent": Cm(0),
                "line_spacing": Pt(28.8),
            },
        }

    def _apply_page_format(self, page_config: Dict) -> None:
        """应用页面格式。"""
        for section in self._doc.sections:
            if "top_margin" in page_config:
                section.top_margin = page_config["top_margin"]
            if "bottom_margin" in page_config:
                section.bottom_margin = page_config["bottom_margin"]
            if "left_margin" in page_config:
                section.left_margin = page_config["left_margin"]
            if "right_margin" in page_config:
                section.right_margin = page_config["right_margin"]

    def _apply_paragraph_format(self, para, style: Dict, ptype: str) -> None:
        """应用段落格式——彻底清理并重新设置排版。

        策略：
        1. 解除段落样式绑定，避免样式继承干扰
        2. 合并所有 run 文本，清空后重建为单个 run（统一字体）
        3. 设置段落级属性（对齐、缩进、行距）
        """
        # ── 1. 解除段落样式引用 ──
        pPr = para._element.get_or_add_pPr()
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is not None:
            pPr.remove(pStyle)

        # ── 2. 收集文本，移除原 run ──
        full_text = para.text

        # 移除所有 run（从后往前删除）
        for run in reversed(para.runs):
            run._element.getparent().remove(run._element)

        # ── 3. 创建一个新 run 并应用全部字体格式 ──
        new_run = para.add_run(full_text) if full_text else para.add_run(" ")

        font_name = style.get("font_name")
        if font_name:
            new_run.font.name = font_name
            rpr = new_run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:ascii"), font_name)
            rfonts.set(qn("w:hAnsi"), font_name)
            rfonts.set(qn("w:eastAsia"), font_name)
            rfonts.set(qn("w:cs"), font_name)

        font_size = style.get("font_size")
        if font_size is not None:
            new_run.font.size = font_size

        bold_value = style.get("bold")
        if bold_value is not None:
            new_run.font.bold = bold_value

        # ── 4. 段落格式 ──
        if "alignment" in style:
            para.alignment = style["alignment"]

        pf = para.paragraph_format
        # 先清除可能存在的缩进
        if "first_line_indent" in style:
            pf.first_line_indent = style["first_line_indent"]
        else:
            pf.first_line_indent = Cm(0)

        if "line_spacing" in style:
            pf.line_spacing = style["line_spacing"]
            pf.line_spacing_rule = 3  # 固定值

        # 清除段前段后间距（Word 默认有段后间距）
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    # ------------------------------------------------------------------
    # 保存
    # ------------------------------------------------------------------

    def save(self, path: Optional[Path] = None) -> Path:
        """保存文档。"""
        output_path = path or self._original_path
        if output_path is None:
            raise ValueError("未指定保存路径")
        self._doc.save(str(output_path))
        return output_path

    @staticmethod
    def backup_original(path: Path) -> Path:
        """备份原始文件，返回备份路径。"""
        backup_path = path.parent / f"{path.stem}_备份{path.suffix}"
        counter = 1
        while backup_path.exists():
            backup_path = path.parent / f"{path.stem}_备份({counter}){path.suffix}"
            counter += 1
        shutil.copy2(str(path), str(backup_path))
        return backup_path
