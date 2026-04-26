# -*- coding: utf-8 -*-
"""模板诊断工具。"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Set

from src.utils.template_path_manager import TemplatePathManager


VARIABLE_PATTERN = re.compile(r"\{\{(\w+)\}\}")


@dataclass
class TemplateDiagnosticIssue:
    """单条模板诊断结果。"""

    level: str
    template_id: str
    template_name: str
    category: str
    message: str
    path: str = ""


@dataclass
class TemplateDiagnosticSummary:
    """模板诊断汇总。"""

    template_count: int
    invalid_paths: int
    duplicate_ids: int
    missing_variables: int
    placeholder_count: int
    issues: List[TemplateDiagnosticIssue]


def _iter_template_file_refs(template: Dict[str, Any]) -> Iterable[str]:
    """遍历模板中引用的 Word 文件路径。"""
    template_file = str(template.get("template_file", "")).strip()
    if template_file:
        yield template_file

    def walk_folder(node: Dict[str, Any]) -> Iterable[str]:
        template_path = str(node.get("template_path", "")).strip()
        if template_path:
            yield template_path
        for child in node.get("subfolders", []) or []:
            if isinstance(child, dict):
                yield from walk_folder(child)
        for child in node.get("children", []) or []:
            if isinstance(child, dict):
                yield from walk_folder(child)

    folder_structure = template.get("folder_structure", {}) or {}
    for folder in folder_structure.get("folders", []) or []:
        if isinstance(folder, dict):
            yield from walk_folder(folder)


def _extract_docx_placeholders(path: Path) -> Set[str]:
    """从 docx 段落中提取变量占位符。"""
    try:
        from docx import Document
    except Exception:
        return set()

    try:
        document = Document(str(path))
    except Exception:
        return set()

    text_parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)
    return set(VARIABLE_PATTERN.findall("\n".join(text_parts)))


def diagnose_templates(
    templates: List[Dict[str, Any]],
    path_manager: Optional[TemplatePathManager] = None,
) -> TemplateDiagnosticSummary:
    """诊断模板路径、重复 ID、变量缺失与 Word 占位符数量。"""
    manager = path_manager or TemplatePathManager()
    issues: List[TemplateDiagnosticIssue] = []
    seen_ids: Dict[str, int] = {}
    placeholder_count = 0
    invalid_paths = 0
    missing_variables = 0

    for template in templates:
        template_id = str(template.get("id", "")).strip()
        template_name = str(template.get("name", "")).strip() or "未命名模板"
        category = str(template.get("category", "")).strip()
        seen_ids[template_id] = seen_ids.get(template_id, 0) + 1

        defined_vars = {
            str(var.get("key", "")).strip()
            for var in template.get("variables", []) or []
            if isinstance(var, dict) and str(var.get("key", "")).strip()
        }
        placeholders: Set[str] = set()

        for ref in _iter_template_file_refs(template):
            resolved = manager.resolve_template_path(ref)
            if resolved is None:
                invalid_paths += 1
                issues.append(TemplateDiagnosticIssue(
                    level="error",
                    template_id=template_id,
                    template_name=template_name,
                    category=category,
                    message="模板文件不存在或路径不安全",
                    path=ref,
                ))
                continue
            if resolved.suffix.lower() == ".docx":
                placeholders.update(_extract_docx_placeholders(resolved))

        placeholder_count += len(placeholders)
        missing = sorted(placeholders - defined_vars)
        if missing:
            missing_variables += len(missing)
            issues.append(TemplateDiagnosticIssue(
                level="warning",
                template_id=template_id,
                template_name=template_name,
                category=category,
                message=f"Word 占位符未在变量列表中定义：{', '.join(missing)}",
            ))

    duplicate_ids = 0
    for template_id, count in seen_ids.items():
        if template_id and count > 1:
            duplicate_ids += count
            issues.append(TemplateDiagnosticIssue(
                level="error",
                template_id=template_id,
                template_name="",
                category="",
                message=f"模板 ID 重复出现 {count} 次",
            ))

    return TemplateDiagnosticSummary(
        template_count=len(templates),
        invalid_paths=invalid_paths,
        duplicate_ids=duplicate_ids,
        missing_variables=missing_variables,
        placeholder_count=placeholder_count,
        issues=issues,
    )
