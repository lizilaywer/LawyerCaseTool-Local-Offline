# -*- coding: utf-8 -*-
"""电子化归档引擎

提供变量替换、文档导出、历史记录等功能。
"""

from pathlib import Path
from typing import Dict, List, Optional, Any
from datetime import datetime
import json
import re

from docxtpl import DocxTemplate

from src.utils.logger import get_logger
from src.config.path_manager import get_path_manager
from src.utils.file_utils import write_json_file


class ArchiveEngine:
    """归档引擎

    功能：
    - 变量替换导出
    - 从文档提取变量
    - 保存/加载历史记录
    """

    # 变量匹配正则
    VARIABLE_PATTERN = re.compile(r'\{\{(\w+)\}\}')

    def __init__(self):
        self._logger = get_logger()
        self._path_manager = get_path_manager()

    def export(
        self,
        template_path: Path,
        variables: Dict[str, str],
        output_path: Path
    ) -> Path:
        """导出文档

        Args:
            template_path: 模板文件路径
            variables: 变量键值对
            output_path: 输出文件路径

        Returns:
            输出文件路径
        """
        try:
            # 加载模板
            doc = DocxTemplate(str(template_path))

            # 准备上下文（跳过空值）
            context = {}
            for key, value in variables.items():
                if value is not None and value != "":
                    context[key] = value

            # 渲染文档
            doc.render(context)

            # 保存文档
            doc.save(str(output_path))

            self._logger.info(f"文档已导出: {output_path}")
            return Path(output_path)

        except Exception as e:
            self._logger.error(f"导出文档失败: {e}")
            raise

    def extract_variables(self, doc_path: Path) -> List[str]:
        """从文档提取变量

        Args:
            doc_path: 文档路径

        Returns:
            变量名列表
        """
        try:
            variables = set()

            from docx import Document

            document = Document(str(doc_path))

            # 遍历所有段落
            for paragraph in document.paragraphs:
                text = paragraph.text
                matches = self.VARIABLE_PATTERN.findall(text)
                variables.update(matches)

            # 遍历所有表格
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text
                        matches = self.VARIABLE_PATTERN.findall(text)
                        variables.update(matches)

            return sorted(list(variables))

        except Exception as e:
            self._logger.error(f"提取变量失败: {e}")
            return []

    def save_history(self, record: Dict[str, Any]) -> None:
        """保存历史记录

        Args:
            record: 历史记录
        """
        try:
            history_file = self._get_history_file()

            # 读取现有历史
            history = []
            if history_file.exists():
                with open(history_file, 'r', encoding='utf-8') as f:
                    history = json.load(f)

            # 添加新记录
            history.append(record)

            # 限制历史记录数量（最多100条）
            if len(history) > 100:
                history = history[-100:]

            # 原子写入
            write_json_file(history_file, history)

            self._logger.info(f"历史记录已保存")

        except Exception as e:
            self._logger.error(f"保存历史记录失败: {e}")

    def load_history(self, limit: int = 20) -> List[Dict[str, Any]]:
        """加载历史记录

        Args:
            limit: 最多返回的记录数

        Returns:
            历史记录列表
        """
        try:
            history_file = self._get_history_file()

            if not history_file.exists():
                return []

            with open(history_file, 'r', encoding='utf-8') as f:
                history = json.load(f)

            # 按时间倒序排序
            history.sort(key=lambda x: x.get("created_at", ""), reverse=True)

            return history[:limit]

        except Exception as e:
            self._logger.error(f"加载历史记录失败: {e}")
            return []

    def clear_history(self) -> None:
        """清空历史记录"""
        try:
            history_file = self._get_history_file()

            if history_file.exists():
                history_file.unlink()

            self._logger.info("历史记录已清空")

        except Exception as e:
            self._logger.error(f"清空历史记录失败: {e}")

    def _get_history_file(self) -> Path:
        """获取历史记录文件路径"""
        config_dir = self._path_manager.config_dir
        archive_dir = config_dir / "archive"

        if not archive_dir.exists():
            archive_dir.mkdir(parents=True)

        return archive_dir / "history.json"
