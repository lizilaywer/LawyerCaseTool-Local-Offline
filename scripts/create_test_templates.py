# -*- coding: utf-8 -*-
"""创建测试 Word 模板文件"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_entrustment_contract_template():
    """创建委托合同模板"""
    doc = Document()

    # 标题
    title = doc.add_heading('委托代理合同', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 合同内容
    doc.add_paragraph('合同编号：{{合同编号}}')
    doc.add_paragraph('签订日期：{{签订日期}}')
    doc.add_paragraph()

    # 甲方（委托人）
    doc.add_heading('甲方（委托人）：', level=2)
    p = doc.add_paragraph()
    p.add_run('姓名/名称：{{委托人姓名}}\n')
    p.add_run('身份证号/统一社会信用代码：{{委托人证件号码}}\n')
    p.add_run('联系电话：{{委托人电话}}\n')
    p.add_run('地址：{{委托人地址}}')

    # 乙方（受托人）
    doc.add_heading('乙方（受托人）：', level=2)
    p = doc.add_paragraph()
    p.add_run('名称：{{律师事务所名称}}\n')
    p.add_run('负责人：{{负责人姓名}}\n')
    p.add_run('联系电话：{{律师事务所电话}}\n')
    p.add_run('地址：{{律师事务所地址}}')

    # 案件信息
    doc.add_heading('案件信息：', level=2)
    p = doc.add_paragraph()
    p.add_run('案号：{{案号}}\n')
    p.add_run('案由：{{案由}}\n')
    p.add_run('对方当事人：{{对方当事人}}\n')
    p.add_run('受理法院：{{受理法院}}')

    # 代理事项
    doc.add_heading('代理事项：', level=2)
    doc.add_paragraph('{{代理事项}}')

    # 代理权限
    doc.add_heading('代理权限：', level=2)
    doc.add_paragraph('{{代理权限}}')

    # 费用
    doc.add_heading('代理费用：', level=2)
    p = doc.add_paragraph()
    p.add_run('代理费：人民币 {{代理费金额}} 元\n')
    p.add_run('支付方式：{{支付方式}}\n')
    p.add_run('支付期限：{{支付期限}}')

    # 承办律师
    doc.add_heading('承办律师：', level=2)
    p = doc.add_paragraph()
    p.add_run('姓名：{{承办律师}}\n')
    p.add_run('联系电话：{{律师电话}}')

    # 签署
    doc.add_paragraph()
    doc.add_paragraph('甲方（签字/盖章）：_____________  日期：{{签署日期}}')
    doc.add_paragraph('乙方（签字/盖章）：_____________  日期：{{签署日期}}')

    return doc


def create_power_of_attorney_template():
    """创建授权委托书模板"""
    doc = Document()

    # 标题
    title = doc.add_heading('授权委托书', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 委托人信息
    doc.add_paragraph('委托人：{{委托人姓名}}')
    doc.add_paragraph('身份证号：{{委托人证件号码}}')
    doc.add_paragraph('联系电话：{{委托人电话}}')
    doc.add_paragraph()

    # 受托人信息
    doc.add_paragraph('受托人：{{受托律师姓名}}')
    doc.add_paragraph('工作单位：{{律师事务所名称}}')
    doc.add_paragraph('联系电话：{{律师电话}}')
    doc.add_paragraph()

    # 案件信息
    doc.add_paragraph('案号：{{案号}}')
    doc.add_paragraph('案由：{{案由}}')
    doc.add_paragraph('审理法院：{{受理法院}}')
    doc.add_paragraph()

    # 委托事项
    doc.add_heading('委托事项：', level=2)
    doc.add_paragraph('{{委托事项}}')

    # 代理权限
    doc.add_heading('代理权限：', level=2)
    doc.add_paragraph('''代理权限为：{{代理权限}}

一般代理权限包括：
1. 代为起诉、应诉
2. 代为承认、放弃、变更诉讼请求
3. 代为进行和解、调解
4. 代为签收法律文书

特别授权权限还包括：
1. 代为提起上诉
2. 代为提起反诉
3. 代为申请执行
4. 代为进行财产保全''')

    # 签署
    doc.add_paragraph()
    doc.add_paragraph('委托人（签字/盖章）：_____________')
    doc.add_paragraph('日期：{{签署日期}}')

    return doc


def create_court_notice_template():
    """创建出庭通知书模板"""
    doc = Document()

    # 标题
    title = doc.add_heading('出庭通知书', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 通知内容
    p = doc.add_paragraph()
    p.add_run('{{委托人姓名}}').bold = True
    p.add_run(' 先生/女士：')

    doc.add_paragraph()
    doc.add_paragraph('关于您与 {{对方当事人}} {{案由}} 一案（案号：{{案号}}），定于 {{开庭日期}} {{开庭时间}} 在 {{开庭地点}} 开庭审理。')

    doc.add_heading('开庭信息：', level=2)
    p = doc.add_paragraph()
    p.add_run('案件名称：{{案件名称}}\n')
    p.add_run('案号：{{案号}}\n')
    p.add_run('开庭时间：{{开庭日期}} {{开庭时间}}\n')
    p.add_run('开庭地点：{{开庭地点}}\n')
    p.add_run('承办律师：{{承办律师}}\n')
    p.add_run('联系电话：{{律师电话}}')

    doc.add_heading('注意事项：', level=2)
    doc.add_paragraph('''1. 请携带本人身份证件
2. 请提前15分钟到达法庭
3. 如有证据需提交，请提前3日告知律师
4. 如需变更开庭时间，请及时与法院联系
5. 联系电话：{{联系电话}}''')

    # 签署
    doc.add_paragraph()
    doc.add_paragraph('{{律师事务所名称}}')
    doc.add_paragraph('{{通知日期}}')

    return doc


def main():
    """主函数"""
    templates_dir = Path(__file__).parent.parent / "templates" / "civil"
    templates_dir.mkdir(parents=True, exist_ok=True)

    # 创建模板文件
    templates = [
        ("委托合同模板.docx", create_entrustment_contract_template),
        ("授权委托书模板.docx", create_power_of_attorney_template),
        ("出庭通知书模板.docx", create_court_notice_template),
    ]

    for filename, create_func in templates:
        output_path = templates_dir / filename
        doc = create_func()
        doc.save(str(output_path))
        print(f"Created template: {output_path}")


if __name__ == "__main__":
    main()
