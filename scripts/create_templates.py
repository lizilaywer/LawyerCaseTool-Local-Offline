# -*- coding: utf-8 -*-
"""创建 Word 模板文件的脚本"""

from pathlib import Path
import sys

# 添加项目根目录到 Python 路径
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))


def create_template(template_path: Path, title: str, variables: list) -> bool:
    """
    创建 Word 模板文件

    Args:
        template_path: 模板文件路径
        title: 模板标题
        variables: 变量列表

    Returns:
        是否成功
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("请先安装 python-docx: pip install python-docx")
        return False

    # 创建文档
    doc = Document()

    # 设置标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 添加变量说明
    info_para = doc.add_paragraph()
    info_run = info_para.add_run("（以下为示例变量，请根据实际情况填写）")
    info_run.font.size = Pt(10)
    info_run.font.italic = True

    doc.add_paragraph()

    # 添加变量表格
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # 表头
    header_cells = table.rows[0].cells
    header_cells[0].text = "项目"
    header_cells[1].text = "内容"

    # 添加变量行
    for var in variables:
        row_cells = table.add_row().cells
        row_cells[0].text = var.get("label", var.get("key", ""))
        row_cells[1].text = "{{" + var.get("key", "") + "}}"

    doc.add_paragraph()

    # 添加说明
    note_para = doc.add_paragraph()
    note_run = note_para.add_run(
        "说明：此模板由案件文件夹管理系统创建，"
        "包含可替换的变量（格式为 {{变量名}}）。"
    )
    note_run.font.size = Pt(9)

    # 保存
    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(template_path))

    return True


def create_all_templates():
    """创建所有默认模板"""
    from src.config.default_templates import (
        CIVIL_TEMPLATE,
        CRIMINAL_TEMPLATE,
        NON_LITIGATION_TEMPLATE
    )

    project_root = Path(__file__).parent.parent

    templates = [
        (CIVIL_TEMPLATE, "templates/civil/template.docx"),
        (CRIMINAL_TEMPLATE, "templates/criminal/template.docx"),
        (NON_LITIGATION_TEMPLATE, "templates/non_litigation/template.docx"),
    ]

    print("开始创建模板文件...")

    for template_config, relative_path in templates:
        template_path = project_root / relative_path
        title = template_config.get("name", "")
        variables = template_config.get("variables", [])

        if create_template(template_path, title, variables):
            print(f"✓ 创建成功: {relative_path}")
        else:
            print(f"✗ 创建失败: {relative_path}")

    print("\n模板创建完成!")


if __name__ == "__main__":
    create_all_templates()
