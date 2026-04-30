# -*- coding: utf-8 -*-
"""自动排版引擎测试用例"""

import os
import sys
import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 添加项目路径
sys.path.insert(0, str(Path(__file__).parent.parent))

from src.core.docx_auto_format import (
    DocxAutoFormatter,
    ParagraphType,
    _get_effective_font_size,
    _is_para_bold,
    _is_para_centered,
    _text_looks_like_body,
    clean_text,
    clean_document,
)


# ═══════════════════════════════════════════════════════════════════
# 辅助函数：创建测试文档
# ═══════════════════════════════════════════════════════════════════

def _make_test_doc() -> Document:
    """创建一个典型法律文书结构。"""
    doc = Document()

    # 设置默认样式
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(14)

    # 1. 大标题（居中，稍大字，加粗）
    title = doc.add_paragraph("民事起诉状")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(18)
        run.font.bold = True

    # 2. 空行
    doc.add_paragraph("")

    # 3. 当事人信息 —— 长段落（以"原告："开头但实际是正文）
    plaintiff = doc.add_paragraph(
        "原告：张三，男，1990年1月1日出生，汉族，住北京市朝阳区某某路100号，"
        "身份证号110101199001011234，联系电话13800000001。"
    )
    for run in plaintiff.runs:
        run.font.size = Pt(14)

    # 4. 被告信息 —— 长段落（类似上面）
    defendant = doc.add_paragraph(
        "被告：李四，男，1988年5月5日出生，汉族，住北京市海淀区某某街200号，"
        "身份证号110102198805051234，联系电话13800000002。"
    )
    for run in defendant.runs:
        run.font.size = Pt(14)

    # 5. 空行
    doc.add_paragraph("")

    # 6. 诉讼请求 —— 短标题（真正的标题）
    claim_title = doc.add_paragraph("诉讼请求：")
    for run in claim_title.runs:
        run.font.size = Pt(16)
        run.font.bold = True

    # 7-8. 列举请求 —— 正文（以"一、"开头但结尾为"；"）
    claim1 = doc.add_paragraph("一、判令被告支付拖欠货款人民币100,000元及逾期利息；")
    for run in claim1.runs:
        run.font.size = Pt(14)

    claim2 = doc.add_paragraph("二、判令被告承担本案全部诉讼费用。")
    for run in claim2.runs:
        run.font.size = Pt(14)

    # 9. 空行
    doc.add_paragraph("")

    # 10. 事实与理由 —— 短标题
    reason_title = doc.add_paragraph("事实与理由：")
    for run in reason_title.runs:
        run.font.size = Pt(16)
        run.font.bold = True

    # 11. 正文段落
    body1 = doc.add_paragraph(
        "原告与被告于2023年6月签订《货物买卖合同》，约定原告向被告供应建材，"
        "被告应于收货后30日内支付货款。原告已按约履行供货义务，但被告至今"
        "未支付货款人民币100,000元，已构成违约。"
    )
    for run in body1.runs:
        run.font.size = Pt(14)

    # 12. 正文段落
    body2 = doc.add_paragraph(
        "原告多次催告被告履行付款义务，被告均以资金周转困难为由推脱，"
        "严重侵害了原告的合法权益。根据《中华人民共和国民法典》第五百七十七条之规定，"
        "被告应当承担继续履行、赔偿损失等违约责任。"
    )
    for run in body2.runs:
        run.font.size = Pt(14)

    return doc


# ═══════════════════════════════════════════════════════════════════
# 单元测试
# ═══════════════════════════════════════════════════════════════════

class TestHelperFunctions:
    """辅助函数测试"""

    def test_text_looks_like_body_long_with_punctuation(self):
        """长文本 + 句末标点 → 正文"""
        assert _text_looks_like_body("原告：张三，男，1990年出生，住北京市朝阳区。") is True

    def test_text_looks_like_body_long_with_semicolon(self):
        """长文本 + 分号结尾 → 正文"""
        assert _text_looks_like_body("一、判令被告支付货款100,000元及利息；") is True

    def test_text_looks_like_body_short_title(self):
        """短文本 → 非正文"""
        assert _text_looks_like_body("诉讼请求：") is False

    def test_text_looks_like_body_short_numbered(self):
        """短编号 → 非正文"""
        assert _text_looks_like_body("一、诉讼请求") is False


class TestParagraphClassification:
    """段落分类测试"""

    def test_main_title_detected(self):
        """大标题正确识别"""
        doc = _make_test_doc()
        formatter = DocxAutoFormatter()

        # 手动注入分析（因为load需要文件路径）
        formatter._doc = doc
        formatter._analyze()

        types = formatter.get_paragraph_types()
        # 第0段："民事起诉状" → MAIN_TITLE
        assert types[0] == ParagraphType.MAIN_TITLE, f"期望 MAIN_TITLE，实际 {types[0]}"

    def test_long_plaintiff_line_is_body(self):
        """长"原告："段落应被归类为正文（非标题）"""
        doc = _make_test_doc()
        formatter = DocxAutoFormatter()
        formatter._doc = doc
        formatter._analyze()

        types = formatter.get_paragraph_types()
        # 第2段：长"原告：" → BODY
        assert types[2] == ParagraphType.BODY, (
            f"长当事人信息行应为 BODY，实际 {types[2]}。"
            f"文本：'{doc.paragraphs[2].text[:50]}…'"
        )

    def test_long_defendant_line_is_body(self):
        """长"被告："段落应被归类为正文（非标题）"""
        doc = _make_test_doc()
        formatter = DocxAutoFormatter()
        formatter._doc = doc
        formatter._analyze()

        types = formatter.get_paragraph_types()
        # 第3段：长"被告：" → BODY
        assert types[3] == ParagraphType.BODY, (
            f"长当事人信息行应为 BODY，实际 {types[3]}。"
            f"文本：'{doc.paragraphs[3].text[:50]}…'"
        )

    def test_claim_title_is_level1(self):
        """"诉讼请求："应识别为一级标题"""
        doc = _make_test_doc()
        formatter = DocxAutoFormatter()
        formatter._doc = doc
        formatter._analyze()

        types = formatter.get_paragraph_types()
        # 第5段：诉讼请求 → LEVEL1_TITLE
        assert types[5] == ParagraphType.LEVEL1_TITLE, (
            f"期望 LEVEL1_TITLE，实际 {types[5]}"
        )

    def test_numbered_claims_are_body(self):
        """"一、判令被告…；"应识别为正文（非标题）"""
        doc = _make_test_doc()
        formatter = DocxAutoFormatter()
        formatter._doc = doc
        formatter._analyze()

        types = formatter.get_paragraph_types()
        assert types[6] == ParagraphType.BODY, (
            f"编号列举项应为 BODY，实际 {types[6]}"
        )
        assert types[7] == ParagraphType.BODY, (
            f"编号列举项应为 BODY，实际 {types[7]}"
        )

    def test_reason_title_is_level1(self):
        """"事实与理由："应识别为一级标题"""
        doc = _make_test_doc()
        formatter = DocxAutoFormatter()
        formatter._doc = doc
        formatter._analyze()

        types = formatter.get_paragraph_types()
        assert types[9] == ParagraphType.LEVEL1_TITLE, (
            f"期望 LEVEL1_TITLE，实际 {types[9]}"
        )

    def test_body_paragraphs_are_body(self):
        """正文段落应识别为正文"""
        doc = _make_test_doc()
        formatter = DocxAutoFormatter()
        formatter._doc = doc
        formatter._analyze()

        types = formatter.get_paragraph_types()
        assert types[10] == ParagraphType.BODY
        assert types[11] == ParagraphType.BODY

    def test_empty_paragraphs_are_empty(self):
        """空段落应识别为空"""
        doc = _make_test_doc()
        formatter = DocxAutoFormatter()
        formatter._doc = doc
        formatter._analyze()

        types = formatter.get_paragraph_types()
        assert types[1] == ParagraphType.EMPTY
        assert types[4] == ParagraphType.EMPTY
        assert types[8] == ParagraphType.EMPTY


class TestFormatApplication:
    """格式应用测试"""

    def test_format_applied_to_all_paragraphs(self):
        """所有段落都被格式化了"""
        doc = _make_test_doc()
        formatter = DocxAutoFormatter()
        formatter._doc = doc

        # 手动设置分类结果
        formatter._paragraph_types = [
            ParagraphType.MAIN_TITLE,   # 0: 民事起诉状
            ParagraphType.EMPTY,        # 1: 空
            ParagraphType.BODY,         # 2: 原告长段落
            ParagraphType.BODY,         # 3: 被告长段落
            ParagraphType.EMPTY,        # 4: 空
            ParagraphType.LEVEL1_TITLE, # 5: 诉讼请求
            ParagraphType.BODY,         # 6: 一、
            ParagraphType.BODY,         # 7: 二、
            ParagraphType.EMPTY,        # 8: 空
            ParagraphType.LEVEL1_TITLE, # 9: 事实与理由
            ParagraphType.BODY,         # 10: 正文
            ParagraphType.BODY,         # 11: 正文
        ]
        formatter._stats = {"most_common_size": 14, "max_size": 18, "min_size": 14,
                            "bold_ratio": 0.25, "total_paragraphs": 12}

        formatter.apply_format()

        # 大标题：居中、方正小标宋、22pt、不加粗
        p0 = doc.paragraphs[0]
        assert p0.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert p0.runs[0].font.size == Pt(22)
        assert p0.runs[0].font.bold is False

        # 正文：两端对齐、仿宋、16pt、首行缩进
        p2 = doc.paragraphs[2]
        assert p2.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        assert p2.runs[0].font.size == Pt(16)
        # 首行缩进 Pt(32) = 2字符（允许 EMU 取整误差）
        indent_emu = p2.paragraph_format.first_line_indent
        assert indent_emu == Pt(32), f"缩进值 {indent_emu} EMU，期望 {Pt(32)} EMU (Pt(32))"

        # 一级标题：左对齐、黑体、16pt、加粗
        p5 = doc.paragraphs[5]
        assert p5.runs[0].font.bold is True

    def test_paragraph_style_detached(self):
        """段落样式引用被清除，格式直接生效"""
        doc = _make_test_doc()
        formatter = DocxAutoFormatter()
        formatter._doc = doc
        formatter._paragraph_types = [ParagraphType.BODY] * len(doc.paragraphs)
        formatter._stats = {"most_common_size": 14, "max_size": 18, "min_size": 14,
                            "bold_ratio": 0.25, "total_paragraphs": 12}
        formatter.apply_format()

        # 检查第一个非空段落是否有样式引用
        p0 = doc.paragraphs[0]
        pPr = p0._element.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr"
        )
        if pPr is not None:
            pStyle = pPr.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle"
            )
            assert pStyle is None, "段落样式引用应该被移除"

    def test_font_names_set_correctly(self):
        """字体名称在所有 run 上正确设置"""
        doc = _make_test_doc()
        formatter = DocxAutoFormatter()
        formatter._doc = doc
        formatter._paragraph_types = [ParagraphType.BODY] * len(doc.paragraphs)
        formatter._stats = {"most_common_size": 14, "max_size": 18, "min_size": 14,
                            "bold_ratio": 0.25, "total_paragraphs": 12}
        formatter.apply_format()

        p2 = doc.paragraphs[2]  # 原来有内容的段落
        if p2.runs:
            rpr = p2.runs[0]._element.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
            )
            if rpr is not None:
                rfonts = rpr.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
                )
                if rfonts is not None:
                    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                    east_asia = rfonts.get(f"{{{ns}}}eastAsia")
                    ascii_f = rfonts.get(f"{{{ns}}}ascii")
                    assert east_asia == "仿宋_GB2312", f"东亚字体应为仿宋_GB2312，实际 {east_asia}"
                    assert ascii_f == "仿宋_GB2312", f"ASCII字体应为仿宋_GB2312, 实际 {ascii_f}"


class TestEndToEnd:
    """端到端测试"""

    def test_full_workflow(self):
        """完整的加载→分析→排版→保存流程"""
        doc = _make_test_doc()

        # 保存到临时文件
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = Path(f.name)
            doc.save(str(tmp_path))

        try:
            # 加载
            formatter = DocxAutoFormatter()
            formatter.load(tmp_path)

            # 验证分类
            types = formatter.get_paragraph_types()
            # 大标题
            assert types[0] == ParagraphType.MAIN_TITLE
            # 长当事人信息 → BODY（关键断言！）
            assert types[2] == ParagraphType.BODY
            assert types[3] == ParagraphType.BODY
            # 标题
            assert types[5] == ParagraphType.LEVEL1_TITLE
            # 编号项 → BODY
            assert types[6] == ParagraphType.BODY
            assert types[7] == ParagraphType.BODY
            # 标题
            assert types[9] == ParagraphType.LEVEL1_TITLE

            # 排版
            formatter.apply_format()

            # 保存
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f_out:
                out_path = Path(f_out.name)
            formatter.save(out_path)

            # 验证输出文件存在且可读
            assert out_path.exists()
            out_doc = Document(str(out_path))
            assert len(out_doc.paragraphs) == len(doc.paragraphs)

            # 清理
            out_path.unlink()

        finally:
            tmp_path.unlink()

    def test_backup_created(self):
        """测试备份文件创建"""
        doc = _make_test_doc()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = Path(f.name)
            doc.save(str(tmp_path))

        try:
            backup_path = DocxAutoFormatter.backup_original(tmp_path)
            assert backup_path.exists()
            assert "_备份" in backup_path.name
            backup_path.unlink()
        finally:
            tmp_path.unlink()


class TestEdgeCases:
    """边界情况测试"""

    def test_short_plaintiff_is_title(self):
        """短"原告："应是标题"""
        doc = Document()
        p = doc.add_paragraph("原告：张三")
        for run in p.runs:
            run.font.size = Pt(14)
            run.font.bold = True

        formatter = DocxAutoFormatter()
        formatter._doc = doc
        formatter._analyze()

        types = formatter.get_paragraph_types()
        assert types[0] == ParagraphType.LEVEL1_TITLE

    def test_centered_bold_short_is_title(self):
        """居中+加粗+短文本→标题"""
        doc = Document()
        p = doc.add_paragraph("证据目录")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(16)
            run.font.bold = True

        formatter = DocxAutoFormatter()
        formatter._doc = doc
        formatter._analyze()

        types = formatter.get_paragraph_types()
        assert types[0] == ParagraphType.MAIN_TITLE

    def test_court_appeal_text_is_body(self):
        """"如不服本判决…"长文本应识别为正文"""
        doc = Document()
        p = doc.add_paragraph(
            "如不服本判决，可在判决书送达之日起十五日内，向本院递交上诉状，"
            "并按对方当事人的人数提出副本，上诉于北京市第二中级人民法院。"
        )
        for run in p.runs:
            run.font.size = Pt(14)

        formatter = DocxAutoFormatter()
        formatter._doc = doc
        formatter._analyze()

        types = formatter.get_paragraph_types()
        assert types[0] == ParagraphType.BODY, (
            f"长'如不服本判决'应为正文，实际 {types[0]}"
        )

    def test_pure_date_is_signature(self):
        """纯日期行→SIGNATURE（落款日期，不应当作标题）"""
        doc = Document()
        p = doc.add_paragraph("2024年1月15日")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(14)

        formatter = DocxAutoFormatter()
        formatter._doc = doc
        formatter._analyze()

        types = formatter.get_paragraph_types()
        assert types[0] == ParagraphType.SIGNATURE


class TestTextCleaning:
    """网页/AI粘贴垃圾字符清洗测试"""

    def test_zero_width_space_removed(self):
        assert clean_text("民事\u200B起诉状") == "民事起诉状"

    def test_no_break_space_to_space(self):
        assert clean_text("原告：\u00A0张三") == "原告： 张三"

    def test_ideographic_space_to_space(self):
        assert clean_text("被告\u3000李四") == "被告 李四"

    def test_direction_control_removed(self):
        assert clean_text("\u202D诉讼\u202C请求\u200E") == "诉讼请求"

    def test_bom_removed(self):
        assert clean_text("\uFEFF事实与理由") == "事实与理由"

    def test_multiple_spaces_merged(self):
        assert clean_text("原告：  张三，  男") == "原告： 张三， 男"

    def test_normal_text_unchanged(self):
        text = "根据《中华人民共和国民法典》第五百七十七条规定，被告应当承担违约责任。"
        assert clean_text(text) == text

    def test_chinese_punctuation_preserved(self):
        """中文标点（含全角分号）不应被清除"""
        text = "合同约定：（一）支付货款100,000元；"
        assert clean_text(text) == text

    def test_smart_quotes_preserved(self):
        """智能引号不应被清除"""
        text = "\u201C民事起诉状\u201D"
        assert clean_text(text) == text

    def test_mixed_garbage_cleaned(self):
        result = clean_text("\uFEFF\u202D\u200B原告\u3000\u3000\u00A0：\u200C\u200E张三\u202C")
        assert "\u200B" not in result
        assert "\uFEFF" not in result
        assert "\u3000" not in result
        assert "\u202D" not in result
        assert "原告" in result
        assert "张三" in result

    def test_md_heading_prefix_removed(self):
        """# 开头的 markdown 标题前缀应被移除"""
        assert clean_text("# 标题") == "标题"
        assert clean_text("## 二级标题") == "二级标题"
        assert clean_text("### 三级标题") == "三级标题"

    def test_double_hash_deleted_anywhere(self):
        """连续2个及以上#无论在哪都删除"""
        assert clean_text("## 标题") == "标题"
        assert clean_text("正##文") == "正文"
        assert clean_text("### 标题") == "标题"
        assert clean_text("##---##") == ""  # --- also matched by HR regex
        assert clean_text("# 单个标题") == "单个标题"

    def test_single_hash_in_middle_kept(self):
        """文本中间的单个#不应被删除（如'#5'等合法用法）"""
        assert "#" in clean_text("第#5号文件")

    def test_md_bold_stripped(self):
        assert clean_text("**加粗**") == "加粗"
        assert clean_text("正常**加粗部分**正常") == "正常加粗部分正常"

    def test_md_italic_stripped(self):
        assert clean_text("*斜体*") == "斜体"

    def test_md_horizontal_rule_cleared(self):
        assert clean_text("***") == ""
        assert clean_text("---") == ""

    def test_md_list_prefix_removed(self):
        assert clean_text("- 列表项") == "列表项"
        assert clean_text("+ 列表项") == "列表项"

    def test_md_link_stripped(self):
        assert clean_text("[链接文字](http://example.com)") == "链接文字"

    def test_document_level_cleaning(self):
        """文档级清洗正确修改 run 文本"""
        doc = Document()
        doc.add_paragraph("\u200B民事\uFEFF起诉状\u200B")
        doc.add_paragraph("正常文本")
        changed = clean_document(doc)
        assert doc.paragraphs[0].text == "民事起诉状"
        assert doc.paragraphs[1].text == "正常文本"
        assert changed >= 1

    def test_cleaning_in_full_workflow(self):
        """端到端：带垃圾字符的文档 → 清洗 → 分类 → 排版"""
        doc = Document()
        p = doc.add_paragraph("\u200B民事起诉状\u200B")
        for r in p.runs:
            r.font.size = Pt(22)
        doc.add_paragraph("原告\u3000：\u00A0张三，男，1990年出生。")
        doc.add_paragraph("诉讼请求")
        doc.add_paragraph("1、判令被告支付货款；")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp = Path(f.name)
        doc.save(str(tmp))

        try:
            formatter = DocxAutoFormatter()
            formatter.load(tmp)

            # 垃圾字符已清除
            assert "\u200B" not in formatter._doc.paragraphs[0].text
            assert "\u3000" not in formatter._doc.paragraphs[1].text

            # 分类仍正确
            types = formatter.get_paragraph_types()
            assert types[0] == ParagraphType.MAIN_TITLE

            formatter.apply_format()

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f2:
                out = Path(f2.name)
            formatter.save(out)

            # 排版后也无垃圾字符
            check_doc = Document(str(out))
            for para in check_doc.paragraphs:
                assert "\u200B" not in para.text
                assert "\uFEFF" not in para.text

            out.unlink()
        finally:
            tmp.unlink()


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
