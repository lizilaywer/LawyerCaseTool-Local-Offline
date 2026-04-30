# -*- coding: utf-8 -*-
"""真实场景自动排版诊断测试

模拟用户实际使用的 Word 文档（从其他来源复制粘贴而来），
检验分类和格式应用是否正确。
"""

import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.core.docx_auto_format import DocxAutoFormatter, ParagraphType


def _make_realistic_doc() -> Document:
    """创建一个更贴近真实法律文书的文档。

    特点：
    - 使用 Word 默认样式（不特别设字体/字号），模拟用户直接粘贴
    - 大标题"民事起诉状"居中但未刻意加大字号
    - 段落标题如"诉讼请求"不一定是加粗
    - 正文混有不同格式来源的内容
    """
    doc = Document()

    # 不修改默认样式 —— 模拟"从其他地方粘贴进来"的场景
    # Word 默认样式通常是 Calibri 11pt 或 宋体 五号(10.5pt)

    # 1. 大标题
    p = doc.add_paragraph("民事起诉状")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 用户可能手动加大了字号
    for r in p.runs:
        r.font.size = Pt(22)

    # 2. 空行
    doc.add_paragraph("")

    # 3. 原告信息（长行）
    doc.add_paragraph(
        "原告：张三，男，1990年1月1日出生，汉族，住北京市朝阳区某某路100号。"
    )

    # 4. 被告信息（长行）
    doc.add_paragraph(
        "被告：李四，男，1988年5月5日出生，汉族，住北京市海淀区某某街200号。"
    )

    # 5. 空行
    doc.add_paragraph("")

    # 6. 诉讼请求
    p = doc.add_paragraph("诉讼请求")
    # 注意：用户可能没加粗，只是单独一行短文本

    # 7-9. 请求列举
    doc.add_paragraph("1、判令被告支付拖欠货款人民币100,000元及逾期利息；")
    doc.add_paragraph("2、判令被告承担本案全部诉讼费用。")
    doc.add_paragraph("")

    # 10. 事实与理由
    doc.add_paragraph("事实与理由")

    # 11-12. 正文
    doc.add_paragraph(
        "原告与被告于2023年6月签订《货物买卖合同》，约定原告向被告供应建材，"
        "被告应于收货后30日内支付货款。原告已按约履行供货义务，但被告至今未支付货款。"
    )
    doc.add_paragraph(
        "原告多次催告被告履行付款义务，被告均以资金周转困难为由推脱。"
        "根据《中华人民共和国民法典》第五百七十七条之规定，被告应当承担违约责任。"
    )

    # 13. 此致
    doc.add_paragraph("此致")
    doc.add_paragraph("北京市朝阳区人民法院")

    # 15. 具状人 + 日期
    p = doc.add_paragraph("具状人：张三")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph("2024年1月15日")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return doc


def _make_docx_no_explicit_formatting() -> Document:
    """模拟一个完全没有任何格式标记的纯文本粘贴文档。"""
    doc = Document()

    # 全部不设置任何格式，模拟纯粘贴
    doc.add_paragraph("民事起诉状")  # 应该是大标题，但可能没居中
    doc.add_paragraph("")
    doc.add_paragraph("原告：张三，男，1990年出生，住北京市朝阳区某某路100号。")
    doc.add_paragraph("被告：李四，男，1988年出生，住北京市海淀区某某街200号。")
    doc.add_paragraph("")
    doc.add_paragraph("诉讼请求")
    doc.add_paragraph("1、判令被告支付货款100,000元；")
    doc.add_paragraph("2、判令被告承担诉讼费用。")
    doc.add_paragraph("")
    doc.add_paragraph("事实与理由")
    doc.add_paragraph("原告与被告签订合同后被告未付款。")
    doc.add_paragraph("此致")
    doc.add_paragraph("北京市朝阳区人民法院")
    doc.add_paragraph("具状人：张三")
    doc.add_paragraph("2024年1月15日")

    return doc


def _make_docx_with_styles() -> Document:
    """模拟使用 Word 内置标题样式的文档。"""
    doc = Document()

    # 用 Word 的标题样式
    doc.add_heading("民事起诉状", level=0)  # Title 样式
    doc.add_paragraph("")
    doc.add_paragraph("原告：张三，男，1990年出生，住北京市朝阳区。")
    doc.add_paragraph("被告：李四，男，1988年出生，住北京市海淀区。")
    doc.add_paragraph("")
    doc.add_heading("诉讼请求", level=1)
    doc.add_paragraph("1、判令被告支付货款100,000元；")
    doc.add_paragraph("2、判令被告承担诉讼费用。")
    doc.add_paragraph("")
    doc.add_heading("事实与理由", level=1)
    doc.add_paragraph("原告与被告签订合同后被告未付款，构成违约。")

    return doc


def _classify_and_print(formatter, doc, label=""):
    """辅助：打印分类结果。"""
    types = formatter.get_paragraph_types()
    print(f"\n{'='*60}")
    print(f"  {label}")
    print(f"{'='*60}")
    for i, (ptype, para) in enumerate(zip(types, doc.paragraphs)):
        text = para.text.strip()[:50] or "(空)"
        align = ""
        if para.alignment is not None:
            align = f" [align={para.alignment}]"
        bold = ""
        if any(r.bold for r in para.runs if r.bold):
            bold = " [bold]"
        font_sz = ""
        for r in para.runs:
            if r.font.size:
                font_sz = f" [size={r.font.size.pt}pt]"
                break
        print(f"  [{i:2d}] {ptype:15s} | {text}{align}{bold}{font_sz}")
    return types


# ═══════════════════════════════════════════════════════════════════
# 诊断测试
# ═══════════════════════════════════════════════════════════════════

def test_realistic_document_classification():
    """测试贴近真实的文档分类。"""
    doc = _make_realistic_doc()
    formatter = DocxAutoFormatter()
    formatter._doc = doc
    formatter._analyze()

    types = _classify_and_print(formatter, doc, "真实文书（手动设置部分格式）")

    # 0: "民事起诉状" 居中 + 22pt → 必须是 MAIN_TITLE
    assert types[0] == ParagraphType.MAIN_TITLE, (
        f"'民事起诉状' 应为 MAIN_TITLE，实际 {types[0]}"
    )

    # 5: "诉讼请求" 短文本 → LEVEL1_TITLE
    assert types[5] == ParagraphType.LEVEL1_TITLE, (
        f"'诉讼请求' 应为 LEVEL1_TITLE，实际 {types[5]}"
    )

    # 9: "事实与理由" → LEVEL1_TITLE
    assert types[9] == ParagraphType.LEVEL1_TITLE, (
        f"'事实与理由' 应为 LEVEL1_TITLE，实际 {types[9]}"
    )


def test_no_formatting_document_classification():
    """测试完全没有格式标记的文档分类。"""
    doc = _make_docx_no_explicit_formatting()
    formatter = DocxAutoFormatter()
    formatter._doc = doc
    formatter._analyze()

    types = _classify_and_print(formatter, doc, "纯文本粘贴（无任何格式）")

    # "民事起诉状" → 至少 LEVEL1_TITLE，理想情况是 MAIN_TITLE
    assert types[0] in (ParagraphType.MAIN_TITLE, ParagraphType.LEVEL1_TITLE), (
        f"'民事起诉状' 应为标题，实际 {types[0]}"
    )

    # "诉讼请求" → LEVEL1_TITLE
    assert types[5] == ParagraphType.LEVEL1_TITLE, (
        f"'诉讼请求' 应为 LEVEL1_TITLE，实际 {types[5]}"
    )

    # "事实与理由" → LEVEL1_TITLE
    assert types[9] == ParagraphType.LEVEL1_TITLE, (
        f"'事实与理由' 应为 LEVEL1_TITLE，实际 {types[9]}"
    )


def test_styled_document_classification():
    """测试使用 Word 标题样式的文档。"""
    doc = _make_docx_with_styles()
    formatter = DocxAutoFormatter()
    formatter._doc = doc
    formatter._analyze()

    types = _classify_and_print(formatter, doc, "Word 标题样式文档")

    # "民事起诉状" (heading level 0) → MAIN_TITLE
    assert types[0] == ParagraphType.MAIN_TITLE, (
        f"Title 样式的'民事起诉状' 应为 MAIN_TITLE，实际 {types[0]}"
    )

    # "诉讼请求" (heading level 1) → LEVEL1_TITLE
    assert types[5] == ParagraphType.LEVEL1_TITLE, (
        f"Heading1 样式的'诉讼请求' 应为 LEVEL1_TITLE，实际 {types[5]}"
    )

    # "事实与理由" (heading level 1) → LEVEL1_TITLE
    assert types[9] == ParagraphType.LEVEL1_TITLE, (
        f"Heading1 样式的'事实与理由' 应为 LEVEL1_TITLE，实际 {types[9]}"
    )

    # 正文不应被误判为标题
    assert types[10] == ParagraphType.BODY, (
        f"正文应为 BODY，实际 {types[10]}"
    )


def test_format_actually_changes_fonts():
    """测试排版后字体是否真正改变。"""
    doc = _make_realistic_doc()
    formatter = DocxAutoFormatter()
    formatter._doc = doc
    formatter._analyze()

    types = formatter.get_paragraph_types()
    formatter.apply_format()

    print("\n--- 排版后字体检查 ---")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()[:40] or "(空)"
        if para.runs:
            run = para.runs[0]
            font_name = run.font.name or "(未设置)"
            font_size = f"{run.font.size.pt}pt" if run.font.size else "(未设置)"
            bold = run.font.bold
            print(f"  [{i:2d}] {types[i]:15s} | font={font_name} size={font_size} bold={bold} | {text}")

    # 大标题：方正小标宋简体 22pt
    p0 = doc.paragraphs[0]
    assert p0.runs[0].font.name == "方正小标宋简体", (
        f"大标题字体应为'方正小标宋简体'，实际'{p0.runs[0].font.name}'"
    )
    assert p0.runs[0].font.size == Pt(22)

    # 一级标题（诉讼请求）：黑体
    # 找到诉讼请求段落
    claim_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == "诉讼请求":
            claim_idx = i
            break
    assert claim_idx is not None, "未找到'诉讼请求'段落"
    p_claim = doc.paragraphs[claim_idx]
    assert p_claim.runs[0].font.name == "黑体", (
        f"一级标题字体应为'黑体'，实际'{p_claim.runs[0].font.name}'"
    )

    # 正文：仿宋_GB2312
    body_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "签订" in para.text:
            body_idx = i
            break
    assert body_idx is not None, "未找到正文段落"
    p_body = doc.paragraphs[body_idx]
    assert p_body.runs[0].font.name == "仿宋_GB2312", (
        f"正文字体应为'仿宋_GB2312'，实际'{p_body.runs[0].font.name}'"
    )


def test_end_to_end_realistic():
    """端到端：创建 → 保存 → 加载 → 分析 → 排版 → 保存 → 验证。"""
    doc = _make_realistic_doc()

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        tmp = Path(f.name)
        doc.save(str(tmp))

    try:
        formatter = DocxAutoFormatter()
        formatter.load(tmp)

        types = formatter.get_paragraph_types()
        print(f"\n--- 端到端分类结果 ---")
        for i, (t, p) in enumerate(zip(types, formatter._doc.paragraphs)):
            print(f"  [{i}] {t:15s} | {p.text.strip()[:50]}")

        formatter.apply_format()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f2:
            out = Path(f2.name)
        formatter.save(out)

        # 重新读取验证
        check_doc = Document(str(out))

        # 大标题
        p0 = check_doc.paragraphs[0]
        assert p0.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert p0.runs[0].font.name == "方正小标宋简体"
        assert p0.runs[0].font.size == Pt(22)

        print("\n--- 端到端排版验证 ---")
        for i, para in enumerate(check_doc.paragraphs):
            text = para.text.strip()[:40] or "(空)"
            if para.runs:
                fn = para.runs[0].font.name or "?"
                fs = f"{para.runs[0].font.size.pt}pt" if para.runs[0].font.size else "?"
                print(f"  [{i}] fn={fn} fs={fs} | {text}")

        out.unlink()
    finally:
        tmp.unlink()


if __name__ == "__main__":
    # 逐个运行诊断测试，打印详细信息
    print("开始诊断测试...")

    for name, fn in [
        ("真实文书分类", test_realistic_document_classification),
        ("纯文本粘贴分类", test_no_formatting_document_classification),
        ("Word样式文档分类", test_styled_document_classification),
        ("字体修改验证", test_format_actually_changes_fonts),
        ("端到端测试", test_end_to_end_realistic),
    ]:
        print(f"\n{'#'*60}")
        print(f"# {name}")
        print(f"{'#'*60}")
        try:
            fn()
            print(f"✓ {name} 通过")
        except AssertionError as e:
            print(f"✗ {name} 失败: {e}")
        except Exception as e:
            print(f"✗ {name} 异常: {type(e).__name__}: {e}")
