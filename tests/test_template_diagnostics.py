# -*- coding: utf-8 -*-
"""模板诊断测试"""

from pathlib import Path

from docx import Document

from src.core.template_diagnostics import diagnose_templates


class _FakeTemplatePathManager:
    def __init__(self, mapping):
        self._mapping = mapping

    def resolve_template_path(self, relative_path: str):
        return self._mapping.get(relative_path)


def test_diagnose_templates_reports_invalid_paths_and_missing_variables(tmp_path):
    """模板诊断应识别失效路径、重复 ID 和未定义占位符。"""
    docx_path = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("委托人：{{client_name}}")
    document.add_paragraph("案号：{{matter_number}}")
    document.save(str(docx_path))

    templates = [
        {
            "id": "civil",
            "name": "民事模板",
            "category": "civil",
            "template_file": "template.docx",
            "variables": [{"key": "client_name"}],
            "folder_structure": {"folders": []},
        },
        {
            "id": "civil",
            "name": "重复模板",
            "category": "civil",
            "template_file": "missing.docx",
            "variables": [],
            "folder_structure": {"folders": []},
        },
    ]
    manager = _FakeTemplatePathManager({"template.docx": docx_path})

    summary = diagnose_templates(templates, manager)

    assert summary.template_count == 2
    assert summary.invalid_paths == 1
    assert summary.duplicate_ids == 2
    assert summary.missing_variables == 1
    assert summary.placeholder_count == 2
    assert any("模板文件不存在" in issue.message for issue in summary.issues)
    assert any("matter_number" in issue.message for issue in summary.issues)
