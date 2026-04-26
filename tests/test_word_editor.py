# -*- coding: utf-8 -*-
"""Word 编辑器测试"""

import shutil
import tempfile
from pathlib import Path
from typing import List

from docx import Document

from src.core.batch_processor import BatchProcessor
from src.core.word_editor import WordEditor


class TestWordEditor:
    """Word 编辑器测试类"""

    def setup_method(self):
        """每个测试方法前执行"""
        self.temp_dir = Path(tempfile.mkdtemp())
        self.editor = WordEditor()

    def teardown_method(self):
        """每个测试方法后执行"""
        self.editor.close()
        if self.temp_dir.exists():
            shutil.rmtree(self.temp_dir)

    def _create_document(self, name: str, paragraphs: List[str]) -> Path:
        """创建测试用 Word 文档。"""
        path = self.temp_dir / name
        document = Document()
        for text in paragraphs:
            document.add_paragraph(text)
        document.save(path)
        return path

    def test_undo_only_reverts_last_replacement(self):
        """测试撤销只回退最后一步替换"""
        path = self._create_document("undo.docx", ["张三 与 李四"])

        assert self.editor.load_document(path) is True
        assert self.editor.replace_text("张三", "{{client_name}}", False)[0] == 1
        assert self.editor.replace_text("李四", "{{opponent_name}}", False)[0] == 1

        assert self.editor.undo() is True
        text = self.editor.extract_all_text()
        assert "{{client_name}}" in text
        assert "{{opponent_name}}" not in text
        assert "李四" in text

    def test_undo_variable_restores_distinct_original_values(self):
        """测试撤销变量时恢复各自原始文本"""
        path = self._create_document("undo_variable.docx", ["申请人：张三", "被申请人：李四"])

        assert self.editor.load_document(path) is True
        assert self.editor.replace_text("张三", "{{client_name}}", False)[0] == 1
        assert self.editor.replace_text("李四", "{{client_name}}", False)[0] == 1

        assert self.editor.undo_variable("client_name", single=False) == 2
        text = self.editor.extract_all_text()
        assert "{{client_name}}" not in text
        assert "张三" in text
        assert "李四" in text

    def test_process_single_respects_cancel_before_start(self):
        """测试单案卷生成在开始前检查取消状态"""
        processor = BatchProcessor()
        processor.cancel()

        result = processor.process_single(
            {
                "variables": [],
                "folder_structure": {
                    "root_name": "测试案卷",
                    "folders": []
                }
            },
            {},
            self.temp_dir
        )

        assert result["success"] is False
        assert result["cancelled"] is True
        assert result["error"] == "已取消"

    def test_save_as_updates_current_path_for_followup_save(self):
        """另存为后，后续保存应继续写入新文件。"""
        source_path = self._create_document("source.docx", ["原始内容"])
        copied_path = self.temp_dir / "copied.docx"

        assert self.editor.load_document(source_path) is True
        assert self.editor.save_as(copied_path) is True
        assert self.editor.get_current_path() == copied_path

        assert self.editor.replace_text("原始内容", "另存后内容", False)[0] == 1
        assert self.editor.save_document() is True

        copied_text = "\n".join(p.text for p in Document(copied_path).paragraphs)
        source_text = "\n".join(p.text for p in Document(source_path).paragraphs)
        assert copied_text == "另存后内容"
        assert source_text == "原始内容"
